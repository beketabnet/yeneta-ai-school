"""
Utility functions for extracting text from various file formats.
Supports PDF, Word (DOCX), and plain text files.
"""

import io
import logging
from typing import Optional, Tuple
from django.core.files.uploadedfile import UploadedFile

logger = logging.getLogger(__name__)


def extract_text_from_file(file: UploadedFile) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text content from uploaded file.
    
    Args:
        file: Django UploadedFile object
        
    Returns:
        Tuple of (extracted_text, error_message)
        If successful, returns (text, None)
        If failed, returns (None, error_message)
    """
    if not file:
        return None, "No file provided"
    
    file_name = file.name.lower()
    
    try:
        # Reset file pointer to beginning
        file.seek(0)
        
        # Handle PDF files
        if file_name.endswith('.pdf'):
            return _extract_from_pdf(file)
        
        # Handle Word documents
        elif file_name.endswith('.docx'):
            return _extract_from_docx(file)
        
        # Handle DOC files (older Word format)
        elif file_name.endswith('.doc'):
            return _extract_from_doc(file)
        
        # Handle plain text files
        elif file_name.endswith('.txt'):
            return _extract_from_txt(file)
        
        else:
            return None, f"Unsupported file format: {file_name.split('.')[-1]}"
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_name}: {str(e)}")
        return None, f"Error processing file: {str(e)}"


def _extract_from_pdf(file: UploadedFile) -> Tuple[Optional[str], Optional[str]]:
    """Extract text from PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader
        
        # Read PDF from file object
        pdf_reader = PdfReader(file)
        
        # Extract text from all pages
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            if page_text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")
        
        if not text_parts:
            return None, "No text content found in PDF"
        
        full_text = "\n\n".join(text_parts)
        return full_text, None
    
    except ImportError:
        return None, "PyPDF2 library not installed"
    except Exception as e:
        logger.error(f"PDF extraction error: {str(e)}")
        return None, f"Failed to extract PDF content: {str(e)}"


def _extract_from_docx(file: UploadedFile) -> Tuple[Optional[str], Optional[str]]:
    """Extract text from DOCX file using python-docx."""
    try:
        from docx import Document
        
        # Read DOCX from file object
        doc = Document(file)
        
        # Extract text from all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        if not paragraphs:
            return None, "No text content found in Word document"
        
        full_text = "\n\n".join(paragraphs)
        return full_text, None
    
    except ImportError:
        return None, "python-docx library not installed"
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        return None, f"Failed to extract Word document content: {str(e)}"


def _extract_from_doc(file: UploadedFile) -> Tuple[Optional[str], Optional[str]]:
    """Extract text from older DOC file format using docx2txt."""
    try:
        import docx2txt
        
        # docx2txt can handle both .doc and .docx
        text = docx2txt.process(file)
        
        if not text or not text.strip():
            return None, "No text content found in document"
        
        return text, None
    
    except ImportError:
        return None, "docx2txt library not installed. Please convert .doc files to .docx format."
    except Exception as e:
        logger.error(f"DOC extraction error: {str(e)}")
        return None, f"Failed to extract document content: {str(e)}"


def _extract_from_txt(file: UploadedFile) -> Tuple[Optional[str], Optional[str]]:
    """Extract text from plain text file."""
    try:
        # Try UTF-8 first
        try:
            text = file.read().decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            file.seek(0)
            text = file.read().decode('latin-1')
        
        if not text or not text.strip():
            return None, "Text file is empty"
        
        return text, None
    
    except Exception as e:
        logger.error(f"TXT extraction error: {str(e)}")
        return None, f"Failed to read text file: {str(e)}"


def get_supported_file_extensions() -> list:
    """Return list of supported file extensions."""
    return ['.pdf', '.docx', '.doc', '.txt']


def is_supported_file(filename: str) -> bool:
    """Check if file extension is supported."""
    if not filename:
        return False
    return any(filename.lower().endswith(ext) for ext in get_supported_file_extensions())
