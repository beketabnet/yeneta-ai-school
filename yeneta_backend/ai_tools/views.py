from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework import status, viewsets
from rest_framework.pagination import PageNumberPagination
from django.http import StreamingHttpResponse, HttpResponse
from django.db.models import Q
import json
import logging
import re

print("DEBUG: RELOADING ai_tools.views MODULE --------------------------------------------------")

from .llm import (
    llm_router,
    cost_analytics,
    cost_tracker,
    ollama_manager,
    vector_store,
    rag_service,
    serp_service,
    LLMRequest,
    TaskType,
    TaskComplexity,
    UserRole,
)
from .models import SavedLessonPlan, LessonPlanRating, SavedRubric, SharedFile, SavedLesson
from .serializers import (
    SavedLessonPlanSerializer,
    SavedLessonPlanListSerializer,
    LessonPlanRatingSerializer,
    SavedRubricSerializer,
    SharedFileSerializer,
    SavedRubricListSerializer,
    SavedLessonSerializer
)
from .rubric_rag_enhancer import RubricRAGEnhancer
from .grader_rag_enhancer import GraderRAGEnhancer
from .rubric_generator_rag_enhancer import RubricGeneratorRAGEnhancer
from .tutor_rag_enhancer import TutorRAGEnhancer
from users.models import User
from communications.models import SharedFileNotification
from academics.models import StudentGrade
from analytics.engagement_models import EngagementSession
from rag.services import query_curriculum_documents

logger = logging.getLogger(__name__)

# Initialize RAG enhancers with ChromaDB client
try:
    import chromadb
    from chromadb.config import Settings as ChromaSettings
    from django.conf import settings
    import os
    
    chroma_path = os.path.join(settings.MEDIA_ROOT, 'chroma_db')
    os.makedirs(chroma_path, exist_ok=True)
    
    chroma_client = chromadb.PersistentClient(
        path=chroma_path,
        settings=ChromaSettings(anonymized_telemetry=False)
    )
    
    # Initialize RAG enhancers
    rubric_rag_enhancer = RubricRAGEnhancer(chroma_client=chroma_client)
    grader_rag_enhancer = GraderRAGEnhancer(chroma_client=chroma_client)
    
    logger.info("✅ RAG enhancers initialized successfully")
except Exception as e:
    logger.warning(f"⚠️ RAG enhancers initialization failed: {e}. Features will work without RAG.")
    rubric_rag_enhancer = None
    grader_rag_enhancer = None


class LessonPlanPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 50


def clean_json_response(content: str) -> str:
    """
    Clean JSON response from LLM by extracting JSON object.
    Handles formats like:
    - ```json { ... } ```
    - ``` { ... } ```
    - ```{ ... }```
    - json { ... }
    - Text before/after JSON
    - LaTeX math syntax with backslashes
    """
    # Remove any markdown code block markers
    content = re.sub(r'```(?:json)?', '', content)
    content = re.sub(r'```', '', content)
    
    # Remove "json {" pattern that sometimes appears
    content = re.sub(r'json\s*\{', '{', content, flags=re.IGNORECASE)
    
    # Try to find JSON object directly
    # Look for the first { and last } that form a valid JSON object
    json_match = re.search(r'\{.*\}', content, re.DOTALL)
    if json_match:
        json_str = json_match.group(0).strip()
        
        # Try to parse and re-serialize to fix formatting issues
        try:
            # Attempt to load and dump to validate/fix JSON
            parsed = json.loads(json_str)
            
            # Clean newlines from all string values in the parsed JSON
            # This prevents fragmented text like "2\n1\n0" displaying as separate lines
            def clean_strings(obj):
                if isinstance(obj, dict):
                    return {k: clean_strings(v) for k, v in obj.items()}
                elif isinstance(obj, list):
                    return [clean_strings(item) for item in obj]
                elif isinstance(obj, str):
                    # Replace newlines with spaces, collapse multiple spaces
                    cleaned = obj.replace('\n', ' ').replace('\r', ' ')
                    cleaned = re.sub(r'\s+', ' ', cleaned)
                    return cleaned.strip()
                else:
                    return obj
            
            cleaned_parsed = clean_strings(parsed)
            return json.dumps(cleaned_parsed)
        except json.JSONDecodeError as e:
            # If parsing fails, try to fix common issues
            logger.warning(f"JSON parsing failed, attempting to fix: {e}")
            
            # Fix unescaped backslashes (common in LaTeX math)
            # Replace single backslashes with double backslashes, but not already escaped ones
            json_str = re.sub(r'(?<!\\)\\(?![\\"nrtbf/])', r'\\\\', json_str)
            
            # Fix unescaped newlines in strings
            json_str = re.sub(r'(?<!\\)\n', ' ', json_str)
            
            # Fix trailing commas before closing brackets/braces
            json_str = re.sub(r',\s*}', '}', json_str)
            json_str = re.sub(r',\s*]', ']', json_str)
            
            # Fix missing commas between array/object elements (common Ollama issue)
            json_str = re.sub(r'}\s*{', '},{', json_str)
            json_str = re.sub(r']\s*\[', '],[', json_str)
            
            # Fix string concatenation with + operator (invalid JSON)
            # Replace patterns like "text" + "text" with "texttext"
            json_str = re.sub(r'"\s*\+\s*"', '', json_str)
            
            # Fix missing commas after closing quotes before opening quotes
            json_str = re.sub(r'"\s+(?=")', '", ', json_str)
            
            # Fix spaces/commas after closing quotes in arrays
            json_str = re.sub(r'"\s*,\s*(?=\])', '"', json_str)
            
            # Try parsing again
            try:
                parsed = json.loads(json_str)
                return json.dumps(parsed)
            except json.JSONDecodeError as e2:
                logger.error(f"JSON parsing still failed after fixes: {e2}")
                # Last resort: try to extract just the criteria array if present
                criteria_match = re.search(r'"criteria"\s*:\s*\[(.*?)\]', json_str, re.DOTALL)
                if criteria_match:
                    logger.warning("Attempting to salvage criteria array from malformed JSON")
                    try:
                        criteria_str = f'{{"criteria": [{criteria_match.group(1)}]}}'
                        parsed = json.loads(criteria_str)
                        return json.dumps(parsed)
                    except:
                        pass
                return json_str
    
    # If still no match, return original content
    return content.strip()


def clean_lesson_content_response(content: str) -> str:
    """
    Clean JSON response specifically for lesson content generation.
    Preserves newlines in string values to maintain Markdown formatting.
    """
    # Remove any markdown code block markers
    content = re.sub(r'```(?:json)?', '', content)
    content = re.sub(r'```', '', content)
    
    # Remove "json {" pattern that sometimes appears
    content = re.sub(r'json\s*\{', '{', content, flags=re.IGNORECASE)
    
    # Try to find JSON object directly
    json_match = re.search(r'\{.*\}', content, re.DOTALL)
    if json_match:
        json_str = json_match.group(0).strip()
        
        try:
            parsed = json.loads(json_str)
            return json.dumps(parsed)
        except json.JSONDecodeError as e:
            # If parsing fails, try to fix common issues
            logger.warning(f"JSON parsing failed in lesson content, attempting to fix: {e}")
            
            # Fix unescaped backslashes
            json_str = re.sub(r'(?<!\\)\\(?![\\"nrtbf/])', r'\\\\', json_str)
            
            # Fix unescaped carriage returns and tabs
            json_str = re.sub(r'(?<!\\)\r', '\\r', json_str)
            json_str = re.sub(r'(?<!\\)\t', '\\t', json_str)
            
            # Remove any remaining control characters (0x00-0x1f) that are invalid in JSON strings
            # We exclude \n, \r, \t because we just escaped them (they are now literal backslash + char)
            # But wait, we replaced the control char with a 2-char string, so the control char is gone.
            # So we can safely remove any remaining control chars.
            json_str = re.sub(r'[\x00-\x1f]', ' ', json_str)
            
            # Fix trailing commas
            json_str = re.sub(r',\s*}', '}', json_str)
            json_str = re.sub(r',\s*]', ']', json_str)
            
            # Fix missing commas between array/object elements
            json_str = re.sub(r'}\s*{', '},{', json_str)
            json_str = re.sub(r']\s*\[', '],[', json_str)
            
            # Fix string concatenation
            json_str = re.sub(r'"\s*\+\s*"', '', json_str)
            
            # Fix missing commas between key-value pairs (common Ollama issue)
            # Look for "value" "key": pattern
            json_str = re.sub(r'"\s+(?="[\w]+"\s*:)', '", ', json_str)
            
            # Try parsing again
            try:
                parsed = json.loads(json_str)
                return json.dumps(parsed)
            except json.JSONDecodeError:
                return json_str
    
    return content.strip()


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def tutor_view(request):
    """AI Tutor endpoint with streaming or JSON response and RAG support."""
    
    message = request.data.get('message', '')
    use_rag = request.data.get('useRAG', False)
    context = request.data.get('context', '')
    stream_response = request.data.get('stream', True)  # Default to streaming response
    subject = request.data.get('subject', '')  # Optional explicit subject from frontend
    grade = request.data.get('grade', '')  # Optional explicit grade from frontend
    # Academic stream parameter (Natural Science/Social Science for Grades 11-12)
    # Frontend sends as 'stream', backend also accepts 'stream_param' for clarity
    academic_stream = request.data.get('stream_param') or request.data.get('stream')
    configured_chapter = request.data.get('chapter', '')
    region = request.data.get('region', None)
    
    if not message:
        return Response(
            {'error': 'Message is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Check for saved tutor configuration
    saved_config = None
    try:
        from .models import TutorConfiguration
        saved_config = TutorConfiguration.objects.filter(user=request.user).first()
        if saved_config:
            saved_config.update_usage()  # Track usage
            # Use saved config if frontend didn't provide explicit values
            if not grade and saved_config.grade:
                grade = saved_config.grade
            if not academic_stream and saved_config.stream:
                academic_stream = saved_config.stream
            if not subject and saved_config.subject:
                subject = saved_config.subject
            if not configured_chapter and saved_config.chapter_input:
                configured_chapter = saved_config.chapter_input
            if not use_rag and saved_config.use_ethiopian_curriculum:
                use_rag = True
            logger.info(f"📚 Using saved tutor configuration for {request.user.username}")
    except Exception as e:
        logger.debug(f"Could not load saved configuration: {e}")
    
    # RAG Context: Query curriculum documents if toggle is ON
    rag_context = ""
    curriculum_sources = []
    rag_status = 'disabled'
    rag_message = None
    
    # Get student's grade level - prefer explicit grade from frontend, fallback to user profile
    student_grade = grade if grade else getattr(request.user, 'grade', None)

    if use_rag:
        try:
            from rag.services import query_curriculum_documents

            

            
            if not student_grade:
                logger.warning(f"⚠️ Student {request.user.username} has no grade level set")
                rag_message = "Your grade level is not set. Please select a grade or update your profile."
                rag_status = 'fallback'
            else:
                logger.info(f"RAG enabled for AI Tutor: Student={request.user.username}, Grade={student_grade}, Subject={subject or 'auto-detect'}")
                
                # Analyze query for better retrieval
                query_analysis = TutorRAGEnhancer.analyze_query_intent(message)
                
                # Infer subject if not provided
                if not subject:
                    subject = TutorRAGEnhancer.infer_subject(message, subject)
                
                # Query curriculum documents if subject is identified
                if subject:
                    try:
                        logger.info(f"Querying curriculum documents: {student_grade} - {subject}")
                        
                        # Extract chapter info for filtering
                        chapter_info = query_analysis.get('chapter_info')
                        chapter_param = str(chapter_info['number']) if chapter_info and chapter_info.get('number') else None
                        
                        # Override chapter from explicit configuration
                        chapter_override = configured_chapter.strip() if isinstance(configured_chapter, str) else ''
                        if chapter_override:
                            logger.info(f"🎯 Using configured chapter focus: '{chapter_override}'")
                            try:
                                from ai_tools.chapter_utils import extract_chapter_number
                                normalized = extract_chapter_number(chapter_override)
                                normalized_number = normalized.get('number') if normalized else None
                                if normalized_number:
                                    chapter_param = str(normalized_number)
                                    query_analysis['chapter_info'] = {
                                        'number': int(normalized_number),
                                        'variants': TutorRAGEnhancer._generate_chapter_variants(int(normalized_number)),
                                        'original': normalized.get('original', chapter_override),
                                        'configured': True
                                    }
                                else:
                                    query_analysis['chapter_info'] = {
                                        'number': None,
                                        'variants': [chapter_override],
                                        'original': chapter_override,
                                        'configured': True
                                    }
                            except Exception as chapter_normalize_error:
                                logger.warning(f"⚠️ Chapter input '{chapter_override}' could not be normalized robustly: {chapter_normalize_error}")
                                query_analysis['chapter_info'] = {
                                    'number': None,
                                    'variants': [chapter_override],
                                    'original': chapter_override,
                                    'configured': True
                                }
                        
                        chapter_for_query = chapter_param if chapter_param else (chapter_override or None)
                        
                        # Use enhanced query for better semantic matching
                        enhanced_query = query_analysis.get('enhanced_query', message)
                        
                        # Enable full chapter extraction if chapter is specified
                        extract_full = bool(chapter_param)
                        
                        # If we have saved config with chapter topics, enhance the query
                        if saved_config and saved_config.chapter_topics:
                            topic_context = ", ".join(saved_config.chapter_topics[:5])
                            enhanced_query = f"{enhanced_query} Topics: {topic_context}"
                            logger.info(f"🎯 Enhanced query with saved chapter topics: {topic_context}")
                        
                        documents = query_curriculum_documents(
                            grade=student_grade,
                            subject=subject,
                            query=enhanced_query,
                            stream=academic_stream,
                            chapter=chapter_for_query,
                            region=region,
                            top_k=5,  # Get more documents for better coverage
                            extract_full_chapter=extract_full  # Extract complete chapter when chapter specified
                        )
                        
                        if documents and len(documents) > 0:
                            logger.info(f"✅ Retrieved {len(documents)} curriculum documents for tutoring")
                            
                            # Extract topics and learning objectives from full chapter documents
                            extracted_topics = []
                            extracted_objectives = []
                            chapter_title = None
                            
                            for doc in documents:
                                if doc.get('full_chapter'):
                                    # Extract topics and learning objectives from full chapter
                                    doc_topics = doc.get('topics', [])
                                    doc_objectives = doc.get('learning_objectives', [])
                                    if doc_topics:
                                        extracted_topics.extend(doc_topics)
                                    if doc_objectives:
                                        extracted_objectives.extend(doc_objectives)
                                    if not chapter_title and doc.get('title'):
                                        chapter_title = doc.get('title')
                            
                            # Remove duplicates while preserving order
                            if extracted_topics:
                                seen = set()
                                unique_topics = []
                                for topic in extracted_topics:
                                    if topic.lower() not in seen:
                                        seen.add(topic.lower())
                                        unique_topics.append(topic)
                                extracted_topics = unique_topics[:10]  # Limit to top 10
                            
                            if extracted_objectives:
                                seen = set()
                                unique_objectives = []
                                for obj in extracted_objectives:
                                    if obj.lower() not in seen:
                                        seen.add(obj.lower())
                                        unique_objectives.append(obj)
                                extracted_objectives = unique_objectives[:8]  # Limit to top 8
                            
                            # Update saved config with extracted metadata if chapter was configured
                            if saved_config and configured_chapter and (extracted_topics or extracted_objectives or chapter_title):
                                try:
                                    from .models import TutorConfiguration
                                    if extracted_topics:
                                        saved_config.chapter_topics = extracted_topics
                                    if extracted_objectives:
                                        # Store objectives - check if model has learning_objectives field
                                        if hasattr(saved_config, 'learning_objectives'):
                                            saved_config.learning_objectives = extracted_objectives
                                        else:
                                            # Fallback: store in chapter_summary
                                            objectives_text = "Learning Objectives:\n" + "\n".join([f"- {obj}" for obj in extracted_objectives[:5]])
                                            if saved_config.chapter_summary:
                                                saved_config.chapter_summary = objectives_text + "\n\n" + saved_config.chapter_summary
                                            else:
                                                saved_config.chapter_summary = objectives_text
                                    if chapter_title and not saved_config.chapter_title:
                                        saved_config.chapter_title = chapter_title
                                    update_fields = ['chapter_topics', 'chapter_summary', 'chapter_title']
                                    if hasattr(saved_config, 'learning_objectives'):
                                        update_fields.append('learning_objectives')
                                    saved_config.save(update_fields=update_fields)
                                    logger.info(f"💾 Updated saved config with extracted metadata: {len(extracted_topics)} topics, {len(extracted_objectives)} objectives")
                                except Exception as update_error:
                                    logger.warning(f"⚠️ Could not update saved config with metadata: {update_error}")
                            
                            # Format context using enhanced formatter
                            # Use larger context for full chapters
                            max_context_chars = 12000 if extract_full else 3000
                            
                            rag_context, curriculum_sources = TutorRAGEnhancer.format_rag_context(
                                documents,
                                query_analysis,
                                max_chars=max_context_chars
                            )
                            
                            # Enhance context with extracted topics and objectives if available
                            if extracted_topics or extracted_objectives:
                                enhancement_note = "\n\n=== CHAPTER METADATA ==="
                                if extracted_topics:
                                    topics_text = ", ".join(extracted_topics[:5])
                                    enhancement_note += f"\n**Key Topics in this Chapter**: {topics_text}"
                                if extracted_objectives:
                                    objectives_text = "\n".join([f"- {obj}" for obj in extracted_objectives[:5]])
                                    enhancement_note += f"\n**Learning Objectives**:\n{objectives_text}"
                                enhancement_note += "\n=== END METADATA ===\n"
                                rag_context = enhancement_note + rag_context
                            
                            if curriculum_sources:
                                rag_status = 'success'
                                logger.info(f"✅ RAG context built from sources: {', '.join(set(curriculum_sources))}")
                            else:
                                logger.warning(f"⚠️ Documents retrieved but no content extracted")
                                rag_context = ""
                                rag_message = "Curriculum documents found but content could not be extracted."
                                rag_status = 'fallback'
                        else:
                            logger.warning(f"⚠️ No curriculum documents found for {student_grade} - {subject}")
                            rag_message = f"No curriculum documents available for {subject} at your grade level."
                            rag_status = 'fallback'
                    
                    except Exception as query_error:
                        logger.error(f"❌ Error querying curriculum documents: {str(query_error)}", exc_info=True)
                        rag_message = f"Error accessing curriculum documents: {str(query_error)}"
                        rag_status = 'fallback'
                else:
                    logger.info(f"ℹ️ Could not identify subject from message: {message[:50]}...")
                    rag_message = "Subject could not be identified from your question. Please select a subject or mention it in your question."
                    rag_status = 'fallback'
        
        except ImportError as import_error:
            logger.error(f"❌ Failed to import RAG services: {str(import_error)}")
            rag_message = "RAG system is not available."
            rag_status = 'fallback'
        
        except Exception as e:
            logger.error(f"❌ Unexpected error in RAG processing: {str(e)}", exc_info=True)
            rag_message = f"Unexpected error: {str(e)}"
            rag_status = 'fallback'
        
        # Log final RAG status
        if rag_context:
            logger.info(f"✅ RAG processing successful - using curriculum content")
        else:
            logger.warning(f"⚠️ RAG processing failed - falling back to AI model only")
            if rag_message:
                logger.warning(f"⚠️ RAG error details: {rag_message}")
    
    # Enhanced system prompt using the new builder method
    enhanced_system_prompt = TutorRAGEnhancer.build_tutor_prompt(
        rag_context=rag_context,
        use_rag=use_rag,
        saved_config=saved_config,
        grade_level=student_grade
    )
    
    # Add RAG status note for logging (not sent to model)
    if use_rag and rag_status == 'success':
        logger.info(f"💬 Tutor request with RAG: sources={len(curriculum_sources)}")
    elif use_rag and rag_status == 'fallback':
        logger.info(f"💬 Tutor request without RAG (fallback): {rag_message}")
    else:
        logger.info(f"💬 Tutor request without RAG (disabled)")

    # Create LLM request
    llm_request = LLMRequest(
        prompt=message,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.TUTORING,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=enhanced_system_prompt,
        temperature=0.7,
        max_tokens=1500,  # Increased for more detailed explanations
        stream=stream_response,
        context_documents=[context] if context else None,
    )
    
    # Generate alert if student message indicates issues (only for students)
    tutor_response_content = ""
    if request.user.role == 'Student':
        try:
            from alerts.alert_generator import AlertGenerator
            
            # For streaming, we'll generate alert based on student message only
            # For non-streaming, we can include tutor response
            if not stream_response:
                # Get tutor response first for non-streaming
                temp_response = llm_router.process_request(llm_request)
                tutor_response_content = temp_response.content
            
            # Check if alert should be generated
            if AlertGenerator.should_generate_alert(request.user, message):
                AlertGenerator.generate_alert_from_tutor_interaction(
                    student=request.user,
                    student_message=message,
                    tutor_response=tutor_response_content,
                    subject=subject,
                    topic=None
                )
                logger.info(f"Alert generated for student {request.user.username} from tutor interaction")
        except Exception as alert_error:
            # Don't fail the request if alert generation fails
            logger.error(f"Failed to generate alert: {alert_error}")
    
    # Check if streaming is requested
    if stream_response:
        # Stream response with RAG metadata in headers
        def generate():
            try:
                for chunk in llm_router.process_request_stream(llm_request):
                    yield chunk
            except Exception as e:
                logger.error(f"❌ AI Tutor streaming error: {e}", exc_info=True)
                yield f"\n\n⚠️ Error: Unable to generate response. Please try again."
        
        response = StreamingHttpResponse(generate(), content_type='text/plain')
        # Add RAG metadata to response headers
        response['X-RAG-Status'] = rag_status
        if rag_status == 'success' and curriculum_sources:
            response['X-RAG-Sources'] = ','.join(set(curriculum_sources))
        if rag_message:
            response['X-RAG-Message'] = rag_message
        return response
    else:
        # Non-streaming JSON response
        try:
            if not tutor_response_content:
                # Get tutor response for non-streaming
                temp_response = llm_router.process_request(llm_request)
                tutor_response_content = temp_response.content
            else:
                # Use cached response from alert generation
                temp_response = type('obj', (object,), {
                    'content': tutor_response_content,
                    'model': 'cached',
                    'input_tokens': 0,
                    'output_tokens': 0,
                    'cost_usd': 0,
                    'latency_ms': 0
                })()
            
            return Response({
                'message': message,
                'response': temp_response.content if hasattr(temp_response, 'content') else tutor_response_content,
                'model': str(temp_response.model) if hasattr(temp_response, 'model') else 'cached',
                'input_tokens': temp_response.input_tokens if hasattr(temp_response, 'input_tokens') else 0,
                'output_tokens': temp_response.output_tokens if hasattr(temp_response, 'output_tokens') else 0,
                'total_tokens': (temp_response.input_tokens + temp_response.output_tokens) if hasattr(temp_response, 'input_tokens') else 0,
                'cost_usd': temp_response.cost_usd if hasattr(temp_response, 'cost_usd') else 0,
                'latency_ms': temp_response.latency_ms if hasattr(temp_response, 'latency_ms') else 0
            })
        except Exception as e:
            logger.error(f"AI Tutor error: {e}")
            return Response(
                {'error': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def lesson_planner_view(request):
    """
    Generate comprehensive lesson plans using AI with UbD-5E-Differentiated framework.
    Follows Ethiopian educational standards with optional RAG support.
    """
    
    # I. Extract all user inputs for contextual lesson generation
    topic = request.data.get('topic', '')
    grade_level = request.data.get('gradeLevel', '')
    objectives = request.data.get('objectives', '')
    duration = request.data.get('duration', 45)
    use_rag = request.data.get('useRAG', False)
    subject = request.data.get('subject', '')
    stream = request.data.get('stream', None)
    
    # Explicit RAG parameters
    explicit_region = request.data.get('region')
    explicit_chapter = request.data.get('chapter')
    
    # Resource Constraint Profile
    resource_constraints = request.data.get('resourceConstraints', {})
    available_materials = resource_constraints.get('availableMaterials', ['Textbook', 'Chalk/Whiteboard', 'Notebooks'])
    internet_access = resource_constraints.get('internetAccess', False)
    electricity_access = resource_constraints.get('electricityAccess', True)
    field_trip_available = resource_constraints.get('fieldTripAvailable', False)
    budget_limit = resource_constraints.get('budgetLimit', 'Very Low')
    class_size = resource_constraints.get('classSize', 40)
    
    # Student Readiness Profile
    student_readiness = request.data.get('studentReadiness', {})
    prior_knowledge = student_readiness.get('averagePriorKnowledge', 'Medium')
    common_misconceptions = student_readiness.get('commonMisconceptions', [])
    special_needs_groups = student_readiness.get('specialNeedsGroups', [])
    
    # Local/Community Context
    local_context = request.data.get('localContext', {})
    region = local_context.get('region', 'Urban')
    dominant_economy = local_context.get('dominantEconomy', '')
    cultural_considerations = local_context.get('culturalConsiderations', [])
    
    # MoE Standard (optional)
    moe_standard_id = request.data.get('moeStandardId', '')
    
    if not topic:
        return Response(
            {'error': 'Topic is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # RAG Context: Query curriculum documents if toggle is ON
    rag_context = ""
    curriculum_sources = []
    rag_error_message = None
    
    if use_rag and grade_level:
        try:
            from rag.services import query_curriculum_documents
            
            logger.info(f"RAG enabled for lesson planning: Grade={grade_level}, Subject={subject}, Topic={topic}")
            
            # Extract subject from topic if not provided
            # Try to infer subject from common patterns
            if not subject:
                subject_keywords = {
                    'math': 'Mathematics',
                    'algebra': 'Mathematics',
                    'geometry': 'Mathematics',
                    'physics': 'Physics',
                    'chemistry': 'Chemistry',
                    'biology': 'Biology',
                    'science': 'General Science',
                    'english': 'English',
                    'amharic': 'Amharic',
                    'history': 'History',
                    'geography': 'Geography',
                    'economics': 'Economics',
                }
                
                topic_lower = topic.lower()
                for keyword, subj in subject_keywords.items():
                    if keyword in topic_lower:
                        subject = subj
                        logger.info(f"Subject inferred from topic: {subject}")
                        break
            
            # Query curriculum documents if subject is identified
            if subject:
                try:
                    from .lesson_planner_rag_enhancer import LessonPlannerRAGEnhancer
                    
                    # Analyze topic for chapter references
                    # Analyze topic for chapter references
                    chapter_info = LessonPlannerRAGEnhancer.analyze_topic_for_chapter(topic, objectives)
                    chapter_param = str(chapter_info['number']) if chapter_info else None
                    
                    # Override with explicit chapter if provided
                    if explicit_chapter:
                        chapter_param = explicit_chapter
                        logger.info(f"Using explicit chapter for RAG: {chapter_param}")
                    
                    # Build enhanced query
                    query_text = LessonPlannerRAGEnhancer.build_lesson_planning_query(
                        topic=topic,
                        objectives=objectives,
                        grade_level=grade_level,
                        subject=subject,
                        chapter_info=chapter_info
                    )
                    
                    logger.info(f"Querying curriculum documents: {grade_level} - {subject}")
                    
                    # Enable full chapter extraction if chapter detected
                    extract_full = bool(chapter_param)
                    
                    documents = query_curriculum_documents(
                        grade=grade_level,
                        subject=subject,
                        query=query_text,
                        stream=stream,
                        chapter=chapter_param,
                        region=explicit_region or region, # Prioritize explicit geographic region
                        top_k=5,
                        extract_full_chapter=extract_full
                    )
                    
                    if documents and len(documents) > 0:
                        logger.info(f"✅ Retrieved {len(documents)} curriculum documents for lesson planning")
                        
                        # Format context using enhancer
                        rag_context, curriculum_sources = LessonPlannerRAGEnhancer.format_lesson_planning_context(
                            documents=documents,
                            topic=topic,
                            chapter_info=chapter_info,
                            max_chars=10000
                        )
                        
                        if curriculum_sources:
                            logger.info(f"✅ RAG context built from sources: {', '.join(set(curriculum_sources))}")
                        else:
                            logger.warning(f"⚠️ Documents retrieved but no content extracted")
                            rag_context = ""
                            rag_error_message = "Curriculum documents found but content could not be extracted."
                    else:
                        logger.warning(f"⚠️ No curriculum documents found for {grade_level} - {subject}")
                        rag_error_message = f"No curriculum documents found for {grade_level} - {subject}. Upload documents in Curriculum Manager to use this feature."
                
                except Exception as query_error:
                    logger.error(f"❌ Error querying curriculum documents: {str(query_error)}", exc_info=True)
                    rag_error_message = f"Error accessing curriculum documents: {str(query_error)}"
            else:
                logger.warning(f"⚠️ Could not identify subject from topic: {topic}")
                rag_error_message = "Subject could not be identified. Please select a subject explicitly."
        
        except ImportError as import_error:
            logger.error(f"❌ Failed to import RAG services: {str(import_error)}")
            rag_error_message = "RAG system is not available. Please contact administrator."
        
        except Exception as e:
            logger.error(f"❌ Unexpected error in RAG processing: {str(e)}", exc_info=True)
            rag_error_message = f"Unexpected error: {str(e)}"
        
        # Log final RAG status
        if rag_context:
            logger.info(f"✅ RAG processing successful - using curriculum content")
        else:
            logger.warning(f"⚠️ RAG processing failed - falling back to AI model only")
            if rag_error_message:
                logger.warning(f"⚠️ RAG error details: {rag_error_message}")
    
    # II. Build enhanced prompt following research standards
    # Role-Based Prompting with specialized Ethiopian context
    role_definition = f"""You are an expert Ethiopian Curriculum Designer specializing in low-resource {'rural' if region == 'Rural' else 'urban'} school environments. 

Your expertise includes:
- Understanding by Design (UbD) framework for backward planning
- 5E Instructional Model (Engage, Explore, Explain, Elaborate, Evaluate) for constructivist learning
- Cognitive Load Theory (CLT) for optimizing student learning
- Differentiated instruction for diverse learner needs
- Ethiopian Ministry of Education (MoE) curriculum standards and competencies
- Resource-adaptive teaching strategies for constrained environments"""

    # Build comprehensive context section
    context_section = f"""
=== LESSON PLANNING CONTEXT ===

**Administrative Details:**
- Grade Level: {grade_level}
- Subject: {subject}
- Topic: {topic}
- Duration: {duration} minutes
- MoE Standard: {moe_standard_id if moe_standard_id else 'General curriculum alignment'}

**Resource Constraint Profile:**
- Available Materials: {', '.join(available_materials)}
- Internet Access: {'Yes' if internet_access else 'No'}
- Electricity: {'Yes' if electricity_access else 'No'}
- Field Trip Capability: {'Yes' if field_trip_available else 'No'}
- Budget: {budget_limit}
- Class Size: {class_size} students

**Student Readiness:**
- Prior Knowledge Level: {prior_knowledge}
- Common Misconceptions: {', '.join(common_misconceptions) if common_misconceptions else 'None specified'}
- Special Needs Groups: {', '.join(special_needs_groups) if special_needs_groups else 'None specified'}

**Local Context:**
- Region: {region}
- Dominant Economy: {dominant_economy if dominant_economy else 'Mixed'}
- Cultural Considerations: {', '.join(cultural_considerations) if cultural_considerations else 'Standard Ethiopian context'}

**Learning Objectives:**
{objectives}

{rag_context}
=== END CONTEXT ===
"""

    # Structured constraints and requirements
    constraints_section = """
**CRITICAL REQUIREMENTS - CONSTRAINT-DRIVEN ACTIVITY SUBSTITUTION:**

1. **Resource Adaptation**: You MUST design ALL activities to work with ONLY the available materials listed above. 
   - If a typical activity requires unavailable resources (e.g., computers, lab equipment, projectors), you MUST substitute with low-cost alternatives using locally available materials
   - Examples: Replace "virtual simulation" with "role-play using local objects", replace "video" with "teacher demonstration", replace "online research" with "textbook analysis"

2. **Active Learning Mandate**: Combat traditional lecture methods by ensuring:
   - Students are actively engaged in ALL phases (not just listening)
   - Include collaborative work, hands-on activities, discussion, and problem-solving
   - Minimize teacher talk time, maximize student activity time

3. **Pedagogical Framework Adherence**:
   - Use UbD backward design: Start with clear learning outcomes, then assessment, then activities
   - Structure instruction using 5E Model: Engage → Explore → Explain → Elaborate → Evaluate
   - Apply CLT principles: Minimize extraneous load, optimize germane load through clear instructions and scaffolding

4. **Differentiation**: Provide adaptations for at least 3 levels:
   - Below grade level (struggling learners)
   - At grade level (typical learners)
   - Above grade level (advanced learners)
"""

    # Enhanced JSON structure specification
    json_structure = """{
    "title": "Engaging lesson title",
    "grade": "Grade level",
    "subject": "Subject name",
    "topic": "Specific topic",
    "duration": 45,
    "moeStandardId": "MoE curriculum code if applicable",
    
    "objectives": [
        "Measurable objective 1 (using action verbs: analyze, evaluate, create)",
        "Measurable objective 2 aligned with MoE competencies"
    ],
    "essentialQuestions": [
        "Big idea question that drives deep understanding",
        "Transfer question connecting to real life"
    ],
    "moeCompetencies": [
        "Relevant MoE competency (e.g., critical thinking, problem-solving)"
    ],
    
    "assessmentPlan": {
        "formativeChecks": [
            "Quick check during Engage phase",
            "Observation during Explore phase",
            "Exit ticket at end"
        ],
        "summativeTask": "Final performance task description",
        "successCriteria": [
            "Student can demonstrate...",
            "Student can explain..."
        ]
    },
    
    "materials": [
        "Only materials from available list",
        "Local/low-cost substitutes"
    ],
    "teacherPreparation": [
        "Specific prep step 1",
        "Specific prep step 2"
    ],
    
    "fiveESequence": [
        {
            "phase": "Engage",
            "duration": 5,
            "activities": ["Hook activity to capture interest"],
            "teacherActions": ["What teacher does"],
            "studentActions": ["What students do actively"],
            "materials": ["Materials for this phase"]
        },
        {
            "phase": "Explore",
            "duration": 15,
            "activities": ["Hands-on exploration with available materials"],
            "teacherActions": ["Facilitate, observe, question"],
            "studentActions": ["Investigate, collaborate, discover"],
            "materials": ["Materials for exploration"]
        },
        {
            "phase": "Explain",
            "duration": 10,
            "activities": ["Students explain findings, teacher clarifies"],
            "teacherActions": ["Guide discussion, introduce vocabulary"],
            "studentActions": ["Share observations, ask questions"],
            "materials": ["Materials for demonstration"]
        },
        {
            "phase": "Elaborate",
            "duration": 10,
            "activities": ["Apply learning to new context"],
            "teacherActions": ["Provide new challenge"],
            "studentActions": ["Apply concepts, solve problems"],
            "materials": ["Materials for application"]
        },
        {
            "phase": "Evaluate",
            "duration": 5,
            "activities": ["Assess understanding"],
            "teacherActions": ["Administer assessment"],
            "studentActions": ["Demonstrate learning"],
            "materials": ["Assessment materials"]
        }
    ],
    
    "differentiationStrategies": [
        {
            "level": "Below Grade Level",
            "contentAdaptations": ["Simplified text", "Visual aids"],
            "processAdaptations": ["More scaffolding", "Peer support"],
            "productAdaptations": ["Alternative demonstration methods"]
        },
        {
            "level": "At Grade Level",
            "processAdaptations": ["Standard instruction with guided practice"]
        },
        {
            "level": "Above Grade Level",
            "contentAdaptations": ["Extension readings"],
            "processAdaptations": ["Independent research"],
            "productAdaptations": ["Creative project options"]
        }
    ],
    
    "homework": "Meaningful practice or preparation for next lesson",
    "extensions": ["Optional enrichment for fast finishers"],
    
    "reflectionPrompts": [
        "Did the low-cost activities engage students effectively?",
        "What misconceptions emerged?",
        "How can I improve differentiation?"
    ]
}"""

    # Complete prompt assembly
    prompt = f"""{role_definition}

{context_section}

{constraints_section}

**YOUR TASK:**
Create a comprehensive, standards-aligned lesson plan that:
1. Follows UbD backward design (outcomes → assessment → instruction)
2. Uses the 5E instructional sequence for active, constructivist learning
3. Adapts ALL activities to work with ONLY the available resources listed
4. Includes differentiation for diverse learners
5. Aligns with Ethiopian MoE curriculum and competencies
6. Uses culturally relevant examples from the local context

**OUTPUT FORMAT:**
Provide your response as a valid JSON object with this exact structure:

{json_structure}

**CRITICAL JSON FORMATTING RULES:** 
- Return ONLY valid JSON - no text before or after
- NO string concatenation with + operator (e.g., "text" + "text" is INVALID)
- NO trailing commas before closing brackets or braces
- ALL string values must be properly quoted
- NO comments or explanations outside the JSON
- Use proper comma separation between all array elements
- Ensure all activities are feasible with the available materials
- Make learning active and student-centered, not lecture-based
- Include specific, actionable instructions for teachers
- Use Ethiopian context and examples throughout

**MULTILINGUAL HANDLING:**
- If the curriculum content or user input is in a non-English language (e.g., Amharic, Oromo):
  1. Identify the language.
  2. Generate the ENTIRE lesson plan in that language.
  3. Ensure all pedagogical terms are appropriately translated or adapted.
  4. Maintain the JSON structure keys in English, but the VALUES should be in the target language (e.g., "title": "የአማርኛ ርዕስ")."""

    # System prompt based on RAG status
    if use_rag and rag_context:
        system_prompt = """You are an expert Ethiopian curriculum specialist and lesson planner with deep knowledge of Understanding by Design (UbD), the 5E Instructional Model, and Cognitive Load Theory. 

You have access to official Ethiopian curriculum content above. Your lesson plans MUST:
1. Strictly align with the provided curriculum materials
2. Follow modern pedagogical frameworks (UbD + 5E + CLT)
3. Adapt to resource constraints through creative substitution
4. Prioritize active, student-centered learning over lectures
5. Include comprehensive differentiation strategies

Base all content decisions on the curriculum materials provided while applying best practices in instructional design.

CRITICAL: Return ONLY valid JSON. No string concatenation with +, no trailing commas, no text outside JSON."""
    else:
        system_prompt = """You are an expert Ethiopian curriculum specialist and lesson planner with deep knowledge of Understanding by Design (UbD), the 5E Instructional Model, and Cognitive Load Theory.

Your lesson plans MUST:
1. Align with Ethiopian Ministry of Education standards and competencies
2. Follow modern pedagogical frameworks (UbD + 5E + CLT)
3. Adapt to resource constraints through creative substitution
4. Prioritize active, student-centered learning over lectures
5. Include comprehensive differentiation strategies

Apply best practices in instructional design while ensuring cultural relevance and feasibility in Ethiopian contexts.

CRITICAL: Return ONLY valid JSON. No string concatenation with +, no trailing commas, no text outside JSON."""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.LESSON_PLANNING,
        complexity=TaskComplexity.ADVANCED,
        system_prompt=system_prompt,
        temperature=0.7,
        max_tokens=4000,  # Increased for comprehensive lesson plans
    )
    
    try:
        # Use robust JSON processing
        lesson_plan = llm_router.process_json_request(llm_request)
        
        # Ensure required fields exist with proper defaults
        if 'title' not in lesson_plan:
            lesson_plan['title'] = f'Lesson Plan: {topic}'
        if 'grade' not in lesson_plan:
            lesson_plan['grade'] = grade_level
        if 'subject' not in lesson_plan:
            lesson_plan['subject'] = subject
        if 'topic' not in lesson_plan:
            lesson_plan['topic'] = topic
        if 'duration' not in lesson_plan:
            lesson_plan['duration'] = duration
        if 'objectives' not in lesson_plan:
            lesson_plan['objectives'] = []
        if 'materials' not in lesson_plan:
            lesson_plan['materials'] = []
        
        # Ensure backward compatibility with legacy activities field
        if 'activities' not in lesson_plan and 'fiveESequence' not in lesson_plan:
            lesson_plan['activities'] = []
        elif 'fiveESequence' in lesson_plan and 'activities' not in lesson_plan:
            # Generate legacy activities from 5E sequence for backward compatibility
            legacy_activities = []
            for phase in lesson_plan.get('fiveESequence', []):
                legacy_activities.append({
                    'duration': phase.get('duration', 5),
                    'description': f"{phase.get('phase', 'Activity')}: {', '.join(phase.get('activities', []))}"
                })
            lesson_plan['activities'] = legacy_activities

    except Exception as e:
        logger.error(f"Lesson generation failed: {e}")
        return Response(
            {'error': 'Failed to generate valid lesson plan. Please try again.'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )
            
    # Add context metadata to response
    lesson_plan['resourceConstraints'] = {
        'availableMaterials': available_materials,
        'internetAccess': internet_access,
        'electricityAccess': electricity_access,
        'fieldTripAvailable': field_trip_available,
        'budgetLimit': budget_limit,
        'classSize': class_size
    }
    
    lesson_plan['studentReadiness'] = {
        'averagePriorKnowledge': prior_knowledge,
        'commonMisconceptions': common_misconceptions,
        'specialNeedsGroups': special_needs_groups
    }
    
    lesson_plan['localContext'] = {
        'region': region,
        'dominantEconomy': dominant_economy,
        'culturalConsiderations': cultural_considerations
    }
    
    # Add curriculum sources and RAG status to response
    if use_rag and curriculum_sources:
        lesson_plan['curriculum_sources'] = list(set(curriculum_sources))
        lesson_plan['rag_enabled'] = True
        lesson_plan['rag_status'] = 'success'
        logger.info(f"✅ Lesson plan generated successfully with RAG: {lesson_plan.get('title')}")
    elif use_rag and not curriculum_sources:
        lesson_plan['rag_enabled'] = False
        lesson_plan['rag_status'] = 'fallback'
        lesson_plan['rag_message'] = rag_error_message or 'No curriculum documents available'
        logger.warning(f"⚠️ Lesson plan generated without RAG (fallback): {lesson_plan.get('title')}")
    else:
        lesson_plan['rag_enabled'] = False
        lesson_plan['rag_status'] = 'disabled'
        logger.info(f"✅ Lesson plan generated without RAG (by choice): {lesson_plan.get('title')}")
    
    return Response(lesson_plan)



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_teacher_note_view(request):
    """
    Generate a high-quality Teacher's Note based on a lesson plan.
    Includes detailed visual aid descriptions and pedagogical guidance.
    Supports RAG for curriculum-aligned content generation.
    """
    
    # Extract parameters
    lesson_plan = request.data.get('lessonPlan', {})
    use_rag = request.data.get('useRAG', False)
    use_rag = request.data.get('useRAG', False)
    chapter = request.data.get('chapter', '')
    region = request.data.get('region', None)
    
    if not lesson_plan:
        return Response(
            {'error': 'Lesson plan data is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
        
    # Extract key details from lesson plan
    topic = lesson_plan.get('topic', '')
    grade_level = lesson_plan.get('grade', '')
    subject = lesson_plan.get('subject', '')
    objectives = lesson_plan.get('objectives', [])
    duration = lesson_plan.get('duration', 45)
    
    # RAG Context
    rag_context = ""
    curriculum_sources = []
    rag_error_message = None
    
    if use_rag and grade_level and subject:
        try:
            from rag.services import query_curriculum_documents
            from .lesson_planner_rag_enhancer import LessonPlannerRAGEnhancer
            
            logger.info(f"RAG enabled for teacher note generation: Grade={grade_level}, Subject={subject}, Topic={topic}")
            
            # Analyze topic for chapter references if not provided
            chapter_info = None
            if chapter:
                # If chapter provided explicitly
                match = re.search(r'\d+', chapter)
                chapter_num = int(match.group()) if match else None
                chapter_info = {'number': chapter_num, 'variants': [chapter]} if chapter_num else None
            else:
                # Analyze topic
                objectives_text = " ".join(objectives) if isinstance(objectives, list) else str(objectives)
                chapter_info = LessonPlannerRAGEnhancer.analyze_topic_for_chapter(topic, objectives_text)
            
            chapter_param = str(chapter_info['number']) if chapter_info else None
            
            # Build enhanced query
            objectives_text = " ".join(objectives) if isinstance(objectives, list) else str(objectives)
            query_text = LessonPlannerRAGEnhancer.build_lesson_planning_query(
                topic=topic,
                objectives=objectives_text,
                grade_level=grade_level,
                subject=subject,
                chapter_info=chapter_info
            )
            
            # Enable full chapter extraction if chapter detected
            extract_full = bool(chapter_param)
            
            documents = query_curriculum_documents(
                grade=grade_level,
                subject=subject,
                query=query_text,
                chapter=chapter_param,
                region=region,
                top_k=5,
                extract_full_chapter=extract_full
            )
            
            if documents and len(documents) > 0:
                # Format context
                rag_context, curriculum_sources = LessonPlannerRAGEnhancer.format_lesson_planning_context(
                    documents=documents,
                    topic=topic,
                    chapter_info=chapter_info,
                    max_chars=12000  # Increased for full note generation
                )
                
                if curriculum_sources:
                    logger.info(f"✅ RAG context built for teacher note from sources: {', '.join(set(curriculum_sources))}")
                else:
                    rag_error_message = "Curriculum documents found but content could not be extracted."
            else:
                rag_error_message = f"No curriculum documents found for {grade_level} - {subject}."
                
        except Exception as e:
            logger.error(f"❌ Error in RAG processing for teacher note: {str(e)}", exc_info=True)
            rag_error_message = f"Error accessing curriculum documents: {str(e)}"
    
    # Build Prompt
    role_definition = f"""You are an Expert Teacher Mentor and Curriculum Specialist.
Your task is to generate a high-quality **Teacher's Note** based on the provided lesson plan.
This note is NOT a script for students, but a comprehensive guide for the teacher to deliver an exceptional lesson."""

    context_section = f"""
=== LESSON CONTEXT ===
Grade: {grade_level}
Subject: {subject}
Topic: {topic}
Duration: {duration} minutes
Objectives:
{json.dumps(objectives, indent=2)}

=== LESSON PLAN STRUCTURE ===
{json.dumps(lesson_plan, indent=2)}

{rag_context}
=== END CONTEXT ===
"""

    task_instructions = """
**YOUR TASK:**
Create a detailed Teacher's Note that empowers the teacher to deliver this lesson effectively.
Focus on pedagogical strategies, deep explanations, and **multimodal engagement**.

**IMPORTANT: Return ONLY a valid JSON object. Do not include any text before or after the JSON.**

**REQUIRED JSON STRUCTURE:**
{
    "title": "Teacher's Note: [Topic]",
    "overview": "Brief summary of the lesson's core value and what the teacher should focus on.",
    "keyConcepts": [
        {
            "concept": "Name of the concept",
            "explanation": "Deep dive explanation for the teacher to master this concept."
        }
    ],
    "visualAids": [
        {
            "title": "Name of the visual aid (e.g., Water Cycle Diagram)",
            "description": "Detailed description of what to draw or show. Be specific (e.g., 'Draw a circle with arrows...').",
            "type": "Diagram/Chart/Photo/Object"
        }
    ],
    "teachingStrategies": [
        {
            "phase": "Hook/Introduction",
            "strategy": "Name of strategy (e.g., Storytelling)",
            "activity": "Specific instructions on what to do."
        },
        {
            "phase": "Core Explanation",
            "strategy": "Analogy/Demonstration",
            "activity": "How to explain the difficult parts using the strategy."
        },
        {
            "phase": "Differentiation",
            "strategy": "Scaffolding",
            "activity": "Specific support for struggling learners."
        }
    ],
    "assessmentQuestions": [
        "Question 1 to check understanding?",
        "Question 2 to provoke deep thinking?"
    ],
    "misconceptions": [
        {
            "misconception": "Common wrong idea",
            "correction": "How to correct it."
        }
    ]
}

**STYLE GUIDELINES:**
*   **Tone**: Mentorship - supportive, expert, and practical.
*   **Quality**: High. Go beyond generic advice. Provide specific examples relevant to the Ethiopian context.

**MULTILINGUAL HANDLING:**
*   **Language Matching**: Generate the Teacher's Note in the SAME LANGUAGE as the provided "LESSON PLAN STRUCTURE". 
*   **Consistency**: If the lesson plan is in Amharic, the Teacher's Note must be in Amharic.
*   **JSON Keys**: Keep JSON keys in English, but values in the target language.
"""

    system_prompt = """You are an Expert Teacher Mentor. Generate a high-quality structured Teacher's Note in JSON format."""

    llm_request = LLMRequest(
        prompt=f"{role_definition}\n\n{context_section}\n\n{task_instructions}",
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.LESSON_PLANNING,
        complexity=TaskComplexity.ADVANCED,
        system_prompt=system_prompt,
        temperature=0.7,
        max_tokens=5000,
    )
    
    try:
        # Use robust JSON processing
        note_data = llm_router.process_json_request(llm_request)
        
        # Construct the response object expected by the frontend
        # We wrap the structured data but also keep a 'content' field for backward compatibility if needed,
        # or we can just return the structured data and let the frontend handle it.
        # The frontend expects { title, content, ... }
        
        # Let's return the structured data directly mixed with metadata
        response_data = {
            'title': note_data.get('title', f"Teacher's Note: {topic}"),
            'overview': note_data.get('overview', ''),
            'keyConcepts': note_data.get('keyConcepts', []),
            'visualAids': note_data.get('visualAids', []),
            'teachingStrategies': note_data.get('teachingStrategies', []),
            'assessmentQuestions': note_data.get('assessmentQuestions', []),
            'misconceptions': note_data.get('misconceptions', []),
            'rag_enabled': bool(rag_context),
            'curriculum_sources': list(set(curriculum_sources)) if curriculum_sources else []
        }
            
        return Response(response_data)

    except Exception as e:
        logger.error(f"Teacher note generation error: {e}", exc_info=True)
        return Response(
            {'error': 'Failed to generate teacher note'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_ai_teacher_lesson_view(request):
    """
    Generate a lesson script/note for the AI Teacher (Student View).
    This serves the "Lesson Creator" for the Virtual Classroom.
    Uses RAG to align with Ethiopian curriculum if available.
    """
    try:
        # Extract parameters
        subject = request.data.get('subject', '')
        grade_level = request.data.get('grade_level', '')
        topic = request.data.get('topic', '')
        duration = request.data.get('duration', 45)
        learning_objectives = request.data.get('learning_objectives', [])
        stream = request.data.get('stream', '')
        chapter = request.data.get('chapter', '')
        use_ethiopian_curriculum = request.data.get('use_ethiopian_curriculum', False)
        region = request.data.get('region', None)
        rag_query = request.data.get('rag_query', '')

        if not topic or not subject:
            return Response(
                {'error': 'Topic and Subject are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
            
        logger.info(f"🤖 Generating AI Teacher Lesson for: {topic} ({grade_level} {subject})")
        
        # RAG Context
        rag_context = ""
        curriculum_sources = []
        rag_enabled = False
        
        if use_ethiopian_curriculum and grade_level and subject:
            try:
                from rag.services import query_curriculum_documents
                from .lesson_planner_rag_enhancer import LessonPlannerRAGEnhancer
                
                # Analyze topic for chapter references if not provided
                chapter_info = None
                if chapter:
                    match = re.search(r'\d+', chapter)
                    chapter_num = int(match.group()) if match else None
                    chapter_info = {'number': chapter_num, 'variants': [chapter]} if chapter_num else None
                else:
                    # Analyze topic
                    objectives_text = " ".join(learning_objectives) if isinstance(learning_objectives, list) else str(learning_objectives)
                    chapter_info = LessonPlannerRAGEnhancer.analyze_topic_for_chapter(topic, objectives_text)
                
                chapter_param = str(chapter_info['number']) if chapter_info else None
                
                # Build query
                query_text = rag_query if rag_query else LessonPlannerRAGEnhancer.build_lesson_planning_query(
                    topic=topic,
                    objectives=learning_objectives,
                    grade_level=grade_level,
                    subject=subject,
                    chapter_info=chapter_info
                )
                
                # Enable full chapter extraction if chapter detected
                extract_full = bool(chapter_param)
                
                documents = query_curriculum_documents(
                    grade=grade_level,
                    subject=subject,
                    query=query_text,
                    chapter=chapter_param,
                    region=region,
                    top_k=5,
                    extract_full_chapter=extract_full
                )
                
                if documents and len(documents) > 0:
                    # Format context
                    rag_context, curriculum_sources = LessonPlannerRAGEnhancer.format_lesson_planning_context(
                        documents=documents,
                        topic=topic,
                        chapter_info=chapter_info,
                        max_chars=8000
                    )
                    rag_enabled = True
                    logger.info(f"✅ RAG context built for AI Teacher lesson")
                else:
                    logger.warning(f"⚠️ No curriculum documents found for AI Teacher lesson")
                    
            except Exception as e:
                logger.error(f"❌ Error in RAG processing for AI Teacher lesson: {str(e)}", exc_info=True)
        
        # Build Prompt for AI Teacher Script
        role_definition = f"""You are an engaging AI Teacher for the 'Yeneta AI School' virtual classroom.
Your goal is to explain the topic '{topic}' to a {grade_level} student in a clear, interactive, and encouraging way.
The explanation should be broken down into 'slides' or 'whiteboard logical segments'."""

        context_section = f"""
=== LESSON CONTEXT ===
Subject: {subject}
Grade: {grade_level}
Topic: {topic}
Stream: {stream if stream else 'General'}
Duration: {duration} minutes
Objectives: {json.dumps(learning_objectives)}
Region: {region if region else 'National'}

{rag_context}
=== END CONTEXT ===
"""

        json_structure = """
{
    "title": "Lesson Title",
    "content": "A detailed, engaging script/explanation of the full lesson content, formatted with Markdown (headings, bullet points, bold text). It should read like a natural lecture/explanation.",
    "objectives": ["Objective 1", "Objective 2"],
    "materials": ["Material 1", "Material 2"],
    "activities": [
        "Interactive Activity 1 Description",
        "Interactive Activity 2 Description"
    ],
    "assessmentPlan": "Description of how to check understanding",
    "teacher_script": "Optional additional guidance for the AI"
}
"""

        task_instructions = f"""
**YOUR TASK:**
Generate a comprehensive lesson script suitable for an AI to present on a virtual whiteboard.
1. **Content**: The 'content' field must be rich, educational, and easy to read. Use Markdown.
2. **Tone**: Encouraging, direct, and adapted to {grade_level} students.
3. **Structure**: Cover Introduction, Core Concepts, Examples, and Conclusion.
4. **Context**: Use the provided RAG context to ensure accuracy and alignment with the Ethiopian curriculum.

**OUTPUT FORMAT:**
Return ONLY a valid JSON object matching the structure below:
{json_structure}
"""

        llm_request = LLMRequest(
            prompt=f"{role_definition}\n\n{context_section}\n\n{task_instructions}",
            user_id=request.user.id,
            user_role=UserRole(request.user.role),
            task_type=TaskType.LESSON_PLANNING, # Using LESSON_PLANNING for similar complexity
            complexity=TaskComplexity.ADVANCED,
            system_prompt="You are an expert AI Teacher. Generate a structured lesson JSON.",
            temperature=0.7,
            max_tokens=4000,
        )
        
        # Process Request
        lesson_data = llm_router.process_json_request(llm_request)
        
        # Augment with metadata
        lesson_data['subject'] = subject
        lesson_data['gradeLevel'] = grade_level
        lesson_data['duration'] = duration
        if rag_enabled:
             lesson_data['rag_status'] = 'success'
             lesson_data['curriculum_sources'] = list(set(curriculum_sources))
        
        return Response(lesson_data)

    except Exception as e:
        logger.error(f"AI Teacher lesson generation error: {e}", exc_info=True)
        # return structured error
        return Response(
            {'error': 'Failed to generate lesson content'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def student_insights_view(request):
    """Generate AI insights for student progress using real-time data and RAG."""
    
    student_data_payload = request.data.get('student', {})
    student_id = student_data_payload.get('id')
    
    if not student_id:
        return Response(
            {'error': 'Student ID is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    try:
        student = User.objects.get(id=student_id)
    except User.DoesNotExist:
        return Response(
            {'error': 'Student not found'},
            status=status.HTTP_404_NOT_FOUND
        )

    # 1. Fetch Real Academic Data
    recent_grades = StudentGrade.objects.filter(student=student).order_by('-created_at')[:10]
    
    # Calculate performance metrics
    subject_performance = {}
    weak_areas = []
    strong_areas = []
    
    for grade in recent_grades:
        if grade.subject not in subject_performance:
            subject_performance[grade.subject] = []
        if grade.percentage is not None:
            subject_performance[grade.subject].append(grade.percentage)
            
    for subject, scores in subject_performance.items():
        avg_score = sum(scores) / len(scores) if scores else 0
        if avg_score < 60:
            weak_areas.append(f"{subject} (Avg: {avg_score:.1f}%)")
        elif avg_score >= 80:
            strong_areas.append(f"{subject} (Avg: {avg_score:.1f}%)")

    # 2. Fetch Real Engagement Data
    recent_sessions = EngagementSession.objects.filter(student=student).order_by('-started_at')[:5]
    engagement_summary = "No recent engagement data."
    if recent_sessions:
        avg_engagement = sum(s.engagement_score for s in recent_sessions) / len(recent_sessions)
        dominant_emotions = [s.dominant_expression for s in recent_sessions]
        engagement_summary = f"Average Engagement Score: {avg_engagement:.1f}/100. Recent emotions: {', '.join(dominant_emotions)}."

    # 3. RAG Integration for Remedial Resources
    remedial_context = ""
    if weak_areas:
        # Pick the first weak subject to find resources for
        weak_subject_name = weak_areas[0].split(' (')[0]
        try:
            # Search for remedial strategies/resources for this subject
            docs = query_curriculum_documents(
                grade=student.grade_level or "Grade 9", # Fallback if not set
                subject=weak_subject_name,
                query=f"remedial strategies common misconceptions {weak_subject_name}",
                region=student.region or "Addis Ababa", # Fallback to default
                top_k=2
            )
            if docs:
                remedial_context = "\nRecommended Remedial Resources from Curriculum:\n"
                for doc in docs:
                    remedial_context += f"- {doc.get('content')[:200]}...\n"
        except Exception as e:
            logger.warning(f"Failed to fetch remedial resources: {e}")

    # 4. Build Enhanced Prompt
    prompt = f"""Analyze the following student performance data and provide actionable, culturally-sensitive insights for the teacher.

Student Profile:
- Name: {student.first_name} {student.last_name} ({student.username})
- Grade Level: {student.grade_level or 'N/A'}

Academic Performance (Real-time Data):
- Recent Grades: {', '.join([f"{g.subject}: {g.percentage:.1f}%" for g in recent_grades if g.percentage is not None])}
- Strong Areas: {', '.join(strong_areas) if strong_areas else 'None identified yet'}
- Areas for Improvement: {', '.join(weak_areas) if weak_areas else 'None identified yet'}

Behavioral & Engagement Profile:
- {engagement_summary}

Curriculum Context (Remedial Resources):
{remedial_context}

Task:
Provide a structured JSON response with:
1. "summary": A brief, professional summary of the student's current status (2-3 sentences).
2. "strengths": List of 2-3 key strengths based on data.
3. "areasForImprovement": List of 2-3 specific areas needing attention.
4. "interventionSuggestions": List of 3 concrete, actionable intervention strategies for the teacher. If remedial resources are provided, incorporate them.
5. "recommendedResources": List of 2 specific resource types or topics to review.

Ensure the tone is supportive and constructive."""
    
    system_prompt = "You are an expert educational psychologist and data analyst specializing in Ethiopian education. Provide actionable, culturally-sensitive insights."
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.STUDENT_INSIGHTS,
        complexity=TaskComplexity.ADVANCED,
        system_prompt=system_prompt,
        temperature=0.6,
        max_tokens=1500,
    )
    
    try:
        # Use robust JSON processing
        insights = llm_router.process_json_request(llm_request)
        return Response(insights)
        
    except Exception as e:
        logger.error(f"Student insights JSON error: {e}")
        # Fallback to structured error response
        return Response({
            'summary': "AI analysis failed to generate structured insights. Please try again.",
            'interventionSuggestions': [],
            'error': str(e)
        })
    


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_rubric_view(request):
    """
    Generate a rubric based on topic, grade level, and subject.
    Uses RAG to align with Ethiopian curriculum if available.
    """
    try:
        # Extract parameters
        topic = request.data.get('topic', '')
        grade_level = request.data.get('grade_level', '')
        subject = request.data.get('subject', '')
        rubric_type = request.data.get('rubric_type', 'analytic')
        num_criteria = int(request.data.get('num_criteria', 5))
        performance_levels = request.data.get('performance_levels', [])
        tone_preference = request.data.get('tone', 'Professional')
        weighting_enabled = request.data.get('weighting', False)
        multimodal_assessment = request.data.get('multimodal', False)
        learning_objectives = request.data.get('learning_objectives', [])
        moe_standard_id = request.data.get('moe_standard_id', '')
        
        # RAG parameters
        use_vector_store = request.data.get('use_vector_store', True)
        chapter_number = request.data.get('chapter_number')
        chapter_range_start = request.data.get('chapter_range_start')
        chapter_range_start = request.data.get('chapter_range_start')
        chapter_range_end = request.data.get('chapter_range_end')
        region = request.data.get('region', None)
        
        if not topic:
            return Response(
                {'error': 'Topic is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
            
        logger.info(f"📝 Generating rubric for: {topic} ({grade_level} {subject})")
        
        # Initialize RAG context
        curriculum_context = None
        rag_context = None
        curriculum_sources = []
        
        if use_vector_store and subject and grade_level:
            try:
                # Analyze topic for chapter detection
                chapter_info = RubricGeneratorRAGEnhancer.analyze_topic_for_chapter(
                    topic=topic,
                    subject=subject,
                    grade_level=grade_level
                )
                
                # Parse chapter parameters
                chapter_range = None
                if chapter_range_start and chapter_range_end:
                    chapter_range = (int(chapter_range_start), int(chapter_range_end))
                elif chapter_number:
                    chapter_number = int(chapter_number)
                elif chapter_info:
                    chapter_number = chapter_info['number']
                    logger.info(f"📖 Using detected chapter: {chapter_number}")
                
                # Build enhanced query
                query_text = RubricGeneratorRAGEnhancer.build_rubric_query(
                    topic=topic,
                    subject=subject,
                    grade_level=grade_level,
                    learning_objectives=learning_objectives,
                    chapter_info=chapter_info
                )
                
                # Query curriculum documents with chapter awareness
                from rag.services import query_curriculum_documents
                
                chapter_param = str(chapter_info['number']) if chapter_info else (str(chapter_number) if chapter_number else None)
                
                documents = query_curriculum_documents(
                    subject=subject,
                    grade=grade_level,
                    query=query_text,
                    chapter=chapter_param,

                    region=region,
                    top_k=10
                )
                
                logger.info(f"📚 Retrieved {len(documents)} curriculum documents for rubric generation")
                
                # Format context using enhancer
                rag_context, curriculum_sources = RubricGeneratorRAGEnhancer.format_rubric_context(
                    documents=documents,
                    topic=topic,
                    chapter_info=chapter_info,
                    max_chars=8000
                )
                
                # Extract content from formatted context
                extracted_objectives = RubricGeneratorRAGEnhancer.extract_learning_objectives(rag_context)
                extracted_concepts = RubricGeneratorRAGEnhancer.extract_key_concepts(rag_context)
                extracted_standards = RubricGeneratorRAGEnhancer.extract_standards(rag_context, subject, grade_level)
                
                # Auto-populate learning objectives if found in curriculum
                if not learning_objectives and extracted_objectives:
                    learning_objectives = extracted_objectives
                    logger.info(f"📚 Auto-populated {len(learning_objectives)} learning objectives from curriculum")
                
                # Auto-populate MoE standard if found
                if not moe_standard_id and extracted_standards:
                    moe_standard_id = extracted_standards[0]
                    logger.info(f"📚 Auto-populated MoE standard: {moe_standard_id}")
                
                # Build curriculum context structure
                curriculum_context = {
                    'success': True if documents else False,
                    'context_available': True if documents else False,
                    'learning_objectives': extracted_objectives,
                    'key_concepts': extracted_concepts,
                    'standards': extracted_standards,
                    'chapter_context': {
                        'chapter_number': chapter_info['number'] if chapter_info else chapter_number,
                        'chapter_title': chapter_info.get('title', f'Chapter {chapter_info["number"]}') if chapter_info else None,
                        'success': True
                    } if (chapter_info or chapter_number) else None,
                    'rag_context': rag_context,
                    'sources': curriculum_sources
                }
            
            except Exception as e:
                logger.warning(f"⚠️ Curriculum context extraction failed: {e}. Continuing without RAG.")
                curriculum_context = None
        
        # Default performance levels if not provided
        if not performance_levels:
            if rubric_type == 'analytic':
                performance_levels = ['Exemplary', 'Proficient', 'Developing', 'Beginning']
            elif rubric_type == 'holistic':
                performance_levels = ['Outstanding', 'Good', 'Satisfactory', 'Needs Improvement']
            elif rubric_type == 'single_point':
                performance_levels = ['Concerns', 'Criteria', 'Advanced']
            else:  # checklist
                performance_levels = ['Yes', 'No', 'Partial']
        
        # Build comprehensive prompt using the Enhancer
        prompt = RubricGeneratorRAGEnhancer.build_rubric_prompt(
            topic=topic,
            subject=subject,
            grade_level=grade_level,
            rubric_type=rubric_type,
            num_criteria=num_criteria,
            performance_levels=performance_levels,
            tone_preference=tone_preference,
            weighting_enabled=weighting_enabled,
            multimodal_assessment=multimodal_assessment,
            learning_objectives=learning_objectives,
            moe_standard_id=moe_standard_id,
            rag_context=rag_context,
            curriculum_context=curriculum_context
        )
        
        system_prompt = """You are an expert assessment designer specializing in Ethiopian education standards. 
Your rubrics are known for clarity, fairness, constructive alignment with learning objectives, and actionable feedback.
You always use concrete, observable language and avoid vague descriptors.
You follow best practices in criterion-referenced assessment and standards-based grading.
You MUST output valid JSON only."""
        
        # Use high complexity for alignment and weighting tasks (Tier 1 LLM)
        complexity = TaskComplexity.ADVANCED if (learning_objectives or weighting_enabled or moe_standard_id) else TaskComplexity.MEDIUM
        
        llm_request = LLMRequest(
            prompt=prompt,
            user_id=request.user.id,
            user_role=UserRole(request.user.role),
            task_type=TaskType.RUBRIC_GENERATION,
            complexity=complexity,
            system_prompt=system_prompt,
            temperature=0.6,
            max_tokens=5000,
        )
        
        try:
            # Use robust JSON processing
            rubric = llm_router.process_json_request(llm_request)
            
            # Validate and enhance rubric structure
            if 'criteria' not in rubric:
                raise ValueError("Invalid rubric structure: missing criteria")
            
            # Calculate total points if not provided
            if 'total_points' not in rubric:
                rubric['total_points'] = 100
            
            # Validate alignment if objectives provided
            alignment_score = 0.0
            if learning_objectives:
                alignment_score = validate_rubric_alignment(rubric, learning_objectives)
                rubric['alignment_validated'] = True
                rubric['alignment_score'] = alignment_score
            
            # Add metadata
            rubric['rubric_type'] = rubric_type
            rubric['grade_level'] = grade_level
            rubric['subject'] = subject
            rubric['weighting_enabled'] = weighting_enabled
            rubric['multimodal_assessment'] = multimodal_assessment
            rubric['performance_levels'] = performance_levels
            
            # NEW: Enhance with curriculum context if available
            if curriculum_context and curriculum_context.get('success'):
                rubric = RubricGeneratorRAGEnhancer.enhance_rubric_with_context(
                    rubric=rubric,
                    curriculum_context=curriculum_context
                )
                logger.info("✨ Rubric enhanced with curriculum context")
            
            return Response(rubric)
        
        except Exception as e:
            logger.error(f"Rubric generation error: {e}")
            return Response(
                {
                    'error': 'Failed to generate valid rubric. Please try again.',
                    'details': str(e)
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
            
    except Exception as e:
        logger.error(f"Rubric view error: {e}", exc_info=True)
        return Response(
            {'error': 'An unexpected error occurred', 'details': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def export_rubric_view(request):
    """
    Export a rubric (not saved) as PDF or DOCX.
    Accepts rubric data and format in request body.
    """
    try:
        rubric_data = request.data
        export_format = rubric_data.get('format', 'pdf')
        
        if export_format not in ['pdf', 'docx']:
            return Response(
                {'error': 'Invalid format. Must be pdf or docx'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if export_format == 'pdf':
            return _export_rubric_as_pdf(rubric_data)
        else:
            return _export_rubric_as_docx(rubric_data)
            
    except Exception as e:
        logger.error(f"Rubric export error: {e}")
        return Response(
            {'error': f'Failed to export rubric: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


def _export_rubric_as_pdf(rubric_data):
    """Helper function to export rubric as PDF"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.units import inch
        from io import BytesIO
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
        )
        elements.append(Paragraph(rubric_data.get('title', 'Rubric'), title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Metadata
        metadata_style = styles['Normal']
        elements.append(Paragraph(f"<b>Topic:</b> {rubric_data.get('topic', 'N/A')}", metadata_style))
        elements.append(Paragraph(f"<b>Grade Level:</b> {rubric_data.get('grade_level', 'N/A')}", metadata_style))
        if rubric_data.get('subject'):
            elements.append(Paragraph(f"<b>Subject:</b> {rubric_data['subject']}", metadata_style))
        elements.append(Paragraph(f"<b>Rubric Type:</b> {rubric_data.get('rubric_type', 'analytic').replace('_', ' ').title()}", metadata_style))
        elements.append(Paragraph(f"<b>Total Points:</b> {rubric_data.get('total_points', 100)}", metadata_style))
        elements.append(Spacer(1, 0.3*inch))
        
        # Learning Objectives
        learning_objectives = rubric_data.get('learning_objectives', [])
        if learning_objectives:
            elements.append(Paragraph("<b>Learning Objectives:</b>", styles['Heading2']))
            for i, obj in enumerate(learning_objectives, 1):
                elements.append(Paragraph(f"{i}. {obj}", metadata_style))
            elements.append(Spacer(1, 0.2*inch))
        
        # Rubric Table
        elements.append(Paragraph("<b>Assessment Criteria:</b>", styles['Heading2']))
        elements.append(Spacer(1, 0.1*inch))
        
        # Build table data
        criteria = rubric_data.get('criteria', [])
        if criteria and len(criteria) > 0:
            # Get performance levels from first criterion
            perf_levels = criteria[0].get('performanceLevels', [])
            
            # Calculate column widths dynamically
            num_perf_levels = len(perf_levels)
            has_weight = rubric_data.get('weighting_enabled')
            
            # Available width (A4 width minus margins)
            available_width = 7.5 * inch  # A4 width minus margins
            
            # Allocate widths
            criterion_width = 1.5 * inch  # Fixed width for criterion column
            weight_width = 0.5 * inch if has_weight else 0
            remaining_width = available_width - criterion_width - weight_width
            perf_level_width = remaining_width / num_perf_levels if num_perf_levels > 0 else 2 * inch
            
            # Build column widths list
            col_widths = [criterion_width]
            if has_weight:
                col_widths.append(weight_width)
            col_widths.extend([perf_level_width] * num_perf_levels)
            
            # Header row
            header = ['Criterion']
            if has_weight:
                header.append('Weight')
            header.extend([level.get('level', '') for level in perf_levels])
            
            table_data = [header]
            
            # Data rows - wrap text in Paragraphs for better formatting
            for criterion in criteria:
                row = [Paragraph(criterion.get('criterion', ''), styles['Normal'])]
                if has_weight:
                    row.append(Paragraph(f"{criterion.get('weight', 0)}%", styles['Normal']))
                
                for level in criterion.get('performanceLevels', []):
                    desc = level.get('description', '')
                    points = level.get('points', '')
                    cell_text = f"{desc}<br/><b>({points} pts)</b>"
                    row.append(Paragraph(cell_text, styles['Normal']))
                
                table_data.append(row)
            
            # Create table with explicit column widths
            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ('WORDWRAP', (0, 0), (-1, -1), True),
            ]))
            
            elements.append(table)
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        # Return PDF response
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        filename = rubric_data.get('title', 'rubric').replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response
        
    except ImportError:
        return Response(
            {'error': 'PDF export requires reportlab library. Please install it: pip install reportlab'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return Response(
            {'error': f'Failed to export PDF: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


def _export_rubric_as_docx(rubric_data):
    """Helper function to export rubric as DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading(rubric_data.get('title', 'Rubric'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Topic: {rubric_data.get('topic', 'N/A')}")
        doc.add_paragraph(f"Grade Level: {rubric_data.get('grade_level', 'N/A')}")
        if rubric_data.get('subject'):
            doc.add_paragraph(f"Subject: {rubric_data['subject']}")
        doc.add_paragraph(f"Rubric Type: {rubric_data.get('rubric_type', 'analytic').replace('_', ' ').title()}")
        doc.add_paragraph(f"Total Points: {rubric_data.get('total_points', 100)}")
        doc.add_paragraph()
        
        # Learning Objectives
        learning_objectives = rubric_data.get('learning_objectives', [])
        if learning_objectives:
            doc.add_heading('Learning Objectives:', 2)
            for i, obj in enumerate(learning_objectives, 1):
                doc.add_paragraph(f"{i}. {obj}")
            doc.add_paragraph()
        
        # Rubric Table
        doc.add_heading('Assessment Criteria:', 2)
        
        criteria = rubric_data.get('criteria', [])
        if criteria and len(criteria) > 0:
            # Get performance levels from first criterion
            perf_levels = criteria[0].get('performanceLevels', [])
            
            # Calculate columns
            num_cols = 1  # Criterion column
            if rubric_data.get('weighting_enabled'):
                num_cols += 1  # Weight column
            num_cols += len(perf_levels)  # Performance level columns
            
            # Create table
            table = doc.add_table(rows=1 + len(criteria), cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            col_idx = 0
            header_cells[col_idx].text = 'Criterion'
            col_idx += 1
            
            if rubric_data.get('weighting_enabled'):
                header_cells[col_idx].text = 'Weight'
                col_idx += 1
            
            for level in perf_levels:
                header_cells[col_idx].text = level.get('level', '')
                col_idx += 1
            
            # Data rows
            for row_idx, criterion in enumerate(criteria, 1):
                row_cells = table.rows[row_idx].cells
                col_idx = 0
                
                row_cells[col_idx].text = criterion.get('criterion', '')
                col_idx += 1
                
                if rubric_data.get('weighting_enabled'):
                    row_cells[col_idx].text = f"{criterion.get('weight', 0)}%"
                    col_idx += 1
                
                for level in criterion.get('performanceLevels', []):
                    desc = level.get('description', '')
                    points = level.get('points', '')
                    row_cells[col_idx].text = f"{desc}\n({points} pts)"
                    col_idx += 1
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return DOCX response
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = rubric_data.get('title', 'rubric').replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response
        
    except ImportError:
        return Response(
            {'error': 'DOCX export requires python-docx library. Please install it: pip install python-docx'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return Response(
            {'error': f'Failed to export DOCX: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def extract_curriculum_content_view(request):
    """
    Extract curriculum content for rubric generation.
    Returns learning objectives, MoE standards, and suggested criteria.
    """
    try:
        # Extract parameters
        subject = request.data.get('subject', '')
        grade_level = request.data.get('grade_level', '')
        topic = request.data.get('topic', '')
        chapter_input = request.data.get('chapter_input', '')
        topic = request.data.get('topic', '')
        chapter_input = request.data.get('chapter_input', '')
        document_type = request.data.get('document_type', 'essay')  # NEW: Document type for topic suggestions
        region = request.data.get('region', None)
        
        if not subject or not grade_level:
            return Response(
                {'error': 'Subject and grade level are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if not rubric_rag_enhancer:
            return Response(
                {'error': 'RAG enhancer not available'},
                status=status.HTTP_503_SERVICE_UNAVAILABLE
            )
        
        # Analyze topic for chapter detection
        chapter_info = RubricGeneratorRAGEnhancer.analyze_topic_for_chapter(
            topic=topic,
            subject=subject,
            grade_level=grade_level
        )
        
        # Parse chapter input
        chapter_number = None
        chapter_range = None
        
        if chapter_input:
            import re
            # Try to parse chapter range (e.g., "3-5", "Chapters 3-5")
            range_match = re.search(r'(\d+)\s*-\s*(\d+)', chapter_input)
            if range_match:
                chapter_range = (int(range_match.group(1)), int(range_match.group(2)))
            else:
                # Try to parse single chapter (e.g., "5", "Chapter 5")
                single_match = re.search(r'\d+', chapter_input)
                if single_match:
                    chapter_number = int(single_match.group())
        
        # Use detected chapter if no explicit input
        if not chapter_number and not chapter_range and chapter_info:
            chapter_number = chapter_info['number']
            logger.info(f"📖 Using detected chapter: {chapter_number}")
        
        logger.info(f"📚 Extracting curriculum content: {subject} {grade_level}, Chapter: {chapter_number or chapter_range or chapter_input}")
        
        # Build enhanced query for RAG
        query_text = RubricGeneratorRAGEnhancer.build_rubric_query(
            topic=topic or f"{subject} curriculum content",
            subject=subject,
            grade_level=grade_level,
            learning_objectives=[],
            chapter_info=chapter_info
        )
        
        # Query curriculum documents with chapter awareness
        from rag.services import query_curriculum_documents
        
        chapter_param = str(chapter_info['number']) if chapter_info else (str(chapter_number) if chapter_number else None)
        
        documents = query_curriculum_documents(
            subject=subject,
            grade=grade_level,
            query=query_text,
            chapter=chapter_param,
            region=region,
            top_k=10
        )
        
        logger.info(f"📚 Retrieved {len(documents)} curriculum documents")
        
        # Format context using enhancer
        rag_context, curriculum_sources = RubricGeneratorRAGEnhancer.format_rubric_context(
            documents=documents,
            topic=topic or f"{subject} curriculum content",
            chapter_info=chapter_info,
            max_chars=8000
        )
        
        # Extract content from formatted context
        learning_objectives = RubricGeneratorRAGEnhancer.extract_learning_objectives(rag_context)
        key_concepts = RubricGeneratorRAGEnhancer.extract_key_concepts(rag_context)
        standards = RubricGeneratorRAGEnhancer.extract_standards(rag_context, subject, grade_level)
        
        # If no key concepts extracted, generate from objectives and topic
        if not key_concepts and learning_objectives:
            # Extract key terms from objectives
            import re
            key_terms = set()
            for obj in learning_objectives[:5]:
                # Extract meaningful words (5+ chars, not common words)
                words = re.findall(r'\b[a-zA-Z]{5,}\b', obj)
                key_terms.update([w.lower() for w in words if w.lower() not in ['about', 'their', 'using', 'given', 'words', 'should', 'students', 'will', 'able']])
            key_concepts = list(key_terms)[:10]
            logger.info(f"📝 Generated {len(key_concepts)} key concepts from objectives")
        
        # Build curriculum context structure
        curriculum_context = {
            'success': True if documents else False,
            'learning_objectives': learning_objectives,
            'key_concepts': key_concepts,
            'standards': standards,
            'chapter_context': {
                'chapter_number': chapter_info['number'] if chapter_info else chapter_number,
                'chapter_title': chapter_info.get('title', f'Chapter {chapter_info["number"]}') if chapter_info else None
            } if (chapter_info or chapter_number) else None
        }
        
        if not curriculum_context.get('success'):
            return Response(
                {
                    'success': False,
                    'error': curriculum_context.get('error', 'Failed to extract content'),
                    'learning_objectives': [],
                    'standards': [],
                    'key_concepts': [],
                    'suggested_criteria_count': 5
                },
                status=status.HTTP_200_OK
            )
        
        # Extract content (already extracted above, just reference from curriculum_context)
        extracted_objectives = curriculum_context.get('learning_objectives', [])
        extracted_standards = curriculum_context.get('standards', [])
        extracted_concepts = curriculum_context.get('key_concepts', [])
        
        # Suggest number of criteria based on content complexity
        suggested_criteria_count = min(max(len(extracted_concepts), 3), 8)
        
        # Generate 5 topic suggestions if requested
        suggested_topics = []
        topic_requested = request.data.get('suggest_topic', False)
        
        logger.info(f"🔍 Topic generation: requested={topic_requested}, key_concepts_count={len(extracted_concepts)}, objectives_count={len(extracted_objectives)}")
        
        # Generate topics if requested and we have either concepts or objectives
        if topic_requested and (extracted_concepts or extracted_objectives):
            try:
                # Build context for topic generation
                chapter_display = ""
                if chapter_number:
                    chapter_display = f"Chapter {chapter_number}"
                elif chapter_range:
                    chapter_display = f"Chapters {chapter_range[0]}-{chapter_range[1]}"
                
                # Map document type to assignment type label
                doc_type_map = {
                    'essay': 'Essay',
                    'examination': 'Examination',
                    'project': 'Project',
                    'group_work': 'Group Work',
                    'lab_report': 'Lab Report',
                    'presentation': 'Presentation',
                    'homework': 'Homework',
                    'quiz': 'Quiz',
                    'creative_writing': 'Creative Writing',
                    'critical_analysis': 'Critical Analysis'
                }
                
                assignment_type = doc_type_map.get(document_type, 'Essay')
                
                # Create prompt for LLM to generate 5 topics specific to document type
                topic_prompt = f"""Generate 5 diverse and engaging {assignment_type} topics for {subject} at {grade_level} level.

Context:
- Subject: {subject}
- Grade Level: {grade_level}
- Document Type: {assignment_type}
{f"- Chapter/Unit: {chapter_display}" if chapter_display else ""}
{f"- Key Concepts: {', '.join(extracted_concepts[:5])}" if extracted_concepts else ""}
{f"- Learning Objectives: {', '.join(extracted_objectives[:3])}" if extracted_objectives else ""}

Requirements:
1. ALL 5 topics MUST be {assignment_type} topics ONLY - do not mix with other assignment types
2. Each topic should start with "{assignment_type}:" followed by a specific, engaging title
3. Topics should be appropriate for {grade_level} students and aligned with the learning objectives
4. Make topics engaging and relevant to students' real-world experiences in Ethiopian context
5. Each topic should directly relate to the key concepts and learning objectives provided
6. Use professional yet accessible language suitable for educational rubrics

Format: Return ONLY 5 {assignment_type} topics, one per line, numbered 1-5. Each must follow this structure:
{assignment_type}: [Specific, Engaging Title Related to Content]

Example format for {assignment_type}:
1. {assignment_type}: [Specific title related to key concepts]
2. {assignment_type}: [Another specific title]
3. {assignment_type}: [Third specific title]
4. {assignment_type}: [Fourth specific title]
5. {assignment_type}: [Fifth specific title]

IMPORTANT: Generate ONLY {assignment_type} topics. Do NOT include Essay, Research Paper, Project, Presentation, or Case Study unless the document type is specifically one of those."""

                # Use LLM to generate topics
                from ai_tools.llm.llm_router import llm_router
                
                logger.info(f"🤖 Calling LLM for topic generation with {len(topic_prompt)} char prompt")
                
                llm_response = llm_router.generate_text(
                    prompt=topic_prompt,
                    max_tokens=1000,
                    temperature=0.8,  # Higher temperature for more creative topics
                    tier_preference='tier2'  # Use Gemini for cost-effective generation
                )
                
                logger.info(f"🤖 LLM Response: success={llm_response.get('success') if llm_response else False}")
                
                if llm_response and llm_response.get('success'):
                    response_text = llm_response.get('response', '')
                    logger.info(f"📝 LLM generated text: {response_text[:200]}...")
                    
                    # Parse the numbered list
                    lines = response_text.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        # Remove numbering (1., 2., etc.) and clean up
                        if line and (line[0].isdigit() or line.startswith('-')):
                            # Remove leading number, dot, dash, and whitespace
                            topic = line.lstrip('0123456789.-) ').strip()
                            if topic and len(topic) > 10:  # Ensure it's a valid topic
                                suggested_topics.append(topic)
                    
                    # Ensure we have exactly 5 topics
                    suggested_topics = suggested_topics[:5]
                    
                    logger.info(f"💡 Generated {len(suggested_topics)} topic suggestions using LLM")
                    
                    # If LLM didn't generate enough topics, use fallback
                    if len(suggested_topics) < 5:
                        logger.warning(f"LLM only generated {len(suggested_topics)} topics, using fallback for remaining")
                        fallback_topics = RubricGeneratorRAGEnhancer.generate_topic_suggestions(
                            content=rag_context,
                            subject=subject,
                            grade_level=grade_level,
                            chapter_info=chapter_info,
                            num_suggestions=5 - len(suggested_topics)
                        )
                        suggested_topics.extend(fallback_topics)
                else:
                    # Fallback: Use enhancer to generate topics
                    error_msg = llm_response.get('error', 'Unknown error') if llm_response else 'No response'
                    logger.warning(f"LLM topic generation failed: {error_msg}, using enhancer fallback")
                    suggested_topics = RubricGeneratorRAGEnhancer.generate_topic_suggestions(
                        content=rag_context,
                        subject=subject,
                        grade_level=grade_level,
                        chapter_info=chapter_info,
                        num_suggestions=5
                    )
                    
            except Exception as e:
                logger.warning(f"Failed to generate topic suggestions: {e}")
                # Fallback: Use enhanced topic generation
                try:
                    suggested_topics = RubricGeneratorRAGEnhancer.generate_topic_suggestions(
                        content=rag_context,
                        subject=subject,
                        grade_level=grade_level,
                        chapter_info=chapter_info,
                        num_suggestions=5
                    )
                    logger.info(f"💡 Generated {len(suggested_topics)} topics using enhanced fallback")
                except Exception as fallback_error:
                    logger.error(f"Enhanced fallback also failed: {fallback_error}")
                    # Last resort: Generate simple topics
                    for i, concept in enumerate(extracted_concepts[:5], 1):
                        if chapter_display:
                            suggested_topics.append(f"{subject} {grade_level}: {chapter_display} - {concept}")
                        else:
                            suggested_topics.append(f"{subject} {grade_level}: {concept}")
        
        logger.info(f"✅ Extracted: {len(extracted_objectives)} objectives, {len(extracted_standards)} standards, {len(extracted_concepts)} concepts")
        
        return Response({
            'success': True,
            'learning_objectives': extracted_objectives,
            'standards': extracted_standards,
            'key_concepts': extracted_concepts,
            'suggested_criteria_count': suggested_criteria_count,
            'suggested_topics': suggested_topics if suggested_topics else [],
            'chapter_context': {
                'chapter_number': curriculum_context.get('chapter_context', {}).get('chapter_number'),
                'chapter_title': curriculum_context.get('chapter_context', {}).get('chapter_title'),
            } if curriculum_context.get('chapter_context') else None
        })
    
    except Exception as e:
        logger.error(f"❌ Content extraction error: {e}", exc_info=True)
        return Response(
            {'error': 'Failed to extract curriculum content', 'details': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def extract_chapters_view(request):
    """
    Extract available chapters from the curriculum vector store.
    """
    try:
        # Extract parameters
        subject = request.data.get('subject', '')
        grade_level = request.data.get('grade_level', '')
        region = request.data.get('region', None)
        stream = request.data.get('stream', None)
        
        if not subject or not grade_level:
            return Response(
                {'error': 'Subject and grade level are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
            
        logger.info(f"📚 Extracting chapters for: {grade_level} {subject} ({region})")
        
        from rag.services import get_available_chapters
        
        chapters = get_available_chapters(
            grade=grade_level,
            subject=subject,
            region=region,
            stream=stream
        )
        
        return Response({
            'success': True,
            'chapters': chapters
        })
        
    except Exception as e:
        logger.error(f"❌ Error extracting chapters: {str(e)}")
        return Response(
            {'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_assignment_description_view(request):
    """
    Generate a detailed assignment description based on topic, document type, and context.
    Returns a comprehensive description suitable for rubrics and grading.
    """
    try:
        # Extract parameters
        topic = request.data.get('topic', '')
        document_type = request.data.get('document_type', 'essay')
        subject = request.data.get('subject', '')
        grade_level = request.data.get('grade_level', '')
        learning_objectives = request.data.get('learning_objectives', [])
        
        if not topic:
            return Response(
                {'error': 'Topic is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Map document type to assignment type label
        doc_type_map = {
            'essay': 'Essay',
            'examination': 'Examination',
            'project': 'Project',
            'group_work': 'Group Work',
            'lab_report': 'Lab Report',
            'presentation': 'Presentation',
            'homework': 'Homework',
            'quiz': 'Quiz',
            'creative_writing': 'Creative Writing',
            'critical_analysis': 'Critical Analysis'
        }
        
        assignment_type = doc_type_map.get(document_type, 'Essay')
        
        # Build comprehensive prompt for description generation
        description_prompt = f"""You are an expert educational content generator. Your task is to write detailed assignment descriptions based ONLY on the provided topic and parameters. Do NOT ask for additional text or context.

Generate a detailed, professional assignment description for the following {assignment_type}.

Assignment Details:
- Topic: {topic}
- Document Type: {assignment_type}
{f"- Subject: {subject}" if subject else ""}
{f"- Grade Level: {grade_level}" if grade_level else ""}
{f"- Learning Objectives: {', '.join(learning_objectives[:3])}" if learning_objectives else ""}

Requirements:
1. Create a clear, comprehensive description (150-250 words) that explains:
   - What students are expected to do
   - The scope and focus of the assignment
   - Key requirements and expectations
   - Learning goals and outcomes
   - Any specific guidelines or constraints

2. Use professional, educational language appropriate for {grade_level if grade_level else 'secondary school'} level

3. Make it specific to the {assignment_type} format and the topic provided

4. Include Ethiopian educational context where relevant

5. Be clear, actionable, and suitable for use in rubrics and grading

Format: Return ONLY the description text, no additional commentary or labels. Write in a professional, instructional tone.

Example structure:
"Students will [main task]. This {assignment_type.lower()} should [specific requirements]. The work must demonstrate [key skills/concepts]. Students are expected to [expectations]. The final submission should [format/length requirements]."

Generate the description now based on your knowledge (do not ask for input text):"""
        
        # Use LLM to generate description
        from ai_tools.llm.llm_router import llm_router
        
        logger.info(f"🤖 Generating assignment description for: {topic}")
        
        llm_response = llm_router.generate_text(
            prompt=description_prompt,
            max_tokens=1000,
            temperature=0.7,  # Balanced creativity and consistency
            tier_preference='tier2'  # Use Gemini for cost-effective generation
        )
        
        if llm_response and llm_response.get('success'):
            description = llm_response.get('response', '').strip()
            
            # Clean up any unwanted prefixes or formatting
            description = description.replace('Description:', '').strip()
            description = description.replace('Assignment Description:', '').strip()
            
            # Ensure reasonable length
            if len(description) < 50:
                logger.warning(f"Generated description too short ({len(description)} chars), using fallback")
                description = f"Students will complete a {assignment_type.lower()} on the topic: {topic}. "
                description += f"This assignment should demonstrate understanding of the key concepts and meet the learning objectives. "
                if subject:
                    description += f"The work should be appropriate for {grade_level} {subject} students. "
                description += f"Students are expected to submit a well-structured, thoughtful {assignment_type.lower()} that addresses all requirements."
            
            logger.info(f"✅ Generated description: {len(description)} characters")
            
            return Response({
                'success': True,
                'description': description,
                'topic': topic,
                'document_type': document_type
            })
        else:
            # Fallback: Generate basic description
            error_msg = llm_response.get('error', 'Unknown error') if llm_response else 'No response'
            logger.warning(f"LLM description generation failed: {error_msg}, using fallback")
            
            description = f"Students will complete a {assignment_type.lower()} on the topic: {topic}. "
            description += f"This assignment should demonstrate understanding of the key concepts and meet the learning objectives. "
            if subject and grade_level:
                description += f"The work should be appropriate for {grade_level} {subject} students. "
            description += f"Students are expected to submit a well-structured, thoughtful {assignment_type.lower()} that addresses all requirements."
            
            return Response({
                'success': True,
                'description': description,
                'topic': topic,
                'document_type': document_type,
                'fallback_used': True
            })
    
    except Exception as e:
        logger.error(f"❌ Description generation error: {e}", exc_info=True)
        return Response(
            {'error': 'Failed to generate assignment description', 'details': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


def validate_rubric_alignment(rubric, learning_objectives):
    """
    Validate constructive alignment between rubric criteria and learning objectives.
    Returns alignment score (0.0 to 1.0).
    """
    try:
        criteria = rubric.get('criteria', [])
        if not criteria or not learning_objectives:
            return 0.0
        
        # Simple heuristic: check for keyword overlap
        # In production, this could use embeddings for semantic similarity
        objective_keywords = set()
        for obj in learning_objectives:
            words = obj.lower().split()
            objective_keywords.update([w for w in words if len(w) > 4])
        
        aligned_criteria = 0
        for criterion in criteria:
            criterion_text = f"{criterion.get('criterion', '')} {criterion.get('description', '')}".lower()
            criterion_words = set([w for w in criterion_text.split() if len(w) > 4])
            
            # Check for overlap
            overlap = len(objective_keywords.intersection(criterion_words))
            if overlap > 0:
                aligned_criteria += 1
        
        alignment_score = aligned_criteria / len(criteria) if criteria else 0.0
        return round(alignment_score, 2)
        
    except Exception as e:
        logger.error(f"Alignment validation error: {e}")
        return 0.0


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def grade_submission_view(request):
    """
    Enhanced Quick Grader - Grade any type of submission with AI.
    Supports essays, assignments, exams, group work, projects, and more.
    Includes RAG-based curriculum context for accurate grading.
    """
    
    submission_id = request.data.get('submission_id')
    custom_rubric = request.data.get('rubric')  # Optional custom rubric
    custom_text = request.data.get('custom_text')  # Optional custom submission text
    assessment_type = request.data.get('assessment_type', 'essay')  # essay, exam, project, group_work, etc.
    
    # NEW: RAG enhancement parameters
    use_vector_store = request.data.get('use_vector_store', False)
    subject = request.data.get('subject')
    grade_level = request.data.get('grade_level')
    chapter_number = request.data.get('chapter_number')
    chapter_range_start = request.data.get('chapter_range_start')
    chapter_range_start = request.data.get('chapter_range_start')
    chapter_range_end = request.data.get('chapter_range_end')
    region = request.data.get('region', None)
    
    if not submission_id and not custom_text:
        return Response(
            {'error': 'Either submission_id or custom_text is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Fetch submission from database or use custom text
    from academics.models import Submission
    from .essay_grader_enhancer import EssayGraderEnhancer
    
    submission = None
    submission_text = None
    assignment_prompt = ""
    rubric_data = None
    student_name = "Student"
    
    if submission_id:
        try:
            submission = Submission.objects.select_related('assignment', 'student').get(id=submission_id)
            submission_text = submission.submitted_text
            assignment_prompt = submission.assignment.description
            rubric_data = custom_rubric if custom_rubric else submission.assignment.rubric
            student_name = submission.student.username
            
            # Get subject and grade from submission if not provided
            if not subject:
                subject = getattr(submission.assignment, 'subject', None)
            if not grade_level:
                grade_level = getattr(submission.student, 'grade', None)
            
            logger.info(f"📝 Grading submission {submission_id} for student {student_name}")
        
        except Submission.DoesNotExist:
            return Response(
                {'error': 'Submission not found'},
                status=status.HTTP_404_NOT_FOUND
            )
    else:
        # Use custom text
        submission_text = custom_text
        assignment_prompt = request.data.get('assignment_description', '')
        rubric_data = custom_rubric
        logger.info(f"📝 Grading custom text submission ({len(submission_text)} chars)")
    
    if not submission_text:
        return Response(
            {'error': 'Submission has no text content'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    if not rubric_data:
        return Response(
            {'error': 'Rubric is required for grading'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Parse and validate rubric
    parsed_rubric = EssayGraderEnhancer.parse_rubric(rubric_data)
    
    logger.info(f"📋 Rubric parsed: {len(parsed_rubric.get('criteria', []))} criteria, {parsed_rubric.get('total_points')} points")
    
    # NEW: Extract grading context from vector stores if enabled
    grading_context = None
    if use_vector_store and grader_rag_enhancer and subject and grade_level:
        try:
            chapter_range = None
            if chapter_range_start and chapter_range_end:
                chapter_range = (int(chapter_range_start), int(chapter_range_end))
            elif chapter_number:
                chapter_number = int(chapter_number)
            
            grading_context = grader_rag_enhancer.extract_grading_context(
                submission_text=submission_text,
                assignment_description=assignment_prompt,
                subject=subject,
                grade=grade_level,
                chapter_number=chapter_number,
                chapter_range=chapter_range,
                region=region,
                use_vector_store=True
            )
            
            if grading_context.get('success'):
                logger.info(f"📚 Extracted grading context: {len(grading_context.get('reference_content', []))} references")
        
        except Exception as e:
            logger.warning(f"⚠️ Grading context extraction failed: {e}. Continuing without RAG.")
            grading_context = None
    
    # Build enhanced grading prompt - use RAG version if available
    if grading_context and grading_context.get('success') and grader_rag_enhancer:
        prompt = grader_rag_enhancer.build_enhanced_grading_prompt(
            submission_text=submission_text,
            rubric=parsed_rubric,
            assignment_description=assignment_prompt,
            grading_context=grading_context,
            assessment_type=assessment_type,
            grade_level=grade_level
        )
        logger.info("📚 Using RAG-enhanced grading prompt")
    else:
        prompt = EssayGraderEnhancer.build_grading_prompt(
            submission_text=submission_text,
            rubric=parsed_rubric,
            assignment_description=assignment_prompt,
            assessment_type=assessment_type,
            grade_level=grade_level
        )
    
    system_prompt = "You are an experienced Ethiopian educator with expertise in fair, comprehensive assessment. Provide detailed, constructive feedback that helps students learn and improve."
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.GRADING,
        complexity=TaskComplexity.ADVANCED,
        system_prompt=system_prompt,
        temperature=0.5,
        max_tokens=1500,
    )
    
    try:
        # Use robust JSON processing
        graded_result = llm_router.process_json_request(llm_request)
        logger.info(f"✅ Grading completed: {graded_result.get('overallScore', 0)}/{parsed_rubric.get('total_points')}")

    except Exception as e:
        logger.error(f"❌ Grading JSON error: {e}")
        # Fallback for grading failure
        graded_result = {
            'overallScore': 0,
            'overallFeedback': "AI grading failed to generate structured output. Please try again.",
            'criteriaFeedback': [],
            'error': str(e)
        }
        
    # Validate grading result
    is_valid, warnings = EssayGraderEnhancer.validate_grading_result(
        graded_result,
        parsed_rubric
    )
    
    if warnings:
        logger.warning(f"⚠️ Grading warnings: {warnings}")
        graded_result['validation_warnings'] = warnings
    
    # Enhance grading result with metadata
    submission_meta = {
        'text': submission_text,
        'student_id': submission.student.id if submission else None,
        'assignment_id': submission.assignment.id if submission else None
    }
    
    enhanced_result = EssayGraderEnhancer.enhance_grading_result(
        graded_result,
        parsed_rubric,
        submission_meta
    )
    
    # NEW: Add RAG-based accuracy validation if available
    if grading_context and grading_context.get('success') and grader_rag_enhancer:
        accuracy_validation = grader_rag_enhancer.validate_submission_accuracy(
            submission_text=submission_text,
            grading_context=grading_context
        )
        
        if accuracy_validation.get('validated'):
            enhanced_result['curriculum_accuracy'] = accuracy_validation
            logger.info(f"✅ Curriculum accuracy: {accuracy_validation.get('coverage_percentage', 0)}%")
    
    # Save grade and feedback to database (only if submission exists)
    if submission:
        submission.grade = enhanced_result.get('overallScore', 0)
        submission.feedback = enhanced_result.get('overallFeedback', '')
        submission.save()
        logger.info(f"💾 Saved grade: {submission.grade}, Performance: {enhanced_result.get('performance_level')}")
    else:
        logger.info(f"✅ Graded custom text: {enhanced_result.get('overallScore', 0)}/{parsed_rubric.get('total_points')}")
    
    return Response(enhanced_result)
    



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def check_authenticity_view(request):
    """Check submission authenticity using AI."""
    
    submission_id = request.data.get('submission_id')
    uploaded_file = request.FILES.get('file')
    is_student_assignment = request.data.get('is_student_assignment') == 'true'
    
    text_to_check = ""
    submission = None
    student_assignment = None
    student_grade = None
    subject = None
    
    # Case 1: Check existing submission
    if submission_id:
        try:
            if is_student_assignment:
                from communications.models import StudentAssignment
                student_assignment = StudentAssignment.objects.get(id=submission_id)
                # Prefer file content, fallback to description
                if student_assignment.file:
                    from .file_utils import extract_text_from_file
                    extracted_text, error = extract_text_from_file(student_assignment.file)
                    if error:
                        return Response({'error': error}, status=status.HTTP_400_BAD_REQUEST)
                    text_to_check = extracted_text
                else:
                    text_to_check = student_assignment.description
                
                # Get metadata
                student_grade = getattr(student_assignment.student, 'grade', None)
                subject = getattr(student_assignment, 'subject', None)
                
            else:
                from academics.models import Submission
                submission = Submission.objects.select_related('assignment', 'student').get(id=submission_id)
                text_to_check = submission.submitted_text
                
                # If text is empty, try to extract from file
                if not text_to_check and submission.submitted_file:
                    from .file_utils import extract_text_from_file
                    extracted_text, error = extract_text_from_file(submission.submitted_file)
                    if error:
                        return Response({'error': error}, status=status.HTTP_400_BAD_REQUEST)
                    text_to_check = extracted_text
                
                # Get metadata
                student_grade = getattr(submission.student, 'grade', None)
                subject = getattr(submission.assignment, 'subject', None)
                
        except (Submission.DoesNotExist if not is_student_assignment else StudentAssignment.DoesNotExist):
            return Response(
                {'error': 'Submission not found'},
                status=status.HTTP_404_NOT_FOUND
            )
            
    # Case 2: Check uploaded file
    elif uploaded_file:
        from .file_utils import extract_text_from_file
        extracted_text, error = extract_text_from_file(uploaded_file)
        if error:
            return Response({'error': error}, status=status.HTTP_400_BAD_REQUEST)
        text_to_check = extracted_text
        
    # Case 3: Check raw text (if provided in body)
    elif request.data.get('text'):
        text_to_check = request.data.get('text')
        
    else:
        return Response(
            {'error': 'No content provided to check'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    if not text_to_check or len(text_to_check.strip()) < 50:
        return Response(
            {'error': 'Text is too short to analyze (minimum 50 characters)'},
            status=status.HTTP_400_BAD_REQUEST
        )
        
    try:
        # Use AuthenticityEnhancer for real analysis
        from .authenticity_enhancer import AuthenticityEnhancer
        
        prompt = AuthenticityEnhancer.analyze_text_authenticity(
            text=text_to_check,
            student_grade=student_grade,
            subject=subject
        )
        
        system_prompt = "You are an expert AI Detection Analyst. You analyze text for patterns typical of AI generation vs human writing."
        
        llm_request = LLMRequest(
            prompt=prompt,
            user_id=request.user.id,
            user_role=UserRole(request.user.role),
            task_type=TaskType.GRADING, # Reuse grading task type or add new one
            complexity=TaskComplexity.ADVANCED,
            system_prompt=system_prompt,
            temperature=0.3, # Low temperature for consistent analysis
            max_tokens=1000,
        )
        
        # Call LLM
        result = llm_router.process_json_request(llm_request)
        
        # Validate result
        validated_result = AuthenticityEnhancer.validate_authenticity_result(result)
        
        # Log the check
        logger.info(f"🔍 Authenticity check: Score={validated_result['originality_score']}, Likelihood={validated_result['ai_likelihood']}")
        
        # Save authenticity scores to database if it's a submission
        if submission:
            submission.authenticity_score = validated_result.get('originality_score', 50)
            submission.ai_likelihood = validated_result.get('ai_likelihood', 'Medium')
            submission.save()
        elif student_assignment:
            student_assignment.authenticity_score = validated_result.get('originality_score', 50)
            student_assignment.ai_likelihood = validated_result.get('ai_likelihood', 'Medium')
            student_assignment.save()
        
        return Response(validated_result)
        
    except Exception as e:
        logger.error(f"Authenticity check failed: {e}", exc_info=True)
        return Response(
            {'error': 'An error occurred while checking the text. Please try again.', 'details': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def evaluate_practice_answer_view(request):
    """Evaluate a practice question answer."""
    
    question = request.data.get('question', '')
    answer = request.data.get('answer', '')
    correct_answer = request.data.get('correct_answer', '')
    
    if not answer:
        return Response(
            {'error': 'Answer is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    prompt = f"""Evaluate the student's answer to this practice question:

**Question:** {question}
**Student's Answer:** {answer}
**Correct Answer:** {correct_answer}

**IMPORTANT: Return ONLY a valid JSON object. Do not include any text before or after the JSON. Do not include explanations outside the JSON.**

Return this exact JSON structure:
{{
    "isCorrect": true/false,
    "score": 0-100,
    "feedback": "Warm, encouraging feedback that acknowledges what the student got right",
    "explanation": "Clear explanation of the concept, why the correct answer is correct, and what makes it complete",
    "hints": ["2-3 helpful hints that guide thinking without giving away the answer"],
    "nextSteps": "Specific, actionable suggestions for what to study or practice next"
}}

**Evaluation Guidelines:**
- Award partial credit for partially correct answers (50-90 points)
- Be generous with encouragement, even for wrong answers
- Focus on understanding, not just correctness
- Provide hints that promote deeper thinking
- Suggest concrete next steps for improvement
- Use simple, clear language appropriate for students
- Celebrate effort and progress

Be warm, supportive, and educational in your JSON content!"""
    
    system_prompt = """You are YENETA, an expert AI tutor specializing in formative assessment.

CRITICAL: You MUST respond with ONLY a valid JSON object. No text before or after the JSON. No markdown code blocks. No explanations outside the JSON.

Your goals:
1. Evaluate student understanding accurately and fairly
2. Provide constructive feedback that builds confidence
3. Guide students toward correct understanding through hints
4. Encourage a growth mindset and love of learning
5. Make learning feel achievable and rewarding

Put all your warm, encouraging, and educational content INSIDE the JSON fields (feedback, explanation, hints, nextSteps)."""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.TUTORING,
        complexity=TaskComplexity.BASIC,
        system_prompt=system_prompt,
        temperature=0.6,
        max_tokens=800,
    )
    
    try:
        # Use robust JSON processing
        feedback = llm_router.process_json_request(llm_request)
        return Response(feedback)
    
    except Exception as e:
        logger.error(f"Practice evaluation JSON error: {e}")
        # Fallback for evaluation failure
        return Response({
            'isCorrect': False,
            'feedback': "AI evaluation failed. Please try again.",
            'error': str(e)
        })
    



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def summarize_conversation_view(request):
    """Summarize a conversation using AI."""
    
    messages = request.data.get('messages', [])
    
    if not messages:
        return Response(
            {'error': 'Messages are required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Build conversation text
    conversation_text = "\n".join([
        f"{msg.get('sender', 'Unknown')}: {msg.get('content', '')}"
        for msg in messages
    ])
    
    prompt = f"""Summarize the following conversation:

{conversation_text}

Provide a JSON response with:
{{
    "summary": "Brief summary",
    "keyPoints": ["point 1", "point 2"],
    "actionItems": ["action 1", "action 2"],
    "sentiment": "positive/neutral/negative",
    "topics": ["topic 1", "topic 2"]
}}"""
    
    system_prompt = "You are an expert at summarizing educational conversations. Be concise and capture key information."
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.CONVERSATION_SUMMARY,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=system_prompt,
        temperature=0.5,
        max_tokens=800,
    )
    
    try:
        # Use robust JSON processing
        summary_data = llm_router.process_json_request(llm_request)
        return Response(summary_data)
    
    except Exception as e:
        logger.error(f"Conversation summary JSON error: {e}")
        # Fallback
        return Response({
            'summary': "AI summary failed to generate structured output.",
            'error': str(e)
        })
    



@api_view(['POST'])
@permission_classes([IsAuthenticated])
def analyze_alert_view(request):
    """Analyze a smart alert using AI."""
    
    alert_id = request.data.get('alert_id')
    alert_data = request.data.get('alert', {})
    alert_text = request.data.get('alert_text', '')
    
    # If alert_id is provided, fetch the alert from database
    if alert_id:
        try:
            from alerts.models import SmartAlert
            alert = SmartAlert.objects.get(id=alert_id)
            alert_data = {
                'id': alert.id,
                'type': alert.type,
                'student_name': alert.student.username if alert.student else 'Unknown',
                'context': f"Priority: {alert.priority}, Status: {alert.status}"
            }
            alert_text = alert.message
        except SmartAlert.DoesNotExist:
            return Response(
                {'error': 'Alert not found'},
                status=status.HTTP_404_NOT_FOUND
            )
    
    if not alert_text and not alert_data:
        return Response(
            {'error': 'Alert data is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    prompt = f"""Analyze the following student alert:

Alert Type: {alert_data.get('type', 'Unknown')}
Student: {alert_data.get('student_name', 'Unknown')}
Alert Text: {alert_text}
Context: {alert_data.get('context', '')}

Provide a JSON response with:
{{
    "sentiment": "positive/neutral/negative/concerning",
    "severity": "low/medium/high/critical",
    "analysis": "Brief analysis",
    "recommendedActions": ["action 1", "action 2"],
    "requiresImmediateAttention": false,
    "suggestedResponse": "Suggested response for teacher"
}}

Be sensitive and professional."""
    
    system_prompt = "You are an expert educational psychologist. Analyze student alerts with care and provide actionable recommendations."
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.ALERT_ANALYSIS,
        complexity=TaskComplexity.ADVANCED,
        system_prompt=system_prompt,
        temperature=0.4,
        max_tokens=1000,
    )
    
    try:
        # Use robust JSON processing
        analyzed_alert = llm_router.process_json_request(llm_request)
        analyzed_alert['id'] = alert_data.get('id')
        return Response(analyzed_alert)
    
    except Exception as e:
        logger.error(f"Alert analysis JSON error: {e}")
        # Fallback
        return Response({
            'id': alert_data.get('id'),
            'sentiment': 'neutral',
            'analysis': "AI analysis failed to generate structured output.",
            'error': str(e)
        })
    



@api_view(['GET'])
@permission_classes([IsAuthenticated])
def system_status_view(request):
    """Get comprehensive system status."""
    
    try:
        # Get all system statuses
        ollama_status = ollama_manager.get_status()
        rag_stats = rag_service.get_stats()
        serp_stats = serp_service.get_stats()
        cost_summary = cost_tracker.get_analytics_summary()
        
        system_status_data = {
            'timestamp': json.dumps(None),  # Will be replaced with actual timestamp
            'ollama': {
                'available': ollama_status.get('available', False),
                'models_installed': ollama_status.get('model_count', 0),
                'required_models': ollama_status.get('required_models', {}),
                'all_required_installed': ollama_status.get('all_required_installed', False),
            },
            'rag': {
                'enabled': rag_stats.get('enabled', False),
                'total_chunks': rag_stats.get('vector_store', {}).get('total_chunks', 0),
                'embedding_dimension': rag_stats.get('vector_store', {}).get('embedding_dimension', 0),
                'status': rag_stats.get('vector_store', {}).get('status', 'unknown'),
            },
            'serp': {
                'enabled': serp_stats.get('enabled', False),
                'api_key_set': serp_stats.get('api_key_set', False),
                'cache_size': serp_stats.get('cache_size', 0),
            },
            'cost': {
                'monthly_cost': cost_summary.get('monthly_cost', 0.0),
                'monthly_budget': cost_summary.get('monthly_budget', 0.0),
                'budget_remaining': cost_summary.get('budget_remaining', 0.0),
                'budget_percentage_used': cost_summary.get('budget_percentage_used', 0.0),
                'total_requests': cost_summary.get('total_requests', 0),
            },
            'health': {
                'ollama': 'healthy' if ollama_status.get('available') else 'unavailable',
                'rag': 'healthy' if rag_stats.get('enabled') and rag_stats.get('vector_store', {}).get('total_chunks', 0) > 0 else 'limited',
                'serp': 'healthy' if serp_stats.get('enabled') else 'disabled',
                'overall': 'operational'
            }
        }
        
        return Response(system_status_data)
    
    except Exception as e:
        logger.error(f"System status error: {e}")
        return Response(
            {'error': 'Failed to get system status'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def cost_analytics_view(request):
    """Get detailed cost analytics."""
    
    try:
        days = int(request.query_params.get('days', 30))
        
        # Get comprehensive report
        report = cost_analytics.get_comprehensive_report(days=days)
        
        return Response(report)
    
    except Exception as e:
        logger.error(f"Cost analytics error: {e}")
        return Response(
            {'error': 'Failed to get cost analytics'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def cost_summary_view(request):
    """Get quick cost summary."""
    
    try:
        summary = cost_tracker.get_analytics_summary()
        recommendations = cost_analytics.get_optimization_recommendations()
        
        return Response({
            'summary': summary,
            'recommendations': recommendations
        })
    
    except Exception as e:
        logger.error(f"Cost summary error: {e}")
        return Response(
            {'error': 'Failed to get cost summary'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def web_search_view(request):
    """Perform web search."""
    
    try:
        query = request.data.get('query', '')
        search_type = request.data.get('type', 'general')  # general, news, scholar
        num_results = request.data.get('num_results', 5)
        
        if not query:
            return Response(
                {'error': 'Query is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if not serp_service.is_available():
            return Response(
                {'error': 'Web search is not available. SERP API key not configured.'},
                status=status.HTTP_503_SERVICE_UNAVAILABLE
            )
        
        # Perform search
        results = serp_service.search(
            query=query,
            num_results=num_results,
            search_type=search_type
        )
        
        return Response({
            'query': query,
            'type': search_type,
            'results': results,
            'count': len(results)
        })
    
    except Exception as e:
        logger.error(f"Web search error: {e}")
        return Response(
            {'error': 'Failed to perform web search'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


# ============================================================================
# PRACTICE LABS - ADAPTIVE AI COACHING SYSTEM
# ============================================================================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_practice_question_view(request):
    """
    Generate adaptive practice questions with RAG support.
    Supports multiple modes: subject-based, random, exam-based (RAG), curriculum-based (RAG).
    Enhanced with PracticeLabRAGEnhancer for robust prompt engineering and context management.
    """
    from ai_tools.practice_lab_rag_enhancer import PracticeLabRAGEnhancer
    
    mode = request.data.get('mode', 'subject')
    subject = request.data.get('subject', '')
    topic = request.data.get('topic', '')
    chapter = request.data.get('chapter', '')  # New: Chapter/Unit/Lesson input
    use_chapter_mode = request.data.get('useChapterMode', False)  # New: Chapter mode flag
    grade_level = request.data.get('gradeLevel', 9)
    difficulty = request.data.get('difficulty', 'medium')
    use_exam_rag = request.data.get('useExamRAG', False)
    use_curriculum_rag = request.data.get('useCurriculumRAG', False)
    stream = request.data.get('stream', '')  # For grade 12 exams
    exam_year = request.data.get('examYear', '')  # For specific exam years
    coach_personality = request.data.get('coachPersonality', 'patient')  # 'patient', 'energetic', 'analyst'
    adaptive_difficulty = request.data.get('adaptiveDifficulty', False)
    practice_mode = request.data.get('practiceMode', 'standard')  # New: standard, exam, game
    region = request.data.get('region', None)
    
    # Handle Matric mode - force RAG usage for Grade 12 exams
    if mode == 'matric':
        grade_level = 12
        use_exam_rag = True
        if not subject:
            return Response(
                {'error': 'Subject is required for Matric Exam mode'},
                status=status.HTTP_400_BAD_REQUEST
            )
    
    # Analyze topic for chapter detection if not explicitly provided
    chapter_info = None
    if not chapter and topic:
        chapter_info = PracticeLabRAGEnhancer.analyze_topic_for_chapter(topic, subject, str(grade_level))
        if chapter_info:
            chapter = str(chapter_info['number'])
            logger.info(f"📖 Auto-detected chapter: {chapter}")

    # Build context based on mode and RAG settings
    rag_context = ""
    rag_status = 'disabled'
    rag_message = None
    curriculum_sources = []
    
    # 1. EXAM RAG (Matric/Model)
    if use_exam_rag and grade_level == 12:
        try:
            from rag.services import query_exam_documents
            
            if not subject:
                rag_message = "Subject is required for exam-based questions."
                rag_status = 'fallback'
            else:
                exam_type = 'Model' if mode == 'model' else 'Matric'
                
                # Build optimized query
                query_text = PracticeLabRAGEnhancer.build_practice_query(
                    topic=topic or chapter or "General",
                    subject=subject,
                    grade_level=str(grade_level),
                    difficulty=difficulty,
                    chapter_info=chapter_info,
                    exam_mode=True
                )
                
                documents = query_exam_documents(
                    exam_type=exam_type,
                    subject=subject,
                    query=query_text,
                    stream=stream if stream and stream != 'N/A' else None,
                    exam_year=exam_year if exam_year else None,
                    chapter=chapter if chapter else None,
                    top_k=3
                )
                
                if documents:
                    # Format context using Enhancer
                    rag_context, curriculum_sources = PracticeLabRAGEnhancer.format_practice_context(
                        documents=documents,
                        topic=topic,
                        chapter_info=chapter_info
                    )
                    rag_status = 'success'
                    rag_message = f"Using {len(documents)} {exam_type} exam documents"
                    
                    # Transform sources for frontend
                    curriculum_sources = [
                        {
                            'type': 'exam',
                            'exam_type': exam_type,
                            'subject': subject,
                            'exam_year': doc.get('metadata', {}).get('exam_year', 'All'),
                            'relevance': 1.0
                        }
                        for doc in documents
                    ]
                else:
                    rag_message = f"No {exam_type} exam documents found. Using AI knowledge."
                    rag_status = 'fallback'
        
        except Exception as e:
            logger.error(f"❌ Error querying exam documents: {str(e)}")
            rag_status = 'fallback'

    # 2. CURRICULUM RAG (Textbooks)
    elif use_curriculum_rag or (use_chapter_mode and chapter):
        try:
            from rag.services import query_curriculum_documents
            
            # Format grade string correctly for vector store lookup (e.g., "Grade 7")
            grade_str = f"Grade {grade_level}" if str(grade_level) != 'KG' else 'KG'
            
            # Build optimized query
            query_text = PracticeLabRAGEnhancer.build_practice_query(
                topic=topic or chapter or "General",
                subject=subject,
                grade_level=str(grade_level),
                difficulty=difficulty,
                chapter_info=chapter_info,
                exam_mode=False
            )
            
            documents = query_curriculum_documents(
                subject=subject,
                grade=grade_str,
                query=query_text,
                chapter=chapter if chapter else None,
                region=region,
                top_k=5
            )
            
            if documents:
                # Format context using Enhancer
                rag_context, sources = PracticeLabRAGEnhancer.format_practice_context(
                    documents=documents,
                    topic=topic,
                    chapter_info=chapter_info
                )
                rag_status = 'success'
                rag_message = f"Using {len(documents)} curriculum documents"
                curriculum_sources = sources
            else:
                rag_message = "No curriculum documents found. Using AI knowledge."
                rag_status = 'fallback'
                
        except Exception as e:
            logger.error(f"❌ Error querying curriculum documents: {str(e)}")
            rag_status = 'fallback'
    
    # Build comprehensive prompt using Enhancer
    prompt = PracticeLabRAGEnhancer.build_practice_prompt(
        topic=topic or chapter or "General",
        subject=subject,
        grade_level=grade_level,
        difficulty=difficulty,
        rag_context=rag_context,
        coach_personality=coach_personality,
        adaptive_difficulty=adaptive_difficulty,
        use_exam_rag=use_exam_rag,
        use_curriculum_rag=use_curriculum_rag,
        chapter_title=f"Chapter {chapter}" if chapter else None,
        practice_mode=practice_mode
    )
    
    system_prompt = """You are YENETA's Practice Labs AI, an expert question generator for Ethiopian students.
Your goal is to create high-quality, culturally relevant, and curriculum-aligned practice questions.
You MUST respond with ONLY valid JSON."""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.QUESTION_GENERATION,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=system_prompt,
        temperature=0.7,
        max_tokens=1500, # Increased for detailed explanations
        use_rag=use_exam_rag or use_curriculum_rag
    )
    
    try:
        # Use robust JSON processing
        question_data = llm_router.process_json_request(llm_request)
        
        # Validate question type consistency
        question_type = question_data.get('questionType', '')
        options = question_data.get('options', [])
        
        # Fix inconsistent question types
        if question_type == 'true_false' and options and len(options) > 2:
            question_data['questionType'] = 'multiple_choice'
        elif question_type == 'multiple_choice' and (not options or len(options) <= 2):
            if not options or len(options) == 0:
                question_data['questionType'] = 'short_answer'
                question_data['options'] = None
            elif len(options) == 2:
                question_data['questionType'] = 'true_false'
                question_data['options'] = None
        
        # Add unique ID
        import uuid
        question_data['id'] = str(uuid.uuid4())
        
        # Add RAG metadata to response
        question_data['ragStatus'] = rag_status
        question_data['ragMessage'] = rag_message
        question_data['curriculumSources'] = curriculum_sources
        
        return Response(question_data)
    
    except Exception as e:
        logger.error(f"Question generation error: {e}")
        return Response(
            {'error': 'Failed to generate question'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def evaluate_practice_answer_adaptive_view(request):
    """
    Evaluate practice answer with adaptive AI coaching.
    Includes personality modes, performance tracking, and motivational feedback.
    """
    
    question = request.data.get('question', '')
    answer = request.data.get('answer', '')
    correct_answer = request.data.get('correctAnswer', '')
    question_type = request.data.get('questionType', 'short_answer')
    difficulty = request.data.get('difficulty', 'medium')
    coach_personality = request.data.get('coachPersonality', 'patient')  # 'patient', 'energetic', 'analyst'
    student_performance = request.data.get('studentPerformance', {})
    
    if not answer:
        return Response(
            {'error': 'Answer is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Extract performance metrics
    correct_count = student_performance.get('correctCount', 0)
    total_count = student_performance.get('totalCount', 0)
    current_streak = student_performance.get('streak', 0)
    current_difficulty = student_performance.get('currentDifficulty', 'medium')
    
    # Build personality-specific system prompt
    personality_prompts = {
        'patient': """You are a PATIENT MENTOR - calm, wise, and Socratic.
- Use gentle, encouraging language
- Ask guiding questions to help students discover answers
- Emphasize learning process over results
- Celebrate small wins and progress
- Use phrases like "Let's think about this together..." or "What if we approach it this way..."
""",
        'energetic': """You are an ENERGETIC COACH - high-energy motivator and cheerleader.
- Use enthusiastic, uplifting language with emojis
- Celebrate effort and persistence loudly
- Turn mistakes into learning opportunities with excitement
- Use phrases like "YES! You're on fire!" or "Let's CRUSH this concept!"
- Keep energy high and positive
""",
        'analyst': """You are an ANALYST - data-driven, precise, and insight-focused.
- Provide detailed, structured feedback
- Reference patterns and trends in performance
- Give specific, actionable improvement steps
- Use metrics and progress indicators
- Use phrases like "Based on your pattern..." or "The data shows..."
"""
    }
    
    personality_prompt = personality_prompts.get(coach_personality, personality_prompts['patient'])
    
    # Build adaptive feedback prompt
    prompt = f"""Evaluate the student's answer with adaptive coaching.

**Question:** {question}
**Question Type:** {question_type}
**Difficulty:** {difficulty}
**Student's Answer:** {answer}
**Correct Answer:** {correct_answer}

**Student Performance Context:**
- Correct: {correct_count}/{total_count} questions
- Current Streak: {current_streak}
- Current Difficulty: {current_difficulty}

{personality_prompt}

**CRITICAL: Return ONLY a valid JSON object. No text before or after. No markdown code blocks.**

**JSON FORMATTING RULES:**
- All text must be in single-line strings (no line breaks within strings)
- Use spaces instead of newlines for readability
- For step-by-step explanations, use numbered format: "1. First step 2. Second step 3. Third step"
- Use markdown formatting within strings: **bold**, numbered lists (1. 2. 3.)
- DO NOT use LaTeX syntax (no $$, no \\begin, no \\frac, no \\)
- Use plain text for math: write "x^2" not "$x^2$", write "1/3" not "\\frac{1}{3}"
- Use Unicode symbols: × ÷ ² ³ √ ≈ ≤ ≥ instead of LaTeX

Return this exact JSON structure:
{{
    "isCorrect": true/false,
    "score": 0-100,
    "feedback": "Warm, personality-appropriate feedback acknowledging what student got right",
    "explanation": "Clear explanation with step-by-step breakdown. Use: 1. **Step One:** Description 2. **Step Two:** Description",
    "hints": ["Hint 1", "Hint 2", "Hint 3"],
    "nextSteps": "Specific suggestions for what to study or practice next",
    "motivationalMessage": "Encouraging message based on performance and streak",
    "difficultyAdjustment": "easier/same/harder",
    "xpEarned": 10-100,
    "skillsImproved": ["skill1", "skill2"]
}}

**Evaluation Guidelines:**
- Award partial credit generously for partially correct answers
- Provide step-by-step explanations
- Reference their performance streak and progress
- Adjust difficulty recommendation based on consistent performance
- Make XP rewards meaningful (harder questions = more XP)
- Be specific about skills demonstrated
- Match your tone to the {coach_personality} personality

Be educational, supportive, and adaptive!"""
    
    system_prompt = f"""You are YENETA's Adaptive AI Coach, specializing in personalized learning feedback.

{personality_prompt}

CRITICAL JSON REQUIREMENTS:
- Respond with ONLY a valid JSON object
- NO text before or after the JSON
- NO markdown code blocks (no ```)
- ALL string values must be single-line (no line breaks)
- Use markdown formatting WITHIN strings for structure (**bold**, 1. 2. 3.)
- NEVER use LaTeX syntax ($$, \\begin, \\frac, \\) - use plain text and Unicode symbols instead

Your goals:
1. Evaluate understanding accurately and fairly
2. Provide adaptive feedback that matches student level
3. Build confidence and growth mindset
4. Track and celebrate progress
5. Recommend appropriate next steps
6. Make learning feel rewarding and achievable

Put all your warm, encouraging, and educational content INSIDE the JSON fields as single-line strings."""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.PRACTICE_EVALUATION,
        complexity=TaskComplexity.BASIC,
        system_prompt=system_prompt,
        temperature=0.7,
        max_tokens=1000,
    )
    
    try:
        # Use robust JSON processing
        feedback = llm_router.process_json_request(llm_request)
        
        # Ensure all required fields exist
        feedback.setdefault('xpEarned', 10)
        feedback.setdefault('skillsImproved', [])
        feedback.setdefault('difficultyAdjustment', 'same')
        feedback.setdefault('motivationalMessage', 'Keep up the great work!')
        
        return Response(feedback)
    
    except Exception as e:
        logger.error(f"Adaptive evaluation error: {e}")
        return Response(
            {'error': 'Failed to evaluate answer'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def get_diagnostic_test_view(request):
    """
    Generate a diagnostic test to assess student's baseline in a subject.
    Returns a set of questions covering different difficulty levels.
    """
    
    subject = request.data.get('subject', '')
    grade_level = request.data.get('gradeLevel', 9)
    num_questions = request.data.get('numQuestions', 5)
    
    if not subject:
        return Response(
            {'error': 'Subject is required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    prompt = f"""Generate a diagnostic test for {subject} at Grade {grade_level} level.

Create {num_questions} questions that:
1. Cover different difficulty levels (easy, medium, hard)
2. Test fundamental concepts in {subject}
3. Help identify student's baseline understanding
4. Are quick to answer (diagnostic, not comprehensive)

**IMPORTANT: Return ONLY a valid JSON object. No text before or after the JSON.**

Return this exact JSON structure:
{{
    "testTitle": "Grade {grade_level} {subject} Diagnostic Assessment",
    "instructions": "Brief instructions for the student",
    "questions": [
        {{
            "id": "q1",
            "question": "Question text",
            "questionType": "multiple_choice" or "true_false" or "short_answer",
            "options": ["option1", "option2", "option3", "option4"],
            "correctAnswer": "The correct answer",
            "difficulty": "easy" or "medium" or "hard",
            "topic": "Specific topic",
            "skillTested": "Specific skill being assessed"
        }}
    ],
    "estimatedTime": "Time in minutes"
}}

Make it comprehensive yet concise for baseline assessment."""
    
    system_prompt = """You are YENETA's Diagnostic Assessment Generator.

CRITICAL: You MUST respond with ONLY a valid JSON object. No text before or after the JSON.

Create balanced diagnostic tests that accurately assess student baseline knowledge."""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.QUESTION_GENERATION,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=system_prompt,
        temperature=0.7,
        max_tokens=2000,
    )
    
    try:
        response = llm_router.process_request(llm_request)
        
        if not response.success:
            return Response(
                {'error': response.error_message},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        try:
            cleaned_content = clean_json_response(response.content)
            test_data = json.loads(cleaned_content)
            return Response(test_data)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {e}")
            return Response(
                {'error': 'Failed to parse diagnostic test'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    except Exception as e:
        logger.error(f"Diagnostic test generation error: {e}")
        return Response(
            {'error': 'Failed to generate diagnostic test'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def evaluate_diagnostic_test_view(request):
    """
    Evaluate a completed diagnostic test and provide detailed feedback.
    """
    
    subject = request.data.get('subject', '')
    grade_level = request.data.get('gradeLevel', 9)
    questions = request.data.get('questions', [])
    answers = request.data.get('answers', {})
    
    if not subject or not questions or not answers:
        return Response(
            {'error': 'Subject, questions, and answers are required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    # Calculate basic metrics
    total_questions = len(questions)
    correct_count = 0
    skill_performance = {}
    
    # Evaluate each answer
    for question in questions:
        q_id = question.get('id')
        user_answer = answers.get(q_id, '')
        correct_answer = question.get('correctAnswer', '')
        skill = question.get('skillTested', 'General')
        
        # Simple comparison (can be enhanced with AI evaluation)
        is_correct = user_answer.strip().lower() == correct_answer.strip().lower()
        if is_correct:
            correct_count += 1
        
        # Track skill performance
        if skill not in skill_performance:
            skill_performance[skill] = {'correct': 0, 'total': 0}
        skill_performance[skill]['total'] += 1
        if is_correct:
            skill_performance[skill]['correct'] += 1
    
    # Build evaluation prompt
    prompt = f"""Analyze this diagnostic test performance for {subject} at Grade {grade_level}.

**Test Results:**
- Total Questions: {total_questions}
- Correct Answers: {correct_count}
- Score: {(correct_count/total_questions*100):.1f}%

**Skill Performance:**
{chr(10).join([f"- {skill}: {data['correct']}/{data['total']} correct" for skill, data in skill_performance.items()])}

**Questions and Answers:**
{chr(10).join([f"Q{i+1} ({q.get('difficulty')}): {q.get('question')[:100]}... | Student: {answers.get(q.get('id'), 'No answer')} | Correct: {q.get('correctAnswer')}" for i, q in enumerate(questions)])}

**IMPORTANT: Return ONLY a valid JSON object. No text before or after the JSON.**

Return this exact JSON structure:
{{
    "overallScore": {(correct_count/total_questions*100):.1f},
    "totalQuestions": {total_questions},
    "correctAnswers": {correct_count},
    "skillAssessment": [
        {{
            "skill": "Skill name",
            "level": "beginner" or "intermediate" or "advanced",
            "score": 0-100,
            "feedback": "Specific feedback for this skill"
        }}
    ],
    "recommendations": ["Specific actionable recommendation 1", "Recommendation 2"],
    "strengthAreas": ["Area where student performed well"],
    "improvementAreas": ["Area needing improvement"],
    "suggestedTopics": ["Topic 1", "Topic 2"]
}}

Provide detailed, actionable feedback based on the performance."""
    
    system_prompt = """You are YENETA's Diagnostic Evaluator.

CRITICAL: You MUST respond with ONLY a valid JSON object. No text before or after the JSON.

Provide constructive, encouraging feedback that identifies both strengths and areas for growth."""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.EVALUATION,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=system_prompt,
        temperature=0.7,
        max_tokens=1500,
    )
    
    try:
        response = llm_router.process_request(llm_request)
        
        if not response.success:
            return Response(
                {'error': response.error_message},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        try:
            cleaned_content = clean_json_response(response.content)
            evaluation_data = json.loads(cleaned_content)
            return Response(evaluation_data)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {e}")
            return Response(
                {'error': 'Failed to parse evaluation'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    except Exception as e:
        logger.error(f"Diagnostic evaluation error: {e}")
        return Response(
            {'error': 'Failed to evaluate diagnostic test'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def get_session_reflection_view(request):
    """
    Generate a reflection summary after a practice session.
    Analyzes performance, identifies patterns, and suggests next steps.
    """
    
    session_data = request.data.get('sessionData', {})
    total_questions = session_data.get('totalQuestions', 0)
    correct_answers = session_data.get('correctAnswers', 0)
    subjects_covered = session_data.get('subjectsCovered', [])
    difficulty_breakdown = session_data.get('difficultyBreakdown', {})
    time_spent = session_data.get('timeSpent', 0)  # in minutes
    
    prompt = f"""Analyze this practice session and provide a comprehensive reflection.

**Session Statistics:**
- Total Questions: {total_questions}
- Correct Answers: {correct_answers}
- Accuracy: {(correct_answers/total_questions*100) if total_questions > 0 else 0:.1f}%
- Subjects Covered: {', '.join(subjects_covered)}
- Difficulty Breakdown: {difficulty_breakdown}
- Time Spent: {time_spent} minutes

**IMPORTANT: Return ONLY a valid JSON object. No text before or after the JSON.**

Return this exact JSON structure:
{{
    "overallPerformance": "Excellent/Good/Fair/Needs Improvement",
    "summary": "Brief summary of the session",
    "strengths": ["strength1", "strength2", "strength3"],
    "areasForImprovement": ["area1", "area2"],
    "patterns": ["pattern1", "pattern2"] - observed learning patterns,
    "recommendations": ["recommendation1", "recommendation2", "recommendation3"],
    "nextSessionFocus": "What to focus on next",
    "motivationalMessage": "Encouraging message for the student",
    "xpEarned": total XP from session,
    "achievementsUnlocked": ["achievement1"] - if any milestones reached,
    "studyTips": ["tip1", "tip2"] - specific study strategies
}}

Provide insightful, actionable feedback that helps the student improve."""
    
    system_prompt = """You are YENETA's Session Reflection Analyzer.

CRITICAL: You MUST respond with ONLY a valid JSON object. No text before or after the JSON.

Provide thoughtful analysis that:
1. Identifies patterns in performance
2. Celebrates progress and effort
3. Gives specific, actionable recommendations
4. Motivates continued learning
5. Suggests optimal next steps"""
    
    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.ANALYSIS,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=system_prompt,
        temperature=0.6,
        max_tokens=1200,
    )
    
    try:
        response = llm_router.process_request(llm_request)
        
        if not response.success:
            return Response(
                {'error': response.error_message},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        try:
            cleaned_content = clean_json_response(response.content)
            reflection = json.loads(cleaned_content)
            return Response(reflection)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {e}")
            return Response(
                {'error': 'Failed to parse session reflection'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    except Exception as e:
        logger.error(f"Session reflection error: {e}")
        return Response(
            {'error': 'Failed to generate session reflection'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


# ============================================================================
# SAVED LESSON PLANS - CRUD, SHARE, EXPORT
# ============================================================================

class SavedLessonPlanViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing saved lesson plans.
    Supports CRUD operations, sharing, rating, and export.
    """
    permission_classes = [IsAuthenticated]
    pagination_class = LessonPlanPagination
    
    def get_serializer_class(self):
        if self.action == 'list':
            return SavedLessonPlanListSerializer
        return SavedLessonPlanSerializer
    
    def get_queryset(self):
        """
        Return lesson plans created by user, shared with user, or public.
        """
        user = self.request.user
        
        # Base queryset: own plans + shared plans + public plans
        queryset = SavedLessonPlan.objects.filter(
            Q(created_by=user) |
            Q(shared_with=user) |
            Q(is_public=True)
        ).distinct()
        
        # Filter by query parameters
        grade = self.request.query_params.get('grade', None)
        subject = self.request.query_params.get('subject', None)
        search = self.request.query_params.get('search', None)
        my_plans = self.request.query_params.get('my_plans', None)
        public_only = self.request.query_params.get('public_only', None)
        
        if grade:
            queryset = queryset.filter(grade=grade)
        if subject:
            queryset = queryset.filter(subject__icontains=subject)
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) |
                Q(topic__icontains=search) |
                Q(tags__contains=[search])
            )
        if my_plans == 'true':
            queryset = queryset.filter(created_by=user)
        if public_only == 'true':
            queryset = queryset.filter(is_public=True)
        
        return queryset.order_by('-created_at')
    
    def perform_create(self, serializer):
        """Set the created_by field to current user"""
        serializer.save(created_by=self.request.user)
    
    def perform_update(self, serializer):
        """Only allow updates by the creator"""
        if serializer.instance.created_by != self.request.user:
            raise PermissionError("You can only edit your own lesson plans")
        serializer.save()
    
    def perform_destroy(self, instance):
        """Only allow deletion by the creator"""
        if instance.created_by != self.request.user:
            raise PermissionError("You can only delete your own lesson plans")
        instance.delete()
    
    @action(detail=True, methods=['post'])
    def use(self, request, pk=None):
        """Increment usage counter when a teacher uses this plan"""
        lesson_plan = self.get_object()
        lesson_plan.increment_usage()
        return Response({'times_used': lesson_plan.times_used})
    
    @action(detail=True, methods=['post'])
    def rate(self, request, pk=None):
        """Rate a lesson plan"""
        lesson_plan = self.get_object()
        rating_value = request.data.get('rating')
        comment = request.data.get('comment', '')
        
        if not rating_value or not (1 <= int(rating_value) <= 5):
            return Response(
                {'error': 'Rating must be between 1 and 5'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Create or update rating
        rating_obj, created = LessonPlanRating.objects.update_or_create(
            lesson_plan=lesson_plan,
            rated_by=request.user,
            defaults={'rating': int(rating_value), 'comment': comment}
        )
        
        # Update lesson plan average rating
        lesson_plan.add_rating(int(rating_value))
        
        serializer = LessonPlanRatingSerializer(rating_obj)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Export lesson plan as PDF"""
        lesson_plan = self.get_object()
        
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            from io import BytesIO
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1e40af'),
                spaceAfter=12,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#1e40af'),
                spaceAfter=6,
                spaceBefore=12
            )
            
            # Title
            story.append(Paragraph(lesson_plan.title, title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Administrative info table
            admin_data = [
                ['Grade:', lesson_plan.grade, 'Subject:', lesson_plan.subject],
                ['Topic:', lesson_plan.topic, 'Duration:', f"{lesson_plan.duration} minutes"],
            ]
            if lesson_plan.moe_standard_id:
                admin_data.append(['MoE Standard:', lesson_plan.moe_standard_id, '', ''])
            
            admin_table = Table(admin_data, colWidths=[1.5*inch, 2*inch, 1.5*inch, 2*inch])
            admin_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e5e7eb')),
                ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#e5e7eb')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(admin_table)
            story.append(Spacer(1, 0.3*inch))
            
            # Learning Objectives
            if lesson_plan.objectives:
                story.append(Paragraph("Learning Objectives", heading_style))
                for obj in lesson_plan.objectives:
                    story.append(Paragraph(f"• {obj}", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Essential Questions
            if lesson_plan.essential_questions:
                story.append(Paragraph("Essential Questions", heading_style))
                for eq in lesson_plan.essential_questions:
                    story.append(Paragraph(f"• {eq}", styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Materials
            if lesson_plan.materials:
                story.append(Paragraph("Materials Needed", heading_style))
                materials_text = ", ".join(lesson_plan.materials)
                story.append(Paragraph(materials_text, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # 5E Instructional Sequence
            if lesson_plan.five_e_sequence:
                story.append(Paragraph("5E Instructional Sequence", heading_style))
                for phase in lesson_plan.five_e_sequence:
                    phase_name = phase.get('phase', 'Activity')
                    duration = phase.get('duration', 0)
                    story.append(Paragraph(f"<b>{phase_name}</b> ({duration} min)", styles['Heading3']))
                    
                    if phase.get('activities'):
                        story.append(Paragraph("<i>Activities:</i>", styles['Normal']))
                        for act in phase['activities']:
                            story.append(Paragraph(f"  • {act}", styles['Normal']))
                    
                    if phase.get('teacherActions'):
                        story.append(Paragraph("<i>Teacher Actions:</i> " + ", ".join(phase['teacherActions']), styles['Normal']))
                    
                    if phase.get('studentActions'):
                        story.append(Paragraph("<i>Student Actions:</i> " + ", ".join(phase['studentActions']), styles['Normal']))
                    
                    story.append(Spacer(1, 0.1*inch))
            
            # Assessment Plan
            if lesson_plan.assessment_plan:
                story.append(Paragraph("Assessment Plan", heading_style))
                ap = lesson_plan.assessment_plan
                
                if ap.get('formativeChecks'):
                    story.append(Paragraph("<b>Formative Checks:</b>", styles['Normal']))
                    for check in ap['formativeChecks']:
                        story.append(Paragraph(f"  • {check}", styles['Normal']))
                
                if ap.get('summativeTask'):
                    story.append(Paragraph(f"<b>Summative Task:</b> {ap['summativeTask']}", styles['Normal']))
                
                if ap.get('successCriteria'):
                    story.append(Paragraph("<b>Success Criteria:</b>", styles['Normal']))
                    for criteria in ap['successCriteria']:
                        story.append(Paragraph(f"  • {criteria}", styles['Normal']))
                
                story.append(Spacer(1, 0.2*inch))
            
            # Differentiation
            if lesson_plan.differentiation_strategies:
                story.append(Paragraph("Differentiation Strategies", heading_style))
                for strategy in lesson_plan.differentiation_strategies:
                    level = strategy.get('level', 'Standard')
                    story.append(Paragraph(f"<b>{level}:</b>", styles['Normal']))
                    
                    if strategy.get('contentAdaptations'):
                        story.append(Paragraph("  Content: " + ", ".join(strategy['contentAdaptations']), styles['Normal']))
                    if strategy.get('processAdaptations'):
                        story.append(Paragraph("  Process: " + ", ".join(strategy['processAdaptations']), styles['Normal']))
                
                story.append(Spacer(1, 0.2*inch))
            
            # Homework
            if lesson_plan.homework:
                story.append(Paragraph("Homework", heading_style))
                story.append(Paragraph(lesson_plan.homework, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
            
            # Reflection Prompts
            if lesson_plan.reflection_prompts:
                story.append(Paragraph("Teacher Reflection Prompts", heading_style))
                for prompt in lesson_plan.reflection_prompts:
                    story.append(Paragraph(f"• {prompt}", styles['Normal']))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            # Return PDF response
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="lesson_plan_{lesson_plan.id}.pdf"'
            return response
            
        except ImportError:
            return Response(
                {'error': 'PDF export requires reportlab library. Please install it.'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        except Exception as e:
            logger.error(f"PDF export error: {e}", exc_info=True)
            return Response(
                {'error': f'Failed to export PDF: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'])
    def export(self, request, pk=None):
        """Generic export for lesson plans (pdf, docx, txt)"""
        lesson_plan = self.get_object()
        export_format = request.query_params.get('format', 'pdf').lower()
        
        if export_format == 'pdf':
            return self.export_pdf(request, pk)
        elif export_format == 'docx':
            # Implement DOCX export for lesson plan
            # For now, fallback to text or implement simple docx
            return self._export_plan_as_docx(lesson_plan)
        else:
            return self._export_plan_as_txt(lesson_plan)

    def _export_plan_as_txt(self, plan):
        content = f"LESSON PLAN: {plan.title}\n"
        content += "=" * 40 + "\n"
        content += f"Grade: {plan.grade}\nSubject: {plan.subject}\nTopic: {plan.topic}\n\n"
        content += f"Objectives:\n"
        for obj in plan.objectives:
            content += f"- {obj}\n"
        content += "\n"
        # Add more fields as needed
        
        response = HttpResponse(content, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="lesson_plan_{plan.id}.txt"'
        return response

    def _export_plan_as_docx(self, plan):
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            doc.add_heading(plan.title, 0)
            doc.add_paragraph(f"Grade: {plan.grade} | Subject: {plan.subject}")
            doc.add_paragraph(f"Topic: {plan.topic}")
            
            doc.add_heading('Objectives', level=1)
            for obj in plan.objectives:
                doc.add_paragraph(obj, style='List Bullet')
                
            # Add other sections...
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="lesson_plan_{plan.id}.docx"'
            return response
        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            return Response({'error': str(e)}, status=500)
    
    @action(detail=True, methods=['post'])
    def duplicate(self, request, pk=None):
        """Create a copy of a lesson plan"""
        original = self.get_object()
        
        # Create a duplicate
        duplicate = SavedLessonPlan.objects.create(
            created_by=request.user,
            title=f"{original.title} (Copy)",
            grade=original.grade,
            subject=original.subject,
            topic=original.topic,
            duration=original.duration,
            moe_standard_id=original.moe_standard_id,
            objectives=original.objectives,
            essential_questions=original.essential_questions,
            enduring_understandings=original.enduring_understandings,
            moe_competencies=original.moe_competencies,
            assessment_plan=original.assessment_plan,
            materials=original.materials,
            teacher_preparation=original.teacher_preparation,
            resource_constraints=original.resource_constraints,
            five_e_sequence=original.five_e_sequence,
            activities=original.activities,
            differentiation_strategies=original.differentiation_strategies,
            homework=original.homework,
            extensions=original.extensions,
            reflection_prompts=original.reflection_prompts,
            student_readiness=original.student_readiness,
            local_context=original.local_context,
            rag_enabled=original.rag_enabled,
            curriculum_sources=original.curriculum_sources,
            tags=original.tags,
        )
        
        serializer = SavedLessonPlanSerializer(duplicate)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['post'])
    def share(self, request, pk=None):
        """Share lesson plan with specific users"""
        lesson_plan = self.get_object()
        user_ids = request.data.get('user_ids', [])
        message = request.data.get('message', '')
        
        if not user_ids:
            return Response(
                {'error': 'At least one user ID is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        from users.models import User
        shared_files = []
        
        for user_id in user_ids:
            try:
                recipient = User.objects.get(id=user_id)
                
                # Don't share with self
                if recipient == request.user:
                    continue
                
                # Create shared file record
                shared_file = SharedFile.objects.create(
                    shared_by=request.user,
                    shared_with=recipient,
                    content_type='lesson_plan',
                    lesson_plan=lesson_plan,
                    message=message
                )
                shared_files.append(shared_file)
                
            except User.DoesNotExist:
                continue
        
        return Response({
            'message': f'Lesson plan shared with {len(shared_files)} user(s)',
            'shared_count': len(shared_files)
        }, status=status.HTTP_201_CREATED)


# ============================================================================
# CHAPTER CONTENT EXTRACTION - RAG-POWERED
# ============================================================================

def number_to_words_map():
    """Map numbers to their word equivalents (1-20 + common larger numbers)"""
    return {
        '1': ['One', 'First', 'I'],
        '2': ['Two', 'Second', 'II'],
        '3': ['Three', 'Third', 'III'],
        '4': ['Four', 'Fourth', 'IV'],
        '5': ['Five', 'Fifth', 'V'],
        '6': ['Six', 'Sixth', 'VI'],
        '7': ['Seven', 'Seventh', 'VII'],
        '8': ['Eight', 'Eighth', 'VIII'],
        '9': ['Nine', 'Ninth', 'IX'],
        '10': ['Ten', 'Tenth', 'X'],
        '11': ['Eleven', 'Eleventh', 'XI'],
        '12': ['Twelve', 'Twelfth', 'XII'],
        '13': ['Thirteen', 'Thirteenth', 'XIII'],
        '14': ['Fourteen', 'Fourteenth', 'XIV'],
        '15': ['Fifteen', 'Fifteenth', 'XV'],
        '16': ['Sixteen', 'Sixteenth', 'XVI'],
        '17': ['Seventeen', 'Seventeenth', 'XVII'],
        '18': ['Eighteen', 'Eighteenth', 'XVIII'],
        '19': ['Nineteen', 'Nineteenth', 'XIX'],
        '20': ['Twenty', 'Twentieth', 'XX'],
    }

def normalize_chapter_input(chapter_input):
    """
    Extract and normalize chapter/unit number from various input formats.
    Returns dict with number and original input.
    """
    chapter_input = chapter_input.strip()
    
    # Common patterns to extract numbers
    patterns = [
        r'(?:chapter|unit|lesson|module)\s*(\d+)',  # "Chapter 3", "Unit 5"
        r'(?:chapter|unit|lesson|module)\s+(one|two|three|four|five|six|seven|eight|nine|ten|eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|nineteen|twenty)',  # "Chapter Three"
        r'^(\d+)$',  # Just "3"
        r'(?:ch|u|l|m)\.?\s*(\d+)',  # "Ch. 3", "U 5"
    ]
    
    # Word to number mapping
    word_to_num = {
        'one': '1', 'two': '2', 'three': '3', 'four': '4', 'five': '5',
        'six': '6', 'seven': '7', 'eight': '8', 'nine': '9', 'ten': '10',
        'eleven': '11', 'twelve': '12', 'thirteen': '13', 'fourteen': '14',
        'fifteen': '15', 'sixteen': '16', 'seventeen': '17', 'eighteen': '18',
        'nineteen': '19', 'twenty': '20',
        'first': '1', 'second': '2', 'third': '3', 'fourth': '4', 'fifth': '5',
        'sixth': '6', 'seventh': '7', 'eighth': '8', 'ninth': '9', 'tenth': '10',
    }
    
    chapter_number = None
    
    for pattern in patterns:
        match = re.search(pattern, chapter_input.lower())
        if match:
            extracted = match.group(1)
            # Convert word to number if needed
            chapter_number = word_to_num.get(extracted.lower(), extracted)
            break
    
    return {
        'number': chapter_number,
        'original': chapter_input
    }

def build_chapter_query_variants(chapter_info):
    """
    Build multiple query variants for flexible matching.
    Handles Chapter/Unit synonyms and numeric/word forms.
    """
    number = chapter_info['number']
    original = chapter_info['original']
    
    if not number:
        # If we couldn't extract a number, use original as-is
        return f"- {original}"
    
    # Section type synonyms
    section_types = ['Chapter', 'Unit', 'Lesson', 'Module', 'Section']
    
    # Get word equivalents for the number
    num_map = number_to_words_map()
    word_forms = num_map.get(number, [])
    
    # Build all possible variants
    variants = []
    
    # Add numeric variants
    for section_type in section_types:
        variants.append(f"{section_type} {number}")
    
    # Add word variants
    for section_type in section_types:
        for word_form in word_forms:
            variants.append(f"{section_type} {word_form}")
    
    # Add original input
    if original not in variants:
        variants.insert(0, original)
    
    # Format as bullet list
    return '\n'.join([f"- {variant}" for variant in variants[:15]])  # Limit to top 15 variants

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def extract_chapter_content_view(request):
    """
    Extract chapter content from curriculum using RAG with full chapter awareness.
    Returns: comprehensive topics, objectives, MoE code, and other metadata.
    Handles various naming conventions: Chapter/Unit, numeric/word forms.
    """
    grade = request.data.get('grade', '')
    subject = request.data.get('subject', '')
    chapter = request.data.get('chapter', '')
    
    if not all([grade, subject, chapter]):
        return Response(
            {'error': 'Grade, subject, and chapter are required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    try:
        from .chapter_assistant_enhancer import ChapterAssistantEnhancer
        from .tutor_rag_enhancer import TutorRAGEnhancer
        from rag.services import query_curriculum_documents
        
        logger.info(f"📚 Chapter Assistant: Extracting content for {grade} {subject} - {chapter}")
        
        # Analyze chapter for normalization
        chapter_info = TutorRAGEnhancer.extract_chapter_info(chapter)
        chapter_param = str(chapter_info['number']) if chapter_info else None
        
        logger.info(f"📖 Chapter detected: {chapter_info}")
        
        # Build enhanced query
        query = ChapterAssistantEnhancer.build_enhanced_extraction_query(
            grade=grade,
            subject=subject,
            chapter=chapter,
            chapter_info=chapter_info
        )
        
        # Enable full chapter extraction if chapter detected
        extract_full = bool(chapter_param)
        
        # Query curriculum documents with full chapter extraction
        documents = query_curriculum_documents(
            grade=grade,
            subject=subject,
            query=query,
            chapter=chapter_param,
            top_k=5,
            extract_full_chapter=extract_full
        )
        
        if not documents or len(documents) == 0:
            # No curriculum content found
            return Response({
                'success': False,
                'message': f'No curriculum content found for {grade} {subject}, Chapter: {chapter}.\n\n'
                          f'Possible reasons:\n'
                          f'1. No curriculum documents have been uploaded for {grade} {subject}\n'
                          f'2. The chapter/unit name doesn\'t match any content in the curriculum\n'
                          f'3. Try a different chapter number or name\n\n'
                          f'Please contact your administrator to upload curriculum documents to the system.',
                'extracted_content': None,
                'suggestions': [
                    'Verify the chapter/unit number is correct',
                    'Try using just the number (e.g., "3" instead of "Chapter 3")',
                    'Check if curriculum documents are uploaded for this grade and subject',
                    'Contact administrator to upload curriculum files'
                ]
            })
        
        logger.info(f"✅ Retrieved {len(documents)} documents for chapter extraction")
        
        # Check if we have full chapter content
        has_full_chapter = any(doc.get('full_chapter', False) for doc in documents)
        
        if has_full_chapter:
            logger.info(f"📖 Full chapter content available - comprehensive extraction enabled")
        
        # Build context from documents
        context_parts = []
        sources = []
        
        for i, doc in enumerate(documents, 1):
            content = doc.get('content', '')
            source = doc.get('source', f'Document {i}')
            is_full_chapter = doc.get('full_chapter', False)
            
            if is_full_chapter:
                chapter_num = doc.get('chapter_number', '')
                chapter_title = doc.get('title', f'Chapter {chapter_num}')
                chunk_count = doc.get('chunk_count', 0)
                
                context_parts.append(f"=== COMPLETE CHAPTER {chapter_num}: {chapter_title} ===")
                context_parts.append(f"Source: {source}")
                context_parts.append(f"Assembled from {chunk_count} chunks")
                context_parts.append("")
                context_parts.append(content)
                context_parts.append("")
            else:
                context_parts.append(f"=== Source {i}: {source} ===")
                context_parts.append(content)
                context_parts.append("")
            
            sources.append(source)
        
        rag_context = "\n".join(context_parts)
        
        # Pre-extract chapter title and objectives directly from RAG context
        # This provides more accurate extraction than relying solely on LLM
        pre_extracted = ChapterAssistantEnhancer.pre_extract_from_rag_context(
            rag_context=rag_context,
            chapter_number=chapter_info.get('number') if chapter_info else None
        )
        
        # Check if explicit objectives were found
        has_explicit_objectives = pre_extracted.get('objectives_count', 0) > 0
        
        logger.info(f"📋 Pre-extracted: title='{pre_extracted.get('chapter_title', 'N/A')}', objectives={pre_extracted.get('objectives_count', 0)}")
        
        if not has_explicit_objectives:
            logger.info(f"⚠️ No explicit objectives section found - AI will generate objectives from chapter content")
        
        # Format extraction prompt using enhancer
        extraction_prompt = ChapterAssistantEnhancer.format_extraction_prompt(
            grade=grade,
            subject=subject,
            chapter=chapter,
            chapter_info=chapter_info,
            rag_context=rag_context,
            has_full_chapter=has_full_chapter,
            has_explicit_objectives=has_explicit_objectives
        )
        
        # Create LLM request
        from .llm.models import LLMRequest, UserRole, TaskComplexity
        
        llm_request = LLMRequest(
            prompt=extraction_prompt,
            user_id=request.user.id,
            user_role=UserRole.TEACHER,  # Assuming teacher role for lesson planning
            task_type=TaskType.LESSON_PLANNING,
            complexity=TaskComplexity.MEDIUM,
            temperature=0.3,
            max_tokens=2000,
            use_rag=False  # We already got RAG context above
        )
        
        # Use the imported llm_router instance with robust JSON processing
        # Use the imported llm_router instance with robust JSON processing
        extracted_content = llm_router.process_json_request(llm_request)
        
        logger.info(f"✅ Extracted content: {len(extracted_content.get('topics', []))} topics, {len(extracted_content.get('objectives', []))} objectives")
        
        # Merge pre-extracted content with LLM-extracted content
        # Strategy depends on whether explicit objectives were found
        
        # Always use pre-extracted title if available (100% accuracy)
        if pre_extracted.get('chapter_title'):
            logger.info(f"🔄 Using pre-extracted chapter title: '{pre_extracted['chapter_title']}'")
            extracted_content['chapter_title'] = pre_extracted['chapter_title']
        elif not extracted_content.get('chapter_title'):
            # If both pre-extraction and LLM failed, use a fallback
            logger.warning(f"⚠️ No chapter title extracted - using fallback")
            extracted_content['chapter_title'] = f"Chapter {chapter}"
        
        # Handle objectives based on source
        if has_explicit_objectives:
            # Textbook has explicit objectives section - prioritize pre-extracted
            # Note: pre_extracted uses key 'objectives' for the formatted objectives
            if pre_extracted.get('objectives') and len(pre_extracted['objectives']) > 0:
                logger.info(f"✅ Using {len(pre_extracted['objectives'])} pre-extracted objectives from textbook")
                extracted_content['objectives'] = pre_extracted['objectives']
                extracted_content['objectives_source'] = 'textbook_explicit'
            else:
                # Fallback to LLM if pre-extraction failed
                logger.warning(f"⚠️ Pre-extraction found objectives section but failed to extract - using LLM objectives")
                extracted_content['objectives_source'] = 'llm_fallback'
        else:
            # No explicit objectives section - use AI-generated objectives
            logger.info(f"🤖 Using {len(extracted_content.get('objectives', []))} AI-generated objectives (no explicit section in textbook)")
            extracted_content['objectives_source'] = 'ai_generated'
            
            # Optionally merge with any partial pre-extracted objectives
            if pre_extracted.get('objectives') and len(pre_extracted['objectives']) > 0:
                logger.info(f"📝 Found {len(pre_extracted['objectives'])} partial objectives - merging with AI-generated")
                # Merge both, removing duplicates while preserving order
                seen = set()
                merged = []
                for obj in pre_extracted['formatted_objectives'] + extracted_content.get('objectives', []):
                    obj_lower = obj.lower().strip()
                    if obj_lower not in seen:
                        seen.add(obj_lower)
                        merged.append(obj)
                extracted_content['objectives'] = merged
                extracted_content['objectives_source'] = 'hybrid'
                logger.info(f"🔄 Merged to {len(merged)} total objectives")
        
        # Validate extracted content
        is_valid, warnings = ChapterAssistantEnhancer.validate_extracted_content(extracted_content)
        
        if warnings:
            logger.warning(f"⚠️ Extraction warnings: {warnings}")
        
        # Enhance extracted content
        extracted_content = ChapterAssistantEnhancer.enhance_extracted_content(
            extracted_content,
            chapter_info
        )
        
        # Add pre-extraction info to metadata
        extracted_content['pre_extraction_used'] = {
            'title_from_pre_extract': bool(pre_extracted.get('chapter_title')),
            'objectives_from_pre_extract': bool(pre_extracted.get('formatted_objectives')),
            'pre_extract_quality': pre_extracted.get('extraction_quality', 'unknown')
        }
        
        # Add source information
        extracted_content['rag_sources'] = [
            {'document': source, 'type': 'full_chapter' if has_full_chapter else 'chunk'}
            for source in sources
        ]
        
        # Add extraction metadata
        extracted_content['extraction_info'] = {
            'has_full_chapter': has_full_chapter,
            'chapter_detected': bool(chapter_info),
            'has_explicit_objectives': has_explicit_objectives,
            'objectives_source': extracted_content.get('objectives_source', 'unknown'),
            'validation_passed': is_valid,
            'warnings': warnings if warnings else None
        }
        
        # Build success message
        success_message = 'Chapter content extracted successfully'
        if not has_explicit_objectives:
            success_message += ' (AI-generated objectives - no explicit objectives section in textbook)'
        if warnings:
            success_message += f' ({len(warnings)} warnings)'
        
        return Response({
            'success': True,
            'message': success_message,
            'extracted_content': extracted_content,
            'rag_enabled': True,
            'full_chapter_extraction': has_full_chapter
        })
    
    except Exception as e:
        logger.error(f"Chapter extraction error: {str(e)}")
        return Response({
            'success': False,
            'error': str(e),
            'message': 'Failed to extract chapter content'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class SavedRubricViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing saved rubrics.
    Supports CRUD operations, filtering, and export functionality.
    """
    serializer_class = SavedRubricSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = LessonPlanPagination
    
    def get_queryset(self):
        """
        Filter rubrics based on user role and query parameters.
        Teachers see their own rubrics and public ones.
        """
        user = self.request.user
        queryset = SavedRubric.objects.all()
        
        # Filter by ownership and visibility
        if user.role == 'Teacher':
            queryset = queryset.filter(
                Q(created_by=user) | Q(is_public=True)
            )
        elif user.role == 'Admin':
            # Admins see all rubrics
            pass
        else:
            # Students and Parents only see public rubrics
            queryset = queryset.filter(is_public=True)
        
        # Apply query parameter filters
        grade_level = self.request.query_params.get('grade_level', None)
        subject = self.request.query_params.get('subject', None)
        rubric_type = self.request.query_params.get('rubric_type', None)
        search = self.request.query_params.get('search', None)
        my_rubrics = self.request.query_params.get('my_rubrics', None)
        
        if grade_level:
            queryset = queryset.filter(grade_level__icontains=grade_level)
        
        if subject:
            queryset = queryset.filter(subject__icontains=subject)
        
        if rubric_type:
            queryset = queryset.filter(rubric_type=rubric_type)
        
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) |
                Q(topic__icontains=search) |
                Q(subject__icontains=search)
            )
        
        if my_rubrics == 'true':
            queryset = queryset.filter(created_by=user)
        
        return queryset.order_by('-created_at')
    
    def get_serializer_class(self):
        """Use list serializer for list action"""
        if self.action == 'list':
            return SavedRubricListSerializer
        return SavedRubricSerializer
    
    def perform_create(self, serializer):
        """Set the creator when saving a new rubric"""
        serializer.save(created_by=self.request.user)
    
    @action(detail=True, methods=['post'])
    def use(self, request, pk=None):
        """Increment usage counter when a teacher uses this rubric"""
        rubric = self.get_object()
        rubric.increment_usage()
        return Response({'times_used': rubric.times_used})
    
    @action(detail=True, methods=['post'])
    def duplicate(self, request, pk=None):
        """Create a copy of an existing rubric"""
        original = self.get_object()
        
        # Create a duplicate
        duplicate = SavedRubric.objects.create(
            created_by=request.user,
            title=f"{original.title} (Copy)",
            topic=original.topic,
            grade_level=original.grade_level,
            subject=original.subject,
            rubric_type=original.rubric_type,
            moe_standard_id=original.moe_standard_id,
            learning_objectives=original.learning_objectives,
            criteria=original.criteria,
            total_points=original.total_points,
            weighting_enabled=original.weighting_enabled,
            multimodal_assessment=original.multimodal_assessment,
            performance_levels=original.performance_levels,
            tone_preference=original.tone_preference,
            tags=original.tags,
            is_public=False,  # Duplicates are private by default
        )
        
        serializer = self.get_serializer(duplicate)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['post'])
    def share(self, request, pk=None):
        """Share rubric with specific users"""
        rubric = self.get_object()
        user_ids = request.data.get('user_ids', [])
        message = request.data.get('message', '')
        
        if not user_ids:
            return Response(
                {'error': 'At least one user ID is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        from users.models import User
        shared_files = []
        
        for user_id in user_ids:
            try:
                recipient = User.objects.get(id=user_id)
                
                # Don't share with self
                if recipient == request.user:
                    continue
                
                # Create shared file record
                shared_file = SharedFile.objects.create(
                    shared_by=request.user,
                    shared_with=recipient,
                    content_type='rubric',
                    rubric=rubric,
                    message=message
                )
                shared_files.append(shared_file)
                
            except User.DoesNotExist:
                continue
        
        return Response({
            'message': f'Rubric shared with {len(shared_files)} user(s)',
            'shared_count': len(shared_files)
        }, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['get'])
    def export_pdf(self, request, pk=None):
        """Export rubric as PDF"""
        rubric = self.get_object()
        
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            from io import BytesIO
            
            # Register Ethiopic-compatible font
            font_name = 'Helvetica'  # Default fallback
            try:
                # Try Ebrima first (Standard on Windows 10/11 for Ethiopic/African scripts)
                ebrima_path = r'C:\Windows\Fonts\ebrima.ttf'
                nyala_path = r'C:\Windows\Fonts\nyala.ttf'
                
                if os.path.exists(ebrima_path):
                    pdfmetrics.registerFont(TTFont('Ebrima', ebrima_path))
                    font_name = 'Ebrima'
                elif os.path.exists(nyala_path):
                    pdfmetrics.registerFont(TTFont('Nyala', nyala_path))
                    font_name = 'Nyala'
                else:
                    logger.warning("No Ethiopic font found (Ebrima/Nyala). Amharic text may not render correctly.")
            except Exception as e:
                logger.warning(f"Failed to register Ethiopic font: {e}")
            
            # Create PDF buffer
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                fontName=font_name,
                textColor=colors.HexColor('#1e40af'),
                spaceAfter=12,
            )
            elements.append(Paragraph(rubric.title, title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Metadata
            metadata_style = ParagraphStyle(
                'Metadata',
                parent=styles['Normal'],
                fontName=font_name,
            )
            elements.append(Paragraph(f"<b>Topic:</b> {rubric.topic}", metadata_style))
            elements.append(Paragraph(f"<b>Grade Level:</b> {rubric.grade_level}", metadata_style))
            if rubric.subject:
                elements.append(Paragraph(f"<b>Subject:</b> {rubric.subject}", metadata_style))
            elements.append(Paragraph(f"<b>Rubric Type:</b> {rubric.get_rubric_type_display()}", metadata_style))
            elements.append(Paragraph(f"<b>Total Points:</b> {rubric.total_points}", metadata_style))
            elements.append(Spacer(1, 0.3*inch))
            
            # Learning Objectives
            heading2_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontName=font_name,
            )
            
            if rubric.learning_objectives:
                elements.append(Paragraph("<b>Learning Objectives:</b>", heading2_style))
                for i, obj in enumerate(rubric.learning_objectives, 1):
                    elements.append(Paragraph(f"{i}. {obj}", metadata_style))
                elements.append(Spacer(1, 0.2*inch))
            
            # Rubric Table
            elements.append(Paragraph("<b>Assessment Criteria:</b>", heading2_style))
            elements.append(Spacer(1, 0.1*inch))
            
            # Build table data
            criteria = rubric.criteria
            if criteria and len(criteria) > 0:
                # Get performance levels from first criterion
                perf_levels = criteria[0].get('performanceLevels', [])
                
                # Calculate column widths dynamically
                num_perf_levels = len(perf_levels)
                has_weight = rubric.weighting_enabled
                
                # Available width (A4 width minus margins)
                available_width = 7.5 * inch  # A4 width minus margins
                
                # Allocate widths
                criterion_width = 1.5 * inch  # Fixed width for criterion column
                weight_width = 0.5 * inch if has_weight else 0
                remaining_width = available_width - criterion_width - weight_width
                perf_level_width = remaining_width / num_perf_levels if num_perf_levels > 0 else 2 * inch
                
                # Build column widths list
                col_widths = [criterion_width]
                if has_weight:
                    col_widths.append(weight_width)
                col_widths.extend([perf_level_width] * num_perf_levels)
                
                # Header row
                header = ['Criterion']
                if has_weight:
                    header.append('Weight')
                
                # Use font-aware styles for header
                header_style = ParagraphStyle('HeaderOne', parent=styles['Normal'], fontName=font_name, fontSize=10, textColor=colors.whitesmoke)
                
                header_row = [Paragraph(f"<b>{h}</b>", header_style) for h in header]
                for level in perf_levels:
                    header_row.append(Paragraph(f"<b>{level.get('level', '')}</b>", header_style))

                table_data = [header_row]
                
                # Data rows - wrap text in Paragraphs for better formatting
                cell_style = ParagraphStyle('Cell', parent=styles['Normal'], fontName=font_name, fontSize=9)
                
                for criterion in criteria:
                    row = [Paragraph(criterion.get('criterion', ''), cell_style)]
                    if has_weight:
                        row.append(Paragraph(f"{criterion.get('weight', 0)}%", cell_style))
                    
                    for level in criterion.get('performanceLevels', []):
                        desc = level.get('description', '')
                        points = level.get('points', '')
                        cell_text = f"{desc}<br/><b>({points} pts)</b>"
                        row.append(Paragraph(cell_text, cell_style))
                    
                    table_data.append(row)
                
                # Create table with explicit column widths
                table = Table(table_data, colWidths=col_widths, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, -1), font_name), # Apply font to table cells
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 7),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    ('WORDWRAP', (0, 0), (-1, -1), True),
                ]))
                
                elements.append(table)
            
            # Build PDF
            doc.build(elements)
            buffer.seek(0)
            
            # Return PDF response
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{rubric.title}.pdf"'
            return response
            
        except ImportError:
            return Response(
                {'error': 'PDF export requires reportlab library'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        except Exception as e:
            logger.error(f"PDF export error: {e}")
            return Response(
                {'error': 'Failed to export PDF'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'])
    def export(self, request, pk=None):
        """Export rubric in multiple formats (PDF, DOCX, TXT)"""
        rubric = self.get_object()
        export_format = request.query_params.get('format', 'pdf').lower()
        
        if export_format not in ['pdf', 'docx', 'txt']:
            return Response(
                {'error': 'Invalid format. Must be pdf, docx, or txt'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Prepare rubric data
        rubric_data = {
            'title': rubric.title,
            'topic': rubric.topic,
            'grade_level': rubric.grade_level,
            'subject': rubric.subject,
            'rubric_type': rubric.rubric_type,
            'total_points': rubric.total_points,
            'learning_objectives': rubric.learning_objectives,
            'criteria': rubric.criteria,
            'weighting_enabled': rubric.weighting_enabled,
            'multimodal_assessment': rubric.multimodal_assessment,
            'alignment_validated': rubric.alignment_validated,
            'alignment_score': rubric.alignment_score,
        }
        
        if export_format == 'pdf':
            return _export_rubric_as_pdf(rubric_data)
        elif export_format == 'docx':
            return _export_rubric_as_docx(rubric_data)
        else:  # txt
            # Generate text content
            text_content = f"{rubric.title}\n"
            text_content += "=" * len(rubric.title) + "\n\n"
            text_content += f"Topic: {rubric.topic}\n"
            text_content += f"Grade Level: {rubric.grade_level}\n"
            if rubric.subject:
                text_content += f"Subject: {rubric.subject}\n"
            text_content += f"Rubric Type: {rubric.get_rubric_type_display()}\n"
            text_content += f"Total Points: {rubric.total_points}\n\n"
            
            if rubric.learning_objectives:
                text_content += "Learning Objectives:\n"
                for i, obj in enumerate(rubric.learning_objectives, 1):
                    text_content += f"{i}. {obj}\n"
                text_content += "\n"
            
            text_content += "Assessment Criteria:\n"
            text_content += "-" * 80 + "\n\n"
            
            for criterion in rubric.criteria:
                text_content += f"Criterion: {criterion.get('criterion', '')}\n"
                if rubric.weighting_enabled:
                    text_content += f"Weight: {criterion.get('weight', 0)}%\n"
                text_content += "\nPerformance Levels:\n"
                for level in criterion.get('performanceLevels', []):
                    text_content += f"  - {level.get('level', '')}: {level.get('description', '')} ({level.get('points', '')} pts)\n"
                text_content += "\n" + "-" * 80 + "\n\n"
            
            response = HttpResponse(text_content, content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{rubric.title}.txt"'
            return response



def _export_lesson_as_pdf(lesson):
    """Helper function to export lesson as PDF"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        from io import BytesIO
        
        # Register Ethiopic-compatible font
        font_name = 'Helvetica'  # Default fallback
        try:
            # Try Ebrima first (Standard on Windows 10/11 for Ethiopic/African scripts)
            ebrima_path = r'C:\Windows\Fonts\ebrima.ttf'
            nyala_path = r'C:\Windows\Fonts\nyala.ttf'
            
            if os.path.exists(ebrima_path):
                pdfmetrics.registerFont(TTFont('Ebrima', ebrima_path))
                font_name = 'Ebrima'
            elif os.path.exists(nyala_path):
                pdfmetrics.registerFont(TTFont('Nyala', nyala_path))
                font_name = 'Nyala'
            else:
                logger.warning("⚠️ No Ethiopic font (Ebrima/Nyala) found. Amharic export may be unreadable.")
                
        except Exception as font_err:
            logger.warning(f"⚠️ Failed to register Ethiopic font: {font_err}")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Update standard styles to use the supported font
        styles['Normal'].fontName = font_name
        styles['Heading1'].fontName = font_name
        styles['Heading2'].fontName = font_name
        styles['Heading3'].fontName = font_name
        styles['BodyText'].fontName = font_name
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=font_name,  # Ensure title supports font
            fontSize=18,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=12,
        )
        elements.append(Paragraph(lesson.title, title_style))
        
        # Metadata
        meta_style = styles['Normal']
        elements.append(Paragraph(f"<b>Subject:</b> {lesson.subject} | <b>Grade:</b> {lesson.grade}", meta_style))
        elements.append(Paragraph(f"<b>Topic:</b> {lesson.topic}", meta_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Content - Process line by line
        lines = lesson.content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                elements.append(Spacer(1, 0.1*inch))
                continue
                
            if line.startswith('# '):
                elements.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '):
                elements.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('### '):
                elements.append(Paragraph(line[4:], styles['Heading3']))
            elif line.startswith('- ') or line.startswith('* '):
                elements.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            else:
                elements.append(Paragraph(line, styles['Normal']))
        
        doc.build(elements)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
        filename = lesson.title.replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response
        
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return Response(
            {'error': f'Failed to export PDF: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

def _export_lesson_as_docx(lesson):
    """Helper function to export lesson as DOCX"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        
        doc = Document()
        
        # Set default font to Ebrima for Ethiopic support
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Ebrima'
        
        doc.add_heading(lesson.title, 0)
        
        doc.add_paragraph(f"Subject: {lesson.subject} | Grade: {lesson.grade}")
        doc.add_paragraph(f"Topic: {lesson.topic}")
        doc.add_paragraph()
        
        # Add content
        lines = lesson.content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                h = doc.add_heading(line[2:], 1)
                h.style.font.name = 'Ebrima'
            elif line.startswith('## '):
                h = doc.add_heading(line[3:], 2)
                h.style.font.name = 'Ebrima'
            elif line.startswith('### '):
                h = doc.add_heading(line[4:], 3)
                h.style.font.name = 'Ebrima'
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
                p.style.font.name = 'Ebrima'
            else:
                p = doc.add_paragraph(line)
                p.style.font.name = 'Ebrima'
                
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = lesson.title.replace(' ', '_')
        response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
        return response
        
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return Response(
            {'error': f'Failed to export DOCX: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

def _export_lesson_as_txt(lesson):
    """Helper function to export lesson as TXT"""
    text_content = f"{lesson.title}\n"
    text_content += "=" * len(lesson.title) + "\n\n"
    text_content += f"Subject: {lesson.subject}\n"
    text_content += f"Grade: {lesson.grade}\n"
    text_content += f"Topic: {lesson.topic}\n\n"
    text_content += lesson.content
    
    response = HttpResponse(text_content, content_type='text/plain')
    filename = lesson.title.replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
    return response



@api_view(['GET'])
@permission_classes([AllowAny])
def debug_export_view(request, pk=None):
    print(f"DEBUG: debug_export_view called for pk={pk}")
    return HttpResponse(f"Debug export view reached for pk={pk}")

class SavedLessonViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing saved generated lessons.
    """
    serializer_class = SavedLessonSerializer
    permission_classes = [IsAuthenticated]
    pagination_class = LessonPlanPagination
    
    def dispatch(self, request, *args, **kwargs):
        print(f"DEBUG: SavedLessonViewSet.dispatch called. Method: {request.method}, URL: {request.path}")
        return super().dispatch(request, *args, **kwargs)
    
    def get_queryset(self):
        user = self.request.user
        queryset = SavedLesson.objects.all()
        
        print(f"DEBUG: SavedLessonViewSet.get_queryset called for user: {user.username} (ID: {user.id})")
        
        if user.role == 'Admin':
            print("DEBUG: User is Admin, returning all")
            pass
        else:
            # Allow users to see their own lessons + public lessons
            queryset = queryset.filter(
                Q(created_by=user) | Q(is_public=True)
            )
            print(f"DEBUG: Filtered queryset for user {user.username}. Count: {queryset.count()}")
            
        logger.info(f"SavedLessonViewSet.get_queryset: User={user.username}, Role={user.role}, Count={queryset.count()}")
        return queryset.order_by('-created_at')
    
    def perform_create(self, serializer):
        logger.info(f"SavedLessonViewSet.perform_create: Creating lesson for {self.request.user.username}")
        serializer.save(created_by=self.request.user)
        
    @action(detail=True, methods=['get'])
    def export(self, request, pk=None):
        print(f"DEBUG: Export action called for pk={pk}")
        try:
            lesson = self.get_object()
            print(f"DEBUG: Lesson found: {lesson.title}")
        except Exception as e:
            print(f"DEBUG: get_object failed: {e}")
            raise e
            
        export_format = request.query_params.get('format', 'pdf').lower()
        
        if export_format == 'pdf':
            return _export_lesson_as_pdf(lesson)
        elif export_format == 'docx':
            return _export_lesson_as_docx(lesson)
        else:
            return _export_lesson_as_txt(lesson)

    @action(detail=True, methods=['post'])
    def share(self, request, pk=None):
        """Share a lesson with other users"""
        lesson = self.get_object()
        user_ids = request.data.get('user_ids', [])
        message = request.data.get('message', '')
        
        if not user_ids:
            return Response(
                {'error': 'No users selected for sharing'},
                status=status.HTTP_400_BAD_REQUEST
            )
            
        shared_count = 0
        for user_id in user_ids:
            try:
                user = User.objects.get(id=user_id)
                
                # Check if already shared
                if not SharedFile.objects.filter(
                    shared_by=request.user,
                    shared_with=user,
                    lesson=lesson
                ).exists():
                    shared_file = SharedFile.objects.create(
                        shared_by=request.user,
                        shared_with=user,
                        content_type='lesson',
                        lesson=lesson,
                        message=message
                    )
                    shared_count += 1
                    
                    # Create notification
                    SharedFileNotification.objects.create(
                        shared_file=shared_file,
                        recipient=user
                    )
            except User.DoesNotExist:
                continue
                
        return Response({
            'message': f'Lesson shared with {shared_count} users',
            'shared_count': shared_count
        })


# ============================================================================
# SHARED FILES VIEWSET
# ============================================================================

class SharedFileViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing shared files between users.
    Supports creating shares, viewing received files, and marking as viewed.
    """
    serializer_class = SharedFileSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        """Filter shared files based on user and query parameters"""
        user = self.request.user
        queryset = SharedFile.objects.all()
        
        # Filter by sent or received
        filter_type = self.request.query_params.get('filter', 'received')
        if filter_type == 'sent':
            queryset = queryset.filter(shared_by=user)
        else:  # received
            queryset = queryset.filter(shared_with=user)
        
        # Filter by content type
        content_type = self.request.query_params.get('content_type')
        if content_type:
            queryset = queryset.filter(content_type=content_type)
        
        # Filter by viewed status
        is_viewed = self.request.query_params.get('is_viewed')
        if is_viewed is not None:
            queryset = queryset.filter(is_viewed=is_viewed.lower() == 'true')
        
        return queryset.select_related('shared_by', 'shared_with', 'lesson_plan', 'rubric')
    
    def perform_create(self, serializer):
        """Set the sender when creating a share and create notification"""
        shared_file = serializer.save(shared_by=self.request.user)
        
        # Create notification for recipient
        from communications.models import SharedFileNotification
        SharedFileNotification.objects.create(
            shared_file=shared_file,
            recipient=shared_file.shared_with
        )
    
    @action(detail=True, methods=['post'])
    def mark_viewed(self, request, pk=None):
        """Mark a shared file as viewed"""
        shared_file = self.get_object()
        
        # Only the recipient can mark as viewed
        if shared_file.shared_with != request.user:
            return Response(
                {'error': 'You can only mark files shared with you as viewed'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        shared_file.mark_as_viewed()
        serializer = self.get_serializer(shared_file)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def unread_count(self, request):
        """Get count of unread shared files"""
        count = SharedFile.objects.filter(
            shared_with=request.user,
            is_viewed=False
        ).count()
        return Response({'count': count})


# ============================================================================
# PRACTICE LABS ENHANCED ENDPOINTS (Research Document Implementation)
# ============================================================================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_two_layer_hints_view(request):
    """
    Generate two-layer hints for a question.
    Layer 1: Minimal one-line hint
    Layer 2: Detailed step-by-step breakdown
    """
    from .practice_labs_service import TwoLayerHintGenerator
    
    question = request.data.get('question', '')
    correct_answer = request.data.get('correctAnswer', '')
    difficulty = request.data.get('difficulty', 'medium')
    
    if not question or not correct_answer:
        return Response(
            {'error': 'Question and correct answer are required'},
            status=status.HTTP_400_BAD_REQUEST
        )
    
    try:
        prompt = TwoLayerHintGenerator.generate_hints_prompt(
            question, correct_answer, difficulty
        )
        
        llm_request = LLMRequest(
            prompt=prompt,
            user_id=request.user.id,
            user_role=UserRole(request.user.role),
            task_type=TaskType.QUESTION_GENERATION,
            complexity=TaskComplexity.BASIC,
            temperature=0.7,
            max_tokens=300,
        )
        
        response = llm_router.process_request(llm_request)
        
        if not response.success:
            return Response(
                {'error': response.error_message},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
        
        cleaned_content = clean_json_response(response.content)
        hints_data = json.loads(cleaned_content)
        
        return Response(hints_data)
    
    except Exception as e:
        logger.error(f"Two-layer hints generation error: {e}")
        return Response(
            {'error': 'Failed to generate hints'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def calculate_zpd_score_view(request):
    """
    Calculate ZPD (Zone of Proximal Development) complexity score.
    Based on student's mastery profile.
    """
    from .practice_labs_service import ZPDCalculator
    
    mastery_percentage = request.data.get('masteryPercentage', 50)
    
    try:
        zpd_score = ZPDCalculator.calculate_zpd_score(mastery_percentage)
        difficulty = ZPDCalculator.zpd_to_difficulty(zpd_score)
        
        return Response({
            'zpdScore': zpd_score,
            'recommendedDifficulty': difficulty,
            'masteryPercentage': mastery_percentage
        })
    
    except Exception as e:
        logger.error(f"ZPD calculation error: {e}")
        return Response(
            {'error': 'Failed to calculate ZPD score'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def detect_misconceptions_view(request):
    """
    Detect misconception patterns from student answer.
    Returns misconceptions and remediation strategies.
    """
    from .practice_labs_service import MisconceptionDetector
    
    subject = request.data.get('subject', '')
    question = request.data.get('question', '')
    student_answer = request.data.get('studentAnswer', '')
    correct_answer = request.data.get('correctAnswer', '')
    is_correct = request.data.get('isCorrect', False)
    
    try:
        misconceptions = MisconceptionDetector.detect_misconceptions(
            subject, question, student_answer, correct_answer, is_correct
        )
        
        misconception_details = []
        for misconception_type in misconceptions:
            remediation = MisconceptionDetector.get_remediation(misconception_type)
            misconception_details.append({
                'type': misconception_type,
                'remediation': remediation
            })
        
        return Response({
            'misconceptions': misconception_details,
            'count': len(misconception_details)
        })
    
    except Exception as e:
        logger.error(f"Misconception detection error: {e}")
        return Response(
            {'error': 'Failed to detect misconceptions'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_badges_view(request):
    """
    Get available badges and student's progress.
    """
    from .practice_labs_service import BadgeSystem
    
    try:
        # In production, fetch from database
        # For now, return badge definitions
        badges = list(BadgeSystem.BADGES.values())
        
        return Response({
            'badges': badges,
            'totalBadges': len(badges)
        })
    
    except Exception as e:
        logger.error(f"Get badges error: {e}")
        return Response(
            {'error': 'Failed to get badges'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_missions_view(request):
    """
    Get daily and weekly missions for student.
    """
    from .practice_labs_service import MissionSystem
    
    try:
        daily_missions = MissionSystem.generate_daily_missions()
        weekly_missions = MissionSystem.generate_weekly_missions()
        
        return Response({
            'dailyMissions': daily_missions,
            'weeklyMissions': weekly_missions
        })
    
    except Exception as e:
        logger.error(f"Get missions error: {e}")
        return Response(
            {'error': 'Failed to get missions'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def extract_file_text_view(request):
    """
    Extract text content from uploaded file (PDF, Word, or text).
    Used for Quick Grader file uploads.
    """
    try:
        from .file_utils import extract_text_from_file
        
        # Get uploaded file
        uploaded_file = request.FILES.get('file')
        if not uploaded_file:
            return Response(
                {'error': 'No file provided'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Extract text from file
        text_content, error = extract_text_from_file(uploaded_file)
        
        if error:
            return Response(
                {'error': error},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if not text_content:
            return Response(
                {'error': 'No text content could be extracted from the file'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        return Response({
            'text': text_content,
            'filename': uploaded_file.name,
            'size': uploaded_file.size
        })
    
    except Exception as e:
        logger.error(f"File text extraction error: {e}")
        return Response(
            {'error': f'Failed to extract text from file: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_tutor_welcome_message_view(request):
    """
    Generate a personalized welcome message for the AI Tutor using LLM.
    Creates a meaningful, well-composed message based on configuration.
    """
    try:
        grade = request.data.get('grade', '')
        subject = request.data.get('subject', '')
        stream = request.data.get('stream', '')
        chapter = request.data.get('chapter', '')
        chapter_title = request.data.get('chapter_title', '')
        chapter_topics = request.data.get('chapter_topics', [])
        chapter_summary = request.data.get('chapter_summary', '')
        learning_objectives = request.data.get('learning_objectives', [])
        use_ethiopian_curriculum = request.data.get('use_ethiopian_curriculum', True)
        
        # If chapter is provided but metadata is missing, try to fetch it from RAG
        if chapter and use_ethiopian_curriculum and (not chapter_topics or not chapter_title):
            try:
                from rag.services import query_curriculum_documents
                from .models import TutorConfiguration
                
                logger.info(f"🔍 Fetching chapter metadata for welcome message: {chapter}")
                
                # Format grade string correctly (e.g., "7" -> "Grade 7")
                grade_str = grade
                if grade and grade.isdigit():
                    grade_str = f"Grade {grade}"
                
                # Query documents
                documents = query_curriculum_documents(
                    grade=grade_str,
                    subject=subject,
                    query=f"{chapter} main topics learning objectives summary",
                    chapter=chapter,
                    top_k=3,
                    extract_full_chapter=True
                )
                
                # Fallback: If no documents found with strict chapter filter, try loose search
                if not documents and chapter:
                    logger.info(f"⚠️ No documents found with strict chapter filter '{chapter}'. Trying loose search...")
                    # Try searching for "Unit X" or "Chapter X" in content
                    loose_query = f"{chapter} main topics"
                    if chapter.isdigit():
                        loose_query = f"Unit {chapter} Chapter {chapter} {loose_query}"
                    
                    documents = query_curriculum_documents(
                        grade=grade_str,
                        subject=subject,
                        query=loose_query,
                        chapter=None, # Disable strict chapter filter
                        top_k=5
                    )
                    
                    # Filter results to ensure they are relevant to the requested chapter
                    if documents:
                        filtered_docs = []
                        chapter_num = ''.join(filter(str.isdigit, str(chapter)))
                        for doc in documents:
                            content = doc.get('content', '').lower()
                            # Check if content mentions the chapter/unit number near the top
                            if chapter_num and (f"unit {chapter_num}" in content[:500] or f"chapter {chapter_num}" in content[:500]):
                                filtered_docs.append(doc)
                            # Or if the title matches
                            elif doc.get('title') and chapter_num and str(chapter_num) in doc.get('title'):
                                filtered_docs.append(doc)
                        
                        if filtered_docs:
                            documents = filtered_docs
                            logger.info(f"✅ Found {len(documents)} documents via loose search")
                
                # Validate the documents - if they don't look like chapter starts, don't use them for specific topics
                # This prevents using random exercises (like "Hospital" for Chapter 6) as the main topic
                valid_documents = []
                if documents:
                    chapter_num = ''.join(filter(str.isdigit, str(chapter)))
                    for doc in documents:
                        content = doc.get('content', '')
                        # Check for title-like patterns or explicit chapter mentions
                        is_likely_title = False
                        
                        # 1. Explicit mention of Unit/Chapter X in first 500 chars
                        if chapter_num and (f"Unit {chapter_num}" in content[:500] or f"UNIT {chapter_num}" in content[:500] or 
                                          f"Chapter {chapter_num}" in content[:500] or f"CHAPTER {chapter_num}" in content[:500]):
                            is_likely_title = True
                        
                        # 2. Check for "UNIT SIX" or "CHAPTER SIX" style (up to 12)
                        word_map = {'1': 'ONE', '2': 'TWO', '3': 'THREE', '4': 'FOUR', '5': 'FIVE', 
                                   '6': 'SIX', '7': 'SEVEN', '8': 'EIGHT', '9': 'NINE', '10': 'TEN',
                                   '11': 'ELEVEN', '12': 'TWELVE'}
                        if not is_likely_title and chapter_num in word_map:
                            word = word_map[chapter_num]
                            if f"UNIT {word}" in content[:500] or f"CHAPTER {word}" in content[:500]:
                                is_likely_title = True
                                
                        # 3. If it has a title that matches
                        if not is_likely_title and doc.get('title'):
                            title = doc.get('title', '').lower()
                            if chapter_num and (f"unit {chapter_num}" in title or f"chapter {chapter_num}" in title):
                                is_likely_title = True
                        
                        if is_likely_title:
                            valid_documents.append(doc)
                        else:
                            logger.debug(f"⚠️ Document content does not look like Chapter {chapter} start: {content[:50]}...")
                    
                    if valid_documents:
                        documents = valid_documents
                        logger.info(f"✅ Kept {len(documents)} valid chapter documents after validation")
                    else:
                        logger.warning(f"⚠️ No documents passed validation for Chapter {chapter}. Ignoring retrieved content to avoid hallucinations.")
                        documents = [] # Clear documents if none are valid
                
                if documents:
                    # Extract metadata from documents
                    extracted_topics = []
                    extracted_objectives = []
                    extracted_title = None
                    
                    for doc in documents:
                        if doc.get('full_chapter'):
                            if not extracted_title and doc.get('title'):
                                extracted_title = doc.get('title')
                            if doc.get('topics'):
                                extracted_topics.extend(doc.get('topics', []))
                            if doc.get('learning_objectives'):
                                extracted_objectives.extend(doc.get('learning_objectives', []))
                    
                    # Update local variables if we found data
                    if extracted_title and not chapter_title:
                        chapter_title = extracted_title
                    
                    if extracted_topics and not chapter_topics:
                        # Deduplicate topics
                        seen = set()
                        chapter_topics = []
                        for t in extracted_topics:
                            if t.lower() not in seen:
                                seen.add(t.lower())
                                chapter_topics.append(t)
                        chapter_topics = chapter_topics[:8]  # Limit to top 8
                        
                    if extracted_objectives and not learning_objectives:
                        # Deduplicate objectives
                        seen = set()
                        learning_objectives = []
                        for o in extracted_objectives:
                            if o.lower() not in seen:
                                seen.add(o.lower())
                                learning_objectives.append(o)
                        learning_objectives = learning_objectives[:5]  # Limit to top 5
                        
                    logger.info(f"✅ Fetched metadata: Title='{chapter_title}', Topics={len(chapter_topics)}, Objectives={len(learning_objectives)}")
                    
                    # Update saved configuration if user is authenticated
                    if request.user.is_authenticated:
                        try:
                            config = TutorConfiguration.objects.filter(user=request.user).first()
                            if config:
                                update_fields = []
                                if chapter_title and not config.chapter_title:
                                    config.chapter_title = chapter_title
                                    update_fields.append('chapter_title')
                                if chapter_topics and not config.chapter_topics:
                                    config.chapter_topics = chapter_topics
                                    update_fields.append('chapter_topics')
                                if learning_objectives and not config.learning_objectives:
                                    config.learning_objectives = learning_objectives
                                    update_fields.append('learning_objectives')
                                
                                if update_fields:
                                    config.save(update_fields=update_fields)
                                    logger.info(f"💾 Updated TutorConfiguration with fetched metadata")
                        except Exception as save_error:
                            logger.warning(f"⚠️ Could not update TutorConfiguration: {save_error}")
            except Exception as e:
                logger.warning(f"⚠️ Failed to fetch chapter metadata: {e}")
        
        # Build context for LLM
        context_parts = []
        if grade:
            context_parts.append(f"Grade: {grade}")
        if stream:
            context_parts.append(f"Stream: {stream}")
        if subject:
            context_parts.append(f"Subject: {subject}")
        if chapter:
            context_parts.append(f"Chapter: {chapter}")
        if chapter_title:
            context_parts.append(f"Chapter Title: {chapter_title}")
        
        # Clean and filter topics - remove messy text
        clean_topics = []
        if chapter_topics and isinstance(chapter_topics, list):
            for topic in chapter_topics:
                if isinstance(topic, str):
                    # Filter out topics that are too long or contain too many special characters
                    topic_clean = topic.strip()
                    # Skip if too long (likely not a topic but content)
                    if len(topic_clean) > 100:
                        continue
                    # Skip if contains too many punctuation marks (likely not a clean topic)
                    if topic_clean.count('.') > 3 or topic_clean.count(',') > 5:
                        continue
                    # Skip if looks like a sentence fragment or activity description
                    if any(phrase in topic_clean.lower() for phrase in [
                        'sit in a group', 'discuss the questions', 'recently, the rate',
                        'they baked', 'drank wine', 'had a wonderful voice'
                    ]):
                        continue
                    if topic_clean and len(topic_clean) >= 5:
                        clean_topics.append(topic_clean)
        
        # Extract learning objectives if available
        clean_objectives = []
        if learning_objectives and isinstance(learning_objectives, list):
            for obj in learning_objectives:
                if isinstance(obj, str):
                    obj_clean = obj.strip()
                    # Filter out objectives that are too long
                    if 15 <= len(obj_clean) <= 200:
                        clean_objectives.append(obj_clean)
        
        # Build prompt for LLM with enhanced focus on topics and objectives
        system_prompt = """You are YENETA, an expert AI Tutor for Ethiopian students following research-based pedagogical principles.
Your role is to create warm, welcoming messages that:
- Make students feel supported and excited to learn
- Clearly communicate what they'll be learning (topics and objectives)
- Align with Understanding by Design (UbD) principles - focus on learning outcomes
- Follow 5E Instructional Model - engage students in active learning
- Apply Cognitive Load Theory - present information clearly and progressively
- Use Socratic Method - encourage discovery through questions
- Emphasize Growth Mindset - celebrate effort and progress
- Maintain Cultural Relevance - use Ethiopian context when appropriate

Keep messages concise (3-5 sentences), friendly, inspiring, and pedagogically sound."""
        
        user_prompt = f"""Create a personalized welcome message for a student starting a tutoring session.

Context:
{chr(10).join(context_parts) if context_parts else 'General tutoring session'}

"""
        
        if chapter:
            user_prompt += f"""The student is focusing on {chapter}"""
            user_prompt += f"""The student is focusing on {chapter}"""
            if chapter_title:
                user_prompt += f": {chapter_title}"
            user_prompt += ".\n\n"
            
            if chapter_title:
                user_prompt += f"CORE THEME: {chapter_title}\n\n"
            
            # Include main topics prominently
            if clean_topics:
                topics_list = "\n".join([f"- {topic}" for topic in clean_topics[:8]])
                user_prompt += f"""Main Topics in this Chapter/Unit/Lesson:
{topics_list}

"""
            
            # Include learning objectives prominently
            if clean_objectives:
                objectives_list = "\n".join([f"- {obj}" for obj in clean_objectives[:6]])
                user_prompt += f"""Learning Objectives (what students will achieve):
{objectives_list}

"""
            elif chapter_summary and len(chapter_summary) < 300:
                user_prompt += f"Chapter overview: {chapter_summary[:200]}...\n\n"
        
        user_prompt += """Create a warm, encouraging, and pedagogically sound welcome message that:
1. Start with "Selam!" (do NOT use greetings in other languages like Japanese, Korean, or Chinese)
2. Introduces yourself as YENETA, their AI Tutor
3. Mentions their grade and subject if provided
4. If a chapter is specified:
   - Clearly state the main topics they'll be exploring (mention 2-3 key topics naturally)
   - Mention the learning objectives/what they'll achieve (1-2 key objectives)
   - Make it engaging and relevant to their learning journey
5. Encourages them to ask questions, upload their work, or practice
6. Keep it concise but informative (3-5 sentences)
7. Use a warm, encouraging tone that builds confidence

IMPORTANT: 
- Center the message on the core topics and objectives
- Don't just list topics - weave them naturally into the message
- Make it clear what the student will learn and achieve
- Align with best teaching practices (active learning, growth mindset)

Write ONLY the welcome message, nothing else."""
        
        # Generate welcome message using LLM
        llm_request = LLMRequest(
            prompt=user_prompt,
            system_prompt=system_prompt,
            user_id=request.user.id,
            user_role=UserRole(request.user.role),
            task_type=TaskType.CONTENT_GENERATION,
            complexity=TaskComplexity.BASIC,
            temperature=0.8,  # Slightly higher for more natural, friendly tone
            max_tokens=1000,  # Allow more tokens for detailed welcome message
        )
        
        response = llm_router.process_request(llm_request)
        
        if response.success and response.content:
            welcome_message = response.content.strip()
            # Clean up any extra formatting
            welcome_message = welcome_message.replace('"', '').replace("'", "")
            # Ensure it ends properly
            if not welcome_message.endswith(('.', '!', '?')):
                welcome_message += '.'
            
            return Response({
                'welcome_message': welcome_message,
                'model': str(response.model)
            })
        else:
            # Fallback to default message
            fallback = "Selam! I'm YENETA, your AI Tutor. How can I help you with your studies today? You can ask me questions, upload a photo of your work, or even send a voice message."
            return Response({
                'welcome_message': fallback,
                'model': 'fallback'
            })
    
    except Exception as e:
        logger.error(f"Error generating welcome message: {e}", exc_info=True)
        # Return default message on error
        return Response({
            'welcome_message': "Selam! I'm YENETA, your AI Tutor. How can I help you with your studies today? You can ask me questions, upload a photo of your work, or even send a voice message.",
            'model': 'fallback'
        })


@api_view(['GET', 'POST', 'PATCH'])
@permission_classes([IsAuthenticated])
def tutor_configuration_view(request):
    """
    Get, create, or update student's AI Tutor configuration.
    Extracts chapter metadata when chapter_input is provided.
    """
    from .models import TutorConfiguration
    from .serializers import TutorConfigurationSerializer
    from ai_tools.chapter_utils import normalize_chapter_input
    
    try:
        if request.method == 'GET':
            # Get existing configuration or return defaults
            try:
                config = TutorConfiguration.objects.get(user=request.user)
                serializer = TutorConfigurationSerializer(config)
                return Response(serializer.data)
            except TutorConfiguration.DoesNotExist:
                # Return default configuration
                return Response({
                    'id': None,
                    'use_ethiopian_curriculum': True,
                    'grade': getattr(request.user, 'grade', '') or '',
                    'stream': '',
                    'subject': '',
                    'chapter_input': '',
                    'chapter_number': None,
                    'chapter_title': None,
                    'chapter_topics': [],
                    'chapter_summary': None,
                    'learning_objectives': [],
                    'created_at': None,
                    'updated_at': None,
                    'last_used_at': None,
                })
        
        elif request.method == 'POST':
            # Create or update configuration
            serializer = TutorConfigurationSerializer(
                data=request.data,
                context={'request': request}
            )
            
            if serializer.is_valid():
                config = serializer.save()
                
                # Extract chapter metadata if chapter_input provided
                chapter_input = config.chapter_input
                if chapter_input and config.grade and config.subject:
                    try:
                        # Normalize chapter input
                        normalized = normalize_chapter_input(chapter_input)
                        chapter_number = normalized.get('number')
                        
                        if chapter_number:
                            config.chapter_number = int(chapter_number)
                            
                            # Extract chapter topics and summary from vector store
                            from rag.services import query_curriculum_documents
                            from .tutor_rag_enhancer import TutorRAGEnhancer
                            
                                # Query for chapter content with broader query for better extraction
                            documents = query_curriculum_documents(
                                grade=config.grade,
                                subject=config.subject,
                                query=f"Chapter {chapter_number} content objectives topics",
                                stream=config.stream if config.stream else None,
                                chapter=str(chapter_number),
                                top_k=5,  # Get more results for better extraction
                                extract_full_chapter=True
                            )
                            
                            if documents and len(documents) > 0:
                                # Extract topics and generate summary using enhanced extractor
                                from .chapter_topic_extractor import ChapterTopicExtractor
                                
                                chapter_data = documents[0]
                                content = chapter_data.get('content', '')
                                
                                # Extract comprehensive chapter metadata
                                metadata = ChapterTopicExtractor.extract_chapter_metadata(
                                    content,
                                    chapter_number=chapter_number
                                )
                                
                                # Use extracted metadata
                                config.chapter_title = metadata.get('title') or chapter_data.get('title', f'Chapter {chapter_number}')
                                config.chapter_summary = metadata.get('summary', content[:300] + "..." if len(content) > 300 else content)
                                
                                # Combine objectives and topics for chapter_topics
                                topics_list = []
                                if metadata.get('objectives'):
                                    topics_list.extend(metadata['objectives'][:5])
                                if metadata.get('topics'):
                                    topics_list.extend(metadata['topics'][:5])
                                
                                # Remove duplicates and limit
                                seen = set()
                                unique_topics = []
                                for topic in topics_list:
                                    topic_lower = topic.lower()
                                    if topic_lower not in seen:
                                        seen.add(topic_lower)
                                        unique_topics.append(topic)
                                
                                config.chapter_topics = unique_topics[:10]
                                
                                config.save()
                                logger.info(f"✅ Extracted chapter metadata for {request.user.username}: Chapter {chapter_number} - {len(unique_topics)} topics")
                    except Exception as e:
                        logger.warning(f"⚠️ Could not extract chapter metadata: {e}")
                
                response_serializer = TutorConfigurationSerializer(config)
                return Response(response_serializer.data, status=status.HTTP_201_CREATED)
            
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        elif request.method == 'PATCH':
            # Update existing configuration
            try:
                config = TutorConfiguration.objects.get(user=request.user)
            except TutorConfiguration.DoesNotExist:
                # Create if doesn't exist
                serializer = TutorConfigurationSerializer(
                    data=request.data,
                    context={'request': request}
                )
                if serializer.is_valid():
                    config = serializer.save()
                    response_serializer = TutorConfigurationSerializer(config)
                    return Response(response_serializer.data, status=status.HTTP_201_CREATED)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
            serializer = TutorConfigurationSerializer(
                config,
                data=request.data,
                partial=True,
                context={'request': request}
            )
            
            if serializer.is_valid():
                config = serializer.save()
                
                # Extract chapter metadata if chapter_input changed
                chapter_input = config.chapter_input
                if chapter_input and config.grade and config.subject:
                    try:
                        normalized = normalize_chapter_input(chapter_input)
                        chapter_number = normalized.get('number')
                        
                        if chapter_number and config.chapter_number != int(chapter_number):
                            config.chapter_number = int(chapter_number)
                            
                            # Extract chapter topics and summary
                            from rag.services import query_curriculum_documents
                            from .tutor_rag_enhancer import TutorRAGEnhancer
                            
                            documents = query_curriculum_documents(
                                grade=config.grade,
                                subject=config.subject,
                                query=f"Chapter {chapter_number} topics summary",
                                stream=config.stream if config.stream else None,
                                chapter=str(chapter_number),
                                top_k=3,
                                extract_full_chapter=True
                            )
                            
                            if documents and len(documents) > 0:
                                # Extract topics and generate summary using enhanced extractor
                                from .chapter_topic_extractor import ChapterTopicExtractor
                                
                                chapter_data = documents[0]
                                content = chapter_data.get('content', '')
                                
                                # Extract comprehensive chapter metadata
                                metadata = ChapterTopicExtractor.extract_chapter_metadata(
                                    content,
                                    chapter_number=chapter_number
                                )
                                
                                # Use extracted metadata
                                config.chapter_title = metadata.get('title') or chapter_data.get('title', f'Chapter {chapter_number}')
                                config.chapter_summary = metadata.get('summary', content[:300] + "..." if len(content) > 300 else content)
                                
                                # Combine objectives and topics for chapter_topics
                                topics_list = []
                                if metadata.get('objectives'):
                                    topics_list.extend(metadata['objectives'][:5])
                                if metadata.get('topics'):
                                    topics_list.extend(metadata['topics'][:5])
                                
                                # Remove duplicates and limit
                                seen = set()
                                unique_topics = []
                                for topic in topics_list:
                                    topic_lower = topic.lower()
                                    if topic_lower not in seen:
                                        seen.add(topic_lower)
                                        unique_topics.append(topic)
                                
                                config.chapter_topics = unique_topics[:10]
                                
                                config.save()
                                logger.info(f"✅ Updated chapter metadata for {request.user.username}: {len(unique_topics)} topics")
                    except Exception as e:
                        logger.warning(f"⚠️ Could not update chapter metadata: {e}")
                
                response_serializer = TutorConfigurationSerializer(config)
                return Response(response_serializer.data)
            
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    except Exception as e:
        logger.error(f"Tutor configuration error: {e}", exc_info=True)
        return Response(
            {'error': f'Failed to process tutor configuration: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def chat_view(request):
    """
    General chat endpoint for Virtual Classroom and other interactive features.
    Supports text, images, and whiteboard context.
    """
    message = request.data.get('message', '')
    context = request.data.get('context', '')
    image_data = request.data.get('image')  # {base64, mimeType}
    whiteboard_json = request.data.get('whiteboardJson')
    whiteboard_image = request.data.get('whiteboardImage')
    use_rag = request.data.get('useRag', False)
    use_web_search = request.data.get('useWebSearch', False)
    
    if not message and not image_data:
        return Response({'error': 'Message or image is required'}, status=status.HTTP_400_BAD_REQUEST)

    # Construct prompt with context
    system_prompt = "You are YENETA, a helpful and knowledgeable AI teacher assistant."
    if context:
        system_prompt += f"\n\nContext: {context}"
    
    # Add whiteboard context if available
    if whiteboard_json:
        # In a real implementation, we might parse the JSON to describe the shapes
        system_prompt += "\n\nThe student is working on a whiteboard."
    
    # Prepare LLM request
    llm_request = LLMRequest(
        prompt=message,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType.TUTORING,
        complexity=TaskComplexity.MEDIUM,
        system_prompt=system_prompt,
        use_rag=use_rag,
        # Handle multimodal inputs if supported by the backend LLM router
        # For now, we'll assume the router handles it or we might need to extend LLMRequest
    )
    
    try:
        response = llm_router.process_request(llm_request)
        return Response({
            'text': response.content,
            'functionCalls': [] # Placeholder for function calling if needed
        })
    except Exception as e:
        logger.error(f"Chat error: {e}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_content_view(request):
    """
    Generate or regenerate content for lessons or other educational materials.
    """
    prompt = request.data.get('prompt', '')
    context = request.data.get('context', '')
    task_type = request.data.get('taskType', 'content_generation')
    
    if not prompt:
        return Response({'error': 'Prompt is required'}, status=status.HTTP_400_BAD_REQUEST)

    llm_request = LLMRequest(
        prompt=prompt,
        user_id=request.user.id,
        user_role=UserRole(request.user.role),
        task_type=TaskType(task_type) if task_type in TaskType._value2member_map_ else TaskType.CONTENT_GENERATION,
        system_prompt=f"Context: {context}" if context else None,
    )

    try:
        response = llm_router.process_request(llm_request)
        return Response({'content': response.content})
    except Exception as e:
        logger.error(f"Content generation error: {e}")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def generate_image_view(request):
    """
    Generate images using available image generation models.
    """
    prompt = request.data.get('prompt', '')
    mode = request.data.get('mode', 'standard') # 'standard', 'hd', etc.
    
    if not prompt:
        return Response({'error': 'Prompt is required'}, status=status.HTTP_400_BAD_REQUEST)

    # Placeholder for actual image generation logic
    # In a real system, this would call DALL-E 3 or similar via llm_router or a specialized service
    
    logger.info(f"Generating image for prompt: {prompt} in mode {mode}")
    
    # Mock response for now
    mock_image_url = f"https://via.placeholder.com/1024x1024.png?text={prompt.replace(' ', '+')}"
    
    return Response({'imageUrl': mock_image_url})
